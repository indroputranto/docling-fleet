#!/usr/bin/env python3
"""
Document Coverage Checker — post-extraction completeness validation.

Compares the source document against the chunks produced by the extraction
pipeline to detect silent content loss.  Two complementary strategies:

  1. Heading-based (DOCX only)
     Opens the source file with python-docx and collects every paragraph
     that uses a Word heading style as the "expected section list".  Checks
     how many of those headings are represented (by key terms) in the chunk
     titles.  Adapted from the audit_all_vessels.py approach.

  2. Word count ratio (all file types)
     Compares total words in the raw-extracted chunks (before AI enrichment)
     against total words in the final chunks (after enrichment).  A ratio
     below WORD_COUNT_THRESHOLD flags possible content loss during enrichment.
     A separate comparison of raw-extracted words against a quick re-read of
     the source body text flags possible loss during extraction itself.

Returns a standardised CoverageResult dict stored as JSON on the Document
record, and surfaced on the preview page.
"""

import re
import logging
from typing import List, Dict, Optional

logger = logging.getLogger(__name__)

# If final chunk word count is below this fraction of raw extracted words,
# flag it as a potential enrichment loss.
ENRICHMENT_WORD_RATIO_THRESHOLD = 0.85

# If raw extracted words are below this fraction of source body words,
# flag possible extraction loss.
EXTRACTION_WORD_RATIO_THRESHOLD = 0.80

# A heading must contribute at least this many significant words (>3 chars)
# to be eligible for the heading coverage check.
MIN_HEADING_WORDS = 1


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _word_count(text: str) -> int:
    return len(text.split())


def _significant_words(text: str) -> List[str]:
    """Return lowercase words longer than 3 characters, stripped of numbers."""
    text = re.sub(r'^\d+[\.\)]\s*', '', text)  # strip leading clause numbers
    return [w.lower() for w in re.findall(r'\b[a-zA-Z]{4,}\b', text)]


def _heading_covered(heading: str, chunk_titles: List[str]) -> bool:
    """
    Check whether a source heading is approximately represented in the chunk
    titles.  Uses a liberal word-overlap approach: if ANY significant word
    from the heading appears in ANY chunk title, it counts as covered.
    This handles cases where the AI enrichment rewrites
    "29. Crew Overtime" → "Fixture Recap - Crew Overtime".
    """
    sig_words = _significant_words(heading)
    if not sig_words:
        return True  # nothing meaningful to match, don't penalise
    for title in chunk_titles:
        title_lower = title.lower()
        if any(w in title_lower for w in sig_words):
            return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# DOCX heading extraction  (adapted from audit_all_vessels.py)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx_headings(file_stream) -> Optional[List[str]]:
    """
    Open a .docx stream with python-docx and return all heading-style
    paragraph texts.  Returns None if python-docx is unavailable or the
    file cannot be read.
    """
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_stream)
        headings = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name if p.style else "").lower()
            if "heading" in style_name:
                headings.append(text)
        return headings
    except Exception as e:
        logger.warning(f"[coverage] Could not extract DOCX headings: {e}")
        return None


def _extract_docx_body_words(file_stream) -> Optional[int]:
    """
    Quick raw word count of all paragraph text in a .docx — used as the
    extraction ground truth for word ratio comparison.
    """
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_stream)
        total = sum(_word_count(p.text) for p in doc.paragraphs if p.text.strip())
        return total
    except Exception as e:
        logger.warning(f"[coverage] Could not count DOCX body words: {e}")
        return None


def _extract_pdf_body_words(file_stream) -> Optional[int]:
    """
    Extract raw word count from a PDF using pdfplumber.
    Used as the extraction ground truth for PDF coverage checking.
    """
    try:
        import pdfplumber
        total = 0
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                total += _word_count(text)
        return total if total > 0 else None
    except Exception as e:
        logger.warning(f"[coverage] Could not extract PDF word count: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def run_coverage_check(
    file_stream,
    filename: str,
    raw_chunks_before_enrichment: List[Dict],
    final_chunks: List[Dict],
) -> Dict:
    """
    Run a completeness check on the extraction + enrichment pipeline output.

    Args:
        file_stream:                  Seekable file-like object for the source
                                      document (will be seeked to 0 before use).
        filename:                     Original filename (used to detect .docx).
        raw_chunks_before_enrichment: Chunks produced by extractor.extract()
                                      BEFORE AI enrichment — the extraction
                                      ground truth.
        final_chunks:                 Chunks after AI enrichment (or the same
                                      list if enrichment was skipped).

    Returns a dict with the following keys:
        ok              bool   — True if no issues detected
        coverage_pct    int    — 0-100, headline coverage score
        method          str    — 'heading+wordcount' | 'wordcount'
        issues          list   — human-readable issue strings
        warnings        list   — non-critical observations
        details         dict   — raw numbers for the UI
    """
    result = {
        "ok": True,
        "coverage_pct": 100,
        "method": "wordcount",
        "issues": [],
        "warnings": [],
        "details": {},
    }

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    is_docx = ext == "docx"

    # ── Word count stats ───────────────────────────────────────────────────────
    raw_words   = sum(_word_count(c.get("body", "")) for c in raw_chunks_before_enrichment)
    final_words = sum(_word_count(c.get("body", "")) for c in final_chunks)

    result["details"]["raw_extracted_words"] = raw_words
    result["details"]["final_chunk_words"]   = final_words
    result["details"]["chunk_count"]         = len(final_chunks)

    # Enrichment word ratio
    if raw_words > 0:
        enrichment_ratio = final_words / raw_words
        result["details"]["enrichment_word_ratio"] = round(enrichment_ratio, 3)
        if enrichment_ratio < ENRICHMENT_WORD_RATIO_THRESHOLD:
            pct_lost = int((1 - enrichment_ratio) * 100)
            result["issues"].append(
                f"AI enrichment may have dropped content — "
                f"{pct_lost}% word count reduction "
                f"({raw_words} → {final_words} words)."
            )
            result["ok"] = False

    # ── DOCX-specific: heading check + source body word count ──────────────────
    if is_docx:
        result["method"] = "heading+wordcount"

        try:
            file_stream.seek(0)
        except Exception:
            pass

        source_words = _extract_docx_body_words(file_stream)
        if source_words is not None:
            result["details"]["source_body_words"] = source_words
            if raw_words > 0 and source_words > 0:
                extraction_ratio = raw_words / source_words
                result["details"]["extraction_word_ratio"] = round(extraction_ratio, 3)
                if extraction_ratio < EXTRACTION_WORD_RATIO_THRESHOLD:
                    pct_lost = int((1 - extraction_ratio) * 100)
                    result["issues"].append(
                        f"Extractor may have missed content — "
                        f"only {100 - pct_lost}% of source body text captured "
                        f"({source_words} source words → {raw_words} extracted)."
                    )
                    result["ok"] = False

        # Heading-based structural check
        try:
            file_stream.seek(0)
        except Exception:
            pass

        headings = _extract_docx_headings(file_stream)
        if headings:
            result["details"]["source_heading_count"] = len(headings)
            chunk_titles = [c.get("title", "") for c in final_chunks if c.get("title")]
            result["details"]["chunk_title_count"] = len(chunk_titles)

            missed = [h for h in headings if not _heading_covered(h, chunk_titles)]
            covered_count = len(headings) - len(missed)
            heading_pct = int(covered_count / len(headings) * 100) if headings else 100

            result["details"]["heading_coverage_pct"] = heading_pct
            result["details"]["missed_headings"]      = missed[:20]  # cap for storage

            if missed:
                result["warnings"].append(
                    f"{len(missed)} source heading(s) may not be represented in "
                    f"chunk titles — review the missed headings list on the preview page."
                )
                # Only flag as an issue if more than 20% of headings are missing
                if heading_pct < 80:
                    result["ok"] = False

            # Use heading coverage as the primary pct for docx
            result["coverage_pct"] = heading_pct
        else:
            # No heading styles found — fall back to word count ratio as pct
            if raw_words > 0 and source_words:
                result["coverage_pct"] = min(
                    100,
                    int(raw_words / source_words * 100)
                )
            result["details"]["source_heading_count"] = 0
            result["warnings"].append(
                "No heading styles detected in source document — "
                "structural coverage check unavailable; word count ratio used instead."
            )
    elif ext == "pdf":
        result["method"] = "wordcount+pdf"

        try:
            file_stream.seek(0)
        except Exception:
            pass

        source_words = _extract_pdf_body_words(file_stream)
        if source_words is not None:
            result["details"]["source_body_words"] = source_words
            if raw_words > 0 and source_words > 0:
                extraction_ratio = raw_words / source_words
                result["details"]["extraction_word_ratio"] = round(extraction_ratio, 3)
                if extraction_ratio < EXTRACTION_WORD_RATIO_THRESHOLD:
                    pct_lost = int((1 - extraction_ratio) * 100)
                    result["issues"].append(
                        f"Extractor may have missed content — "
                        f"only {100 - pct_lost}% of source text captured "
                        f"({source_words} source words → {raw_words} extracted). "
                        f"The PDF may contain images, tables, or scanned pages."
                    )
                    result["ok"] = False
                # Use extraction ratio as primary pct
                result["coverage_pct"] = min(100, int(extraction_ratio * 100))
            elif source_words > 0:
                result["details"]["source_body_words"] = source_words
        else:
            # pdfplumber couldn't extract — warn that audit was limited
            result["warnings"].append(
                "Could not read source PDF for extraction check — "
                "word count audit unavailable (possibly a scanned/image PDF)."
            )

    else:
        # Other formats (xlsx, etc.): enrichment ratio only
        if raw_words > 0:
            result["coverage_pct"] = min(
                100,
                int(final_words / raw_words * 100)
            )

    if result["issues"] or result["warnings"]:
        logger.info(
            f"[coverage] '{filename}': pct={result['coverage_pct']}% "
            f"issues={len(result['issues'])} warnings={len(result['warnings'])}"
        )

    return result
