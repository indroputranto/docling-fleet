#!/usr/bin/env python3
"""
Document text extraction for the upload pipeline.

Supports:
  .docx — python-docx; clause/heading-aware chunking
  .pdf  — PyMuPDF (font-aware) + pdfplumber fallback; section-aware chunking
  .xlsx — openpyxl; one chunk per sheet

Each extractor returns a list of dicts:
  [{"title": str | None, "body": str}, ...]

Chunking strategy for maritime documents:
  PDFs:  1. Use PyMuPDF get_text("rawdict") for text + per-glyph boxes and
             get_text("dict")-compatible span metadata (font size, bold flag),
             — lines visually larger or bold → section header boundary
             — regex-based clause detection as secondary signal
         2. Filter "junk" chunks (cargo diagrams, slot-plan labels)
         3. Fallback: pdfplumber + regex-only chunking
         4. Last resort: fixed-size chunks of MAX_CHUNK_WORDS words
  DOCX:  Heading styles + clause regex
"""

import re
import unicodedata
import logging
from typing import IO, List, Dict, Optional, Tuple

logger = logging.getLogger(__name__)

MAX_CHUNK_WORDS = 1500  # soft word limit per chunk before splitting
                        # Raised from 600: charter-party clauses can be 800–1200 words
                        # and must stay unified for quality Pinecone indexing.
                        # 1500 words ≈ ~2000 tokens, well within ada-002's 8191-token limit.


# ─────────────────────────────────────────────────────────────────────────────
# PDF text cleaning — ligature & encoding artifact repair
# ─────────────────────────────────────────────────────────────────────────────

# BIMCO SmartCon and similar commercial charter-party PDFs embed proprietary
# fonts whose ToUnicode CMap tables incorrectly map common ligature glyphs to
# Latin Extended codepoints instead of the correct ASCII sequences.
#
# Observed mappings (confirmed on BIMCO SmartCon GENCON 2022):
#   U+019F  Ɵ  → "ti"   (LATIN CAPITAL LETTER O WITH MIDDLE TILDE used for ti-ligature)
#   U+014C  Ō  → "ft"   (LATIN CAPITAL LETTER O WITH MACRON used for ft-ligature)
#   U+01A9  Ʃ  → "tt"   (LATIN CAPITAL LETTER ESH used for tt-ligature)
#
# Add further entries here if new artifacts are discovered in other PDF sources.
_BIMCO_LIGATURE_MAP: dict = {
    '\u019F': 'ti',   # Ɵ → ti
    '\u014C': 'ft',   # Ō → ft
    '\u01A9': 'tt',   # Ʃ → tt
}

# Build a str.translate()-compatible table (codepoint → replacement string)
_PDF_TRANSLATE_TABLE = str.maketrans(_BIMCO_LIGATURE_MAP)


def _clean_pdf_text(text: str) -> str:
    """
    Repair encoding artifacts from PDFs with broken ToUnicode CMap entries.

    Three-pass strategy:
      1. NFKC normalization — decomposes standard Unicode ligatures
         (ﬁ→fi, ﬂ→fl, ﬀ→ff, ﬃ→ffi, ﬄ→ffl, ﬅ/ﬆ→st) into ASCII pairs.
      2. Character substitution — corrects BIMCO SmartCon font mis-encodings
         where ligature glyphs are mapped to wrong Latin Extended codepoints.
      3. Ligature space-join — fixes cases where PyMuPDF renders a ligature
         glyph as two separate span fragments with a space between them.
         Observed in BIMCO SmartCon PDFs:
           "Dura ti on"  → "Duration"
           "condi ti on" → "condition"
           "par ti es"   → "parties"
           "cer ti fy"   → "certify"
           "no ti ce"    → "notice"
           "addi ti onal"→ "additional"
         Also handles "ft" and "tt" ligature splits ("draf t ing" → "drafting").
    """
    # Pass 1: standard Unicode ligatures (U+FB00–U+FB06)
    text = unicodedata.normalize('NFKC', text)
    # Pass 2: BIMCO-specific wrong-codepoint mappings
    text = text.translate(_PDF_TRANSLATE_TABLE)
    # Pass 3a: join word fragments split by a FLOATING ligature substring
    # Handles ligature as middle fragment:  "no ti ce" → "notice"
    # Also handles ligature as a standalone word at token start:
    #   "  ti me" → "time",  "fi nal" → "final"
    text = re.sub(r'(\w) (ti|ft|tt|ffi|ffl|fi|fl) (\w)', r'\1\2\3', text)
    # Standalone ligature fragment followed by a word (word-boundary anchor)
    text = re.sub(r'\b(ti|tt|ffi|ffl)\s+([a-z])', r'\1\2', text)

    # Pass 3b: join PREFIX + merged-suffix splits where fitz absorbed the
    # ligature into the suffix but left a space before it.
    # e.g. "Dura tion" → "Duration", "Instruc tions" → "Instructions",
    #      "Communica tions" → "Communications", "obliga tion" → "obligation"
    # The suffix list covers the most common 'ti'-ligature-affected endings.
    _TI_SUFFIXES = (
        r"tions?|tives?|tively|tings?|tional|tionally|"
        r"tified|tifying|tifiable|tification|tifications|"
        r"tities|tity|tion\b"
    )
    text = re.sub(
        rf'([A-Za-z]{{2,}}) ({_TI_SUFFIXES})',
        lambda m: m.group(1) + m.group(2),
        text,
    )
    return text


# ─────────────────────────────────────────────────────────────────────────────
# Clause / section header detection  (regex fallback for contract PDFs)
# ─────────────────────────────────────────────────────────────────────────────

# Matches common maritime clause patterns:
#   "CLAUSE 1", "Clause 14", "1. DEFINITIONS", "14 — LIEN", "PART II", "ANNEX A"
#
# NOTE: The numbered pattern requires a SPACE then a LETTER after the separator
# so that decimal numbers (6.438, 16.0m, 3.5 mt/m2) are NOT mistaken for
# clause headers.  "1." alone (no following word) is intentionally excluded
# because in vessel description PDFs single numbers ending with "." are common
# data values, not structural headers.
_CLAUSE_RE = re.compile(
    r"^\s*(?:"
    r"(?:CLAUSE|Clause)\s+\d+"              # CLAUSE 1 / Clause 14
    r"|(?:PART|Part)\s+(?:[IVX]+|\d+)"      # PART II / Part 3
    r"|(?:ANNEX|Annex|APPENDIX|Appendix)\s+\S+"  # ANNEX A / Appendix B
    r"|\d{1,3}\s*[.\-—]\s+[A-Za-z]"        # "1. Definitions" / "14 — Lien"
    r")",                                   # NOTE: space required before letter
    re.MULTILINE,
)

# Matches a left-margin line number that was merged into a content span by
# PyMuPDF when the number and content happen to be in the same text line object.
# Pattern: 1–3 digits at the very start, followed by whitespace and a letter.
# Upper bound of 3 digits (≤999) safely avoids stripping genuine numeric content
# like "12580 MT" (5 digits) or "9611 GT" (4 digits).
# The positive lookahead (?=[A-Za-z]) ensures we never strip a number that is
# itself the start of the content (e.g. "1.5 knots" would NOT match because
# the next char after "1" is "." not a space; "12580 metric" won't match —
# 5 digits exceeds \d{1,3}).
_LN_PREFIX_RE = re.compile(r'^\d{1,3}\s+(?=[A-Za-z])')

# When the last seen top-level clause is e.g. 9 but a line begins "47. The …",
# it is almost always a **merged right-margin line number** (47) plus the next
# sentence — not charter-party clause 47.  Real consecutive clauses rarely jump
# by this much; when they do, titles are capitalised ("47. Liens"), not prose.
_CLAUSE_JUMP_RUNON_MIN_GAP = 20


def _is_inline_clause_citation_sentence(line: str) -> bool:
    """
    True when a line *opens* with ``Clause N`` but continues as a sentence
    (reference to another clause / rider text), not as a standalone title.

    Examples that MUST NOT be headers:
        "Clause 47. The Owners reserve…"
        "Clause 12 the charterers shall…"   (unlikely typography; guarded anyway)

    Standalone titles like ``Clause 14`` or ``Clause 14 — Hire`` are False here.
    """
    s = line.strip()
    if not re.match(r"(?i)^(?:clause|CLAUSE)\s+\d+", s):
        return False
    # Standalone clause label (possibly with a trailing period only)
    if re.match(r"(?i)^(?:clause|CLAUSE)\s+\d+\.?\s*$", s):
        return False
    # Dash / em-dash title on the same line: "Clause 14 — Hire Payment"
    if re.match(r"(?i)^(?:clause|CLAUSE)\s+\d+\s*[—\-–]\s*\S", s):
        return False
    # Sentence continues after "Clause N" / "Clause N." — word after "the" starts
    # lowercase → body prose (PDF may still use "The" when it should be "the").
    m = re.search(r"(?i)(?:clause|CLAUSE)\s+\d+\.?\s+the\s+(\w)", s)
    if m and m.group(1).islower():
        return True
    return False


def _looks_like_merged_margin_clause_number(
    clause_num: int,
    max_clause_seen: int,
    title_rest: str,
) -> bool:
    """
    Detect ``N. <prose>`` that is likely a margin line number pasted onto the
    start of the next sentence, not a real top-level clause heading.
    """
    if max_clause_seen <= 0:
        return False
    if clause_num <= max_clause_seen:
        return False
    gap = clause_num - max_clause_seen
    if gap < _CLAUSE_JUMP_RUNON_MIN_GAP:
        return False
    rest = (title_rest or "").strip()
    if not rest:
        return False

    # Sentence-style openers after "N. " (merged right-margin line number + body).
    m_the = re.match(r"(?i)^the\s+(\w+)", rest)
    if m_the:
        w = m_the.group(1)
        if w[0].islower():
            return True
        # "The Owners reserve…" — title-case but long clause jump + long line → body
        if gap >= 30 and len(rest) > 50:
            return True
        return False

    m_kw = re.match(
        r"(?i)^(if|where|when)\s+(\w)",
        rest,
    )
    if m_kw and m_kw.group(2)[0].islower():
        return True

    if re.match(r"(?i)^in order to\s+\w", rest):
        return True
    if re.match(r"(?i)^in the event that\s+\w", rest):
        return True

    return False


def _is_clause_header(text: str) -> bool:
    s = text.strip()
    if _is_inline_clause_citation_sentence(s):
        return False
    return bool(_CLAUSE_RE.match(s))


# ─────────────────────────────────────────────────────────────────────────────
# DOCX run-level formatting — strikethrough preservation
# ─────────────────────────────────────────────────────────────────────────────

def _paragraph_text_with_strikethrough(para) -> str:
    """
    Return the text of a python-docx Paragraph with strikethrough runs
    wrapped in ~~double-tilde~~ markdown notation.

    Checks both <w:strike> (single) and <w:dstrike> (double) elements on each
    run's rPr, matching the behaviour of process_vessel_new.py.  Plain runs
    are returned as-is so the result is drop-in compatible with para.text.
    """
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    result_parts: List[str] = []
    strike_buf:   List[str] = []
    in_strike = False

    for run in para.runs:
        text = run.text
        if not text:
            continue

        is_struck = False
        rPr = getattr(run._element, "rPr", None)
        if rPr is not None:
            is_struck = (
                rPr.find(f"{{{W_NS}}}strike")  is not None or
                rPr.find(f"{{{W_NS}}}dstrike") is not None
            )

        if is_struck:
            if not in_strike:
                in_strike = True
            strike_buf.append(text)
        else:
            if in_strike:
                result_parts.append(f"~~{''.join(strike_buf)}~~")
                strike_buf = []
                in_strike = False
            result_parts.append(text)

    # Flush any trailing strikethrough
    if strike_buf:
        result_parts.append(f"~~{''.join(strike_buf)}~~")

    return "".join(result_parts)


def _is_docx_subsection_label(text: str) -> bool:
    """
    Detect paragraph-style subsection labels used in vessel-description DOCX
    files where section headers carry no heading style but are identifiable
    because they end with ':' and have no value after the colon.

    Examples that match:
        "General Information:"  "Tonnage:"  "Propulsion & Maneuvering:"
        "Hold and Hatch Sizes:"  "RoRo Features:"  "Container Capacity:"

    Examples that do NOT match:
        "Call Sign: PEVT"          (has value after ':')
        "DWAT (closed/open): 4540" (has value after ':')
        "Grain fitted"             (does not end with ':')
        "conditions:"              (starts lowercase → likely a sentence fragment)
        "This is a very long sentence that ends with a colon:"  (>80 chars)
    """
    stripped = text.strip()
    if not stripped or not stripped.endswith(':'):
        return False
    # Must be short — real labels are terse
    if len(stripped) > 80:
        return False
    # Must start with an uppercase letter (section titles, not sentence fragments)
    if stripped[0].islower():
        return False
    # Must not contain ': ' followed by content (that would be a key-value pair)
    if ': ' in stripped[:-1]:   # ignore the final ':'
        return False
    return True


def _split_into_chunks(segments: List[Dict], max_words: int = MAX_CHUNK_WORDS) -> List[Dict]:
    """
    Given a list of {title, body} where body may be very long,
    split any chunk exceeding max_words into sub-chunks.

    Args:
        segments:  list of {title, body} dicts from an extraction pass.
        max_words: soft word ceiling per chunk.  Defaults to MAX_CHUNK_WORDS
                   (1500) for general documents.  Pass a higher value (e.g.
                   5000) for clause-presplit paths where splitting mid-clause
                   is worse than a large chunk — the hard ceiling is the
                   embedding model's token limit (~6 000 words for
                   text-embedding-3-small at 8 191 tokens).
    """
    result = []
    for seg in segments:
        words = seg["body"].split()
        if len(words) <= max_words:
            result.append(seg)
        else:
            # Split into sub-chunks; carry the original title on the first only
            parts = [
                words[i : i + max_words]
                for i in range(0, len(words), max_words)
            ]
            for idx, part in enumerate(parts):
                suffix = f" (part {idx + 1})" if len(parts) > 1 else ""
                result.append({
                    "title": (seg["title"] or "") + suffix if idx == 0 else f"{seg['title'] or 'Continued'} (part {idx + 1})",
                    "body":  " ".join(part),
                })
    return result


# ─────────────────────────────────────────────────────────────────────────────
# DOCX extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx(file_obj: IO[bytes]) -> List[Dict]:
    """
    Extract chunks from a .docx file.

    Strategy:
    - Iterate paragraphs in document order.
    - A paragraph is a chunk boundary if it uses a Heading style OR
      matches the clause header pattern.
    - Table cells are appended as plain text after the paragraph that
      precedes them.
    """
    try:
        import docx as python_docx
    except ImportError:
        from docx import Document as _D
        python_docx = type("_m", (), {"Document": _D})()

    from docx import Document as DocxDocument
    doc = DocxDocument(file_obj)

    chunks: List[Dict] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []

    def _flush():
        body = "\n".join(current_lines).strip()
        if body:
            chunks.append({"title": current_title, "body": body})

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]  # "p" or "tbl"

        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(block, doc)
            # Use run-level extraction to preserve ~~strikethrough~~ markers.
            # Fall back to para.text for paragraphs without run structure.
            rich_text = _paragraph_text_with_strikethrough(para)
            text = rich_text.strip()
            if not text:
                continue

            # For heading detection, compare against plain text (no ~~ noise)
            plain_text = para.text.strip()
            style_name = (para.style.name or "").lower()
            is_heading = (
                "heading" in style_name
                or _is_clause_header(plain_text)
                or _is_docx_subsection_label(plain_text)
            )

            if is_heading:
                _flush()
                # Headings are stored as plain text — strike markers on a
                # heading title would be confusing and hurt AI enrichment.
                current_title = plain_text
                current_lines = []
            else:
                current_lines.append(text)

        elif tag == "tbl":
            # Extract table cells preserving strikethrough formatting
            from docx.table import Table
            try:
                tbl = Table(block, doc)
                rows = []
                for row in tbl.rows:
                    cells = []
                    for cell in row.cells:
                        # Join paragraphs within each cell, preserving strike
                        cell_text = " ".join(
                            _paragraph_text_with_strikethrough(p).strip()
                            for p in cell.paragraphs
                            if p.text.strip()
                        )
                        cells.append(cell_text)
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    current_lines.append("\n".join(rows))
            except Exception:
                pass

    _flush()

    if not chunks:
        # Fallback: whole document as one chunk
        full_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        chunks = [{"title": None, "body": full_text}]

    return _split_into_chunks(chunks)


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — helpers
# ─────────────────────────────────────────────────────────────────────────────

def _is_junk_body(body: str) -> bool:
    """
    Return True when the chunk body is extraction noise rather than useful text.

    Catches:
    - Cargo slot diagrams / hold plans: position labels like A B C … and
      bay/row numbers 1 2 3 … extracted as individual lines from graphical
      vector drawings.  Heuristic: >60 % of non-empty lines are ≤2 chars.
    - Trivially short chunks (fewer than 5 content words total).
    """
    lines = [l.strip() for l in body.split("\n") if l.strip()]
    if not lines:
        return True
    total_words = sum(len(l.split()) for l in lines)
    if total_words < 3:
        return True
    short = sum(1 for l in lines if len(l) <= 2)
    return short / len(lines) > 0.60


def _chunk_lines(
    lines_meta: List[Dict],
    body_size: float,
) -> List[Dict]:
    """
    Walk a list of {text, size, is_bold} line dicts and group them into
    {title, body} chunks using font metadata + clause-regex signals.

    Header signals (any one triggers a new chunk):
      1. Font size >= body_size * 1.12  (visually larger than body text)
      2. All spans bold + line is short (<120 chars) + doesn't start with digit
         (short bold lines are captions/labels; long bold lines are body prose)
      3. Line matches _CLAUSE_RE (numbered clauses in contract PDFs)
    """
    chunks: List[Dict] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []

    header_threshold = body_size * 1.12

    def _flush():
        body = "\n".join(current_lines).strip()
        if body and not _is_junk_body(body):
            chunks.append({"title": current_title, "body": body})

    for lm in lines_meta:
        text = lm["text"]
        size = lm["size"]
        is_bold = lm["is_bold"]

        is_header = (
            size >= header_threshold
            or (is_bold and len(text) < 120 and not text[0].isdigit())
            or _is_clause_header(text)
        )

        if is_header:
            _flush()
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

    _flush()
    return chunks


def _detect_column_split(elements: List[Dict], page_width: float) -> Optional[float]:
    """
    Detect the x-coordinate of a vertical column gap in a two-column PDF layout.

    Algorithm (bin-density):
      1. Divide the page width into 20pt bins and count elements per bin.
      2. Restrict search to the central 15–80% of page width (skip margins).
      3. Find the longest contiguous run of "sparse" bins (≤2% of total elements).
      4. A run of ≥3 bins (≥60pt) qualifies as a column gap.
      5. Return the midpoint of that gap; None if no gap found.

    Robust to isolated cross-column titles (e.g. "Ship's particulars" centred
    over the gap) because a single element is counted as sparse (≤2%).
    """
    if not elements:
        return None

    BIN = 20.0
    n_bins = int(page_width / BIN) + 2
    counts = [0] * n_bins
    for e in elements:
        b = int(e["x"] / BIN)
        if 0 <= b < n_bins:
            counts[b] += 1

    lo_b = int(page_width * 0.15 / BIN)
    hi_b = int(page_width * 0.80 / BIN)
    sparse_threshold = max(1, len(elements) * 0.02)

    best_len, best_start = 0, None
    cur_len, cur_start = 0, None

    for i in range(lo_b, hi_b + 1):
        if counts[i] <= sparse_threshold:
            cur_len += 1
            if cur_start is None:
                cur_start = i
            if cur_len > best_len:
                best_len, best_start = cur_len, cur_start
        else:
            cur_len, cur_start = 0, None

    # Require at least 3 bins = 60pt gap
    if best_len >= 3 and best_start is not None:
        gap_left = best_start * BIN
        gap_right = (best_start + best_len) * BIN
        return (gap_left + gap_right) / 2.0

    return None


def _clause_number(text: str) -> Optional[int]:
    """
    If `text` is a numbered clause header like "14. Hire Payment" or
    "14 — Hire Payment", return the clause number as an int.
    Returns None for non-numbered headers (PART II, ANNEX A, bold section
    labels, etc.) so those can always create a new chunk.
    """
    m = re.match(r"^\s*(\d{1,3})\s*[.\-—]", text.strip())
    return int(m.group(1)) if m else None


def _column_to_chunks(elements: List[Dict], snap: float = 3.0) -> List[Dict]:
    """
    Convert a list of positioned text elements into {title, body} chunks.

    Works for both single-column and per-column element sets.

    Steps:
      1. Sort elements by (y, x) — top-to-bottom, left-to-right.
      2. Group elements whose y positions are within `snap` pt into the same
         visual row (handles sub-pixel baseline differences).
      3. Determine body font size (modal rounded size).
      4. For each row, decide if it is a section header:
           - All spans bold  AND  line < 120 chars  AND  starts with uppercase
             letter  →  header (handles OCEAN7-style spec sheets)
           - Max font size ≥ body_size * 1.12  →  header (font-size-based)
           - Line matches clause regex  →  header (BIMCO/contract PDFs)
           - Monotonic clause guard: if a numbered clause header has a number
             LOWER than the highest clause number seen so far, it is a list
             item within the current clause body, not a new top-level clause.
             e.g. "2. the Vessel shall..." appearing inside clause 9 is a list
             item, not the re-appearance of clause 2.
      5. Flush completed chunks, filter junk bodies.
    """
    from collections import Counter

    if not elements:
        return []

    # Step 1 & 2: sort and group into visual rows
    sorted_els = sorted(elements, key=lambda e: (e["y"], e["x"]))
    rows: List[List[Dict]] = []
    current_row: List[Dict] = []
    row_y: Optional[float] = None

    for el in sorted_els:
        if row_y is None or abs(el["y"] - row_y) <= snap:
            current_row.append(el)
            if row_y is None:
                row_y = el["y"]
        else:
            if current_row:
                rows.append(current_row)
            current_row = [el]
            row_y = el["y"]
    if current_row:
        rows.append(current_row)

    # Step 3: modal body font size
    size_counts: Counter = Counter(
        round(e["size"] * 2) / 2 for e in elements
    )
    body_size = size_counts.most_common(1)[0][0] if size_counts else 10.0
    header_threshold = body_size * 1.12

    # Step 4 & 5: build chunks
    chunks: List[Dict] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []
    max_clause_seen: int = 0   # monotonic guard for numbered clauses

    def _flush_col():
        body = "\n".join(current_lines).strip()
        if body and not _is_junk_body(body):
            chunks.append({"title": current_title, "body": body})

    for row in rows:
        txt = _clean_pdf_text(" ".join(e["text"] for e in row)).strip()
        if not txt:
            continue

        if _is_inline_clause_citation_sentence(txt):
            current_lines.append(txt)
            continue

        is_bold = all(e["is_bold"] for e in row)
        max_size = max(e["size"] for e in row)
        first_char = txt[0] if txt else ""

        is_header = (
            (is_bold and len(txt) < 120 and first_char.isalpha() and first_char.isupper())
            or max_size >= header_threshold
            or _is_clause_header(txt)
        )

        # Monotonic clause guard: reject numbered headers whose number is
        # lower than the highest clause number seen so far — they are list
        # items inside the current clause, not new top-level clauses.
        if is_header and _is_clause_header(txt):
            num = _clause_number(txt)
            if num is not None:
                if num <= max_clause_seen:
                    # Back-reference: treat as body text, not a new header
                    is_header = False
                else:
                    rm = re.match(
                        r"^\s*\d{1,3}\s*[.\-—]\s*(.+)$", txt.strip()
                    )
                    rest = rm.group(1) if rm else ""
                    if _looks_like_merged_margin_clause_number(
                        num, max_clause_seen, rest
                    ):
                        is_header = False
                    else:
                        max_clause_seen = num

        if is_header:
            _flush_col()
            current_title = txt
            current_lines = []
        else:
            current_lines.append(txt)

    _flush_col()
    return chunks


def _fallback_fixed_chunks(text_lines: List[str]) -> List[Dict]:
    """Last-resort: split plain text into fixed-size word chunks."""
    words = " ".join(text_lines).split()
    chunks = []
    for i in range(0, len(words), MAX_CHUNK_WORDS):
        part = " ".join(words[i : i + MAX_CHUNK_WORDS])
        chunk_num = i // MAX_CHUNK_WORDS + 1
        chunks.append({"title": f"Section {chunk_num}", "body": part})
    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Left-margin line-number detection & stripping
# ─────────────────────────────────────────────────────────────────────────────

def _detect_margin_line_numbers(
    elements: List[Dict],
    *,
    margin_side: str = "left",
) -> Optional[Dict]:
    """
    Detect sequential margin line numbers in charter-party style PDFs (PyMuPDF path).

    Many BIMCO / NYPE forms print integers (1, 2, 3 …) in the **left** or
    **right** margin. PyMuPDF captures these as ordinary text elements.

    *margin_side* ``\"left\"``: cluster in left 25 % of text width (original behaviour).
    *margin_side* ``\"right\"``: cluster in right 25 % — same numeric heuristics,
    mirrored geometry (e.g. Ocean7 working-copy layout).

    Returns a dict with ``strip_ids``, ``x_cluster``, or None.
    """
    from collections import Counter

    margin_side = (margin_side or "left").lower()
    if margin_side not in ("left", "right"):
        return None

    if not elements:
        return None

    # Step 1: candidate elements — bare 1-3 digit integers only
    numeric_els = [
        e for e in elements
        if re.fullmatch(r'\d{1,3}', e['text'].strip())
    ]
    if len(numeric_els) < 10:
        return None

    # Step 2: find modal x-position bin (15 pt resolution)
    BIN = 15.0
    x_bins: Counter = Counter(int(e['x'] / BIN) for e in numeric_els)
    modal_bin, modal_count = x_bins.most_common(1)[0]
    if modal_count < 10:
        return None

    cluster = [e for e in numeric_els if int(e['x'] / BIN) == modal_bin]
    cluster_x = modal_bin * BIN

    # Step 3: cluster must sit in the outer quarter on the chosen side
    all_x = [e['x'] for e in elements]
    x_min, x_max = min(all_x), max(all_x)
    text_width = x_max - x_min
    if text_width > 0:
        rel_pos = (cluster_x - x_min) / text_width
        if margin_side == "left":
            if rel_pos > 0.25:
                return None
        else:
            if rel_pos < 0.75:
                return None

    # Step 4: numbers must cover a meaningful range
    nums = sorted(int(e['text'].strip()) for e in cluster)
    span = nums[-1] - nums[0]
    if span < 15:
        return None

    # Step 5: density — at least 50 % of integers in the range are present
    density = len(set(nums)) / (span + 1)
    if density < 0.50:
        return None

    # Step 6: monotonicity — when read top-to-bottom the numbers should
    # be non-decreasing with at most 20 % inversions
    cluster_by_y = sorted(cluster, key=lambda e: e['y'])
    seq = [int(e['text'].strip()) for e in cluster_by_y]
    if len(seq) > 1:
        inversions = sum(1 for i in range(len(seq) - 1) if seq[i] > seq[i + 1])
        if inversions / (len(seq) - 1) > 0.20:
            return None

    logger.info(
        "[extractor] %s-margin line numbers detected: "
        "x≈%.0fpt, range %d–%d, %d elements to strip",
        margin_side, cluster_x, nums[0], nums[-1], len(cluster),
    )
    return {
        'strip_ids': {id(e) for e in cluster},
        'x_cluster': cluster_x,
    }


def _try_strip_left_prefix_line_numbers(lines: List[str]) -> Optional[List[str]]:
    """Apply leading-integer line stripping; return new lines or None if not applicable."""
    _prefix = re.compile(r'^(\d{1,3})(\s+.*|)$')

    candidate_nums: list = []
    for line in lines:
        m = _prefix.match(line.strip())
        if m:
            candidate_nums.append(int(m.group(1)))

    non_empty = [l for l in lines if l.strip()]
    if len(candidate_nums) < 10:
        return None
    if not non_empty or len(candidate_nums) / len(non_empty) < 0.25:
        return None

    nums = sorted(candidate_nums)
    span = nums[-1] - nums[0]
    if span < 15:
        return None
    if len(set(candidate_nums)) / (span + 1) < 0.40:
        return None

    logger.info(
        "[extractor] pdfplumber: stripping left-margin line numbers "
        "(range %d–%d, %d/%d lines)",
        nums[0], nums[-1], len(candidate_nums), len(non_empty),
    )

    _prefix_c = re.compile(r'^(\d{1,3})(\s+.*|)$')
    cleaned: list = []
    for line in lines:
        stripped = line.strip()
        m = _prefix_c.match(stripped)
        if m:
            remainder = m.group(2).strip()
            cleaned.append(remainder)
        else:
            cleaned.append(line)
    return cleaned


def _strip_bare_sequential_line_rows(lines: List[str]) -> Optional[List[str]]:
    """
    Remove rows that are only a 1-4 digit integer when they form a running
    line-number column (typical of **right** margin: each number on its own
    extracted line). These fail the left-prefix stripper's coverage ratio when
    most lines are normal prose.

    Two acceptance modes (either qualifies):

    Strict (clean digital PDFs): strong density (≥40 % of integers in span),
    very few inversions, ≥10 hits.

    Lenient (OCR'd / scanned PDFs): on a Toshiba MFP scan only ~1 in 3 line
    numbers survive OCR, so density-on-span will land near 0.20–0.35. Accept
    when there are still ≥30 hits, the values are strictly monotonic
    (inversions ≤ 5 %), and the median gap between consecutive values is
    small (≤ 8) — i.e. these really are sequential margin numbers, not
    coincidental data values.
    """
    bare_indices: List[int] = []
    bare_values: List[int] = []
    for i, line in enumerate(lines):
        s = line.strip()
        # Accept up to 4 digits so we don't silently miss line 1000+ in long
        # negotiated charter parties (this PDF runs to line 1131).
        if re.fullmatch(r'\d{1,4}', s):
            bare_indices.append(i)
            bare_values.append(int(s))
    if len(bare_values) < 10:
        return None

    span = max(bare_values) - min(bare_values)
    if span < 15:
        return None

    density = len(set(bare_values)) / (span + 1)
    if len(bare_values) > 1:
        inversions_frac = sum(
            1 for i in range(len(bare_values) - 1) if bare_values[i] > bare_values[i + 1]
        ) / (len(bare_values) - 1)
    else:
        inversions_frac = 0.0

    sorted_vals = sorted(bare_values)
    gaps = [
        sorted_vals[i + 1] - sorted_vals[i] for i in range(len(sorted_vals) - 1)
    ]
    median_gap = sorted(gaps)[len(gaps) // 2] if gaps else 0

    accept_strict = density >= 0.40 and inversions_frac <= 0.20
    accept_lenient = (
        len(bare_values) >= 30
        and density >= 0.18
        and inversions_frac <= 0.05
        and median_gap <= 8
    )
    if not (accept_strict or accept_lenient):
        return None

    logger.info(
        "[extractor] Stripping bare line-number rows (margin column, often right): "
        "range %d–%d, %d rows removed (density=%.2f, median_gap=%d, mode=%s)",
        min(bare_values), max(bare_values), len(bare_indices),
        density, median_gap,
        "strict" if accept_strict else "lenient",
    )
    skip = set(bare_indices)
    return [line for i, line in enumerate(lines) if i not in skip]


# Lines repeated on every page of many BIMCO / owner “working copy” charter PDFs
# (footer band is often ≥8pt so MIN_FONT filtering in fitz-presplit does not remove it).
#
# We accept three vendor formats:
#   - SmartCon working copy:    "CP ID: 12345", "CP Date: …", "Vessel: …"
#   - Chinsay working copy:     "Chinsay ID: 139705"  (and the OCR-mangled
#                               "ChinsaylD: 139705" variant where l→I)
#   - Page-of footers:          "Page 1 of 35"  AND  the trailing-text variant
#                               "Page 1 of 35 - MORGENSTOND 1-19 May 2021"
_CP_PDF_FOOTER_LINE_RES: Tuple[re.Pattern, ...] = (
    re.compile(r"^CP ID:\s*\d+\s*$", re.I),
    re.compile(r"^CP Date:\s*.+$", re.I),
    re.compile(r"^Vessel:\s*.+$", re.I),
    # Chinsay working-copy: tolerate extra whitespace and OCR ID/lD mangling
    re.compile(r"^Chinsay\s*[Il]\s*[Dd]\s*[:.]?\s*\d+\s*$", re.I),
    # Page footer: accept either a clean "Page N of M" or a longer line with a
    # trailing dash + descriptor ("Page 2 of 35 - MORGENSTOND 1-19 May 2021").
    re.compile(r"^Page\s+\d+\s+of\s+\d+(?:\s*[-–—].*)?\s*$", re.I),
    re.compile(r"^WORKING COPY\s*$", re.I),
)

# Single-line SmartCon / owner working-copy footer glued into body text (between
# “1. …” and “2. …”) — all metadata on one physical line in the PDF.
_CP_INLINE_FOOTER_RE = re.compile(
    r"[ \t]*CP ID:\s*\d+\s+CP Date:\s*.+?\s+Vessel:\s*.+?\s+Page\s+\d+\s+of\s+\d+"
    r"(?:[ \t]+WORKING COPY)?[ \t]*",
    re.I,
)


def _strip_cp_pdf_repeating_footer_lines(text: str) -> str:
    """Drop known working-copy footer lines (inserted once per PDF page in the text layer)."""
    if not text:
        return text
    t = text
    t, n_inline = _CP_INLINE_FOOTER_RE.subn(" ", t)
    if n_inline:
        logger.info(
            "[extractor] Collapsed %d inline working-copy footer blob(s) (CP ID … Page N of M)",
            n_inline,
        )
        t = "\n".join(re.sub(r" {2,}", " ", ln) for ln in t.splitlines())
    removed = 0
    out: List[str] = []
    for line in t.splitlines():
        s = line.strip()
        if s and any(rx.match(s) for rx in _CP_PDF_FOOTER_LINE_RES):
            removed += 1
            continue
        out.append(line)
    if removed:
        logger.info("[extractor] Removed %d working-copy / CP footer line(s)", removed)
    return "\n".join(out)


def _strip_trailing_margin_line_numbers(text: str) -> str:
    """
    Remove right-margin line indices glued to the end of body lines, e.g.
    ``...human; 1``, ``...Vessel. 12``, or (when the layout is pervasive)
    ``...crew 2``. Safe patterns run always; an aggressive `` \\d{1,3}$ ``
    pass runs only when many lines share the same artefact.
    """
    if not text:
        return text
    lines = text.splitlines()
    non_empty = [ln.strip() for ln in lines if ln.strip()]
    n = len(non_empty)
    hits = 0
    aggressive_tail = False
    if n >= 8:
        for s in non_empty:
            if (len(s) >= 28 and re.search(r";\s*\d{1,3}\s*$", s)) or (
                len(s) >= 55 and re.search(r"[.!?]\s+\d{1,3}\s*$", s)
            ):
                hits += 1
                continue
            # Long prose line ending with a tight “␠12” tail (no 4-digit year)
            if len(s) >= 58 and re.search(r"\s\d{1,3}\s*$", s) and not re.search(
                r"\d{4}\s*$", s
            ):
                hits += 1
        aggressive_tail = hits >= 8 and (hits / n) >= 0.12

    if aggressive_tail:
        logger.info(
            "[extractor] Stripping trailing right-margin line numbers (aggressive tail pass; "
            "%.0f%% of lines matched artefact pattern)",
            100.0 * hits / n,
        )

    out: List[str] = []
    for line in lines:
        s = line.rstrip()
        if not s:
            out.append(line)
            continue
        lead = line[: len(line) - len(line.lstrip("\t "))]
        orig = s
        if len(s) >= 28:
            s = re.sub(r";\s*\d{1,3}\s*$", "", s).rstrip()
        if len(s) >= 55:
            s = re.sub(r"([.!?])\s+\d{1,3}\s*$", r"\1", s)
        if aggressive_tail and len(s) >= 45 and not re.search(r"\d{4}\s*$", s):
            s = re.sub(r"\s+\d{1,3}\s*$", "", s)
        if s != orig:
            out.append(lead + s)
        else:
            out.append(line)
    return "\n".join(out)


def _strip_working_copy_watermark(text: str) -> str:
    """
    Remove diagonal "WORKING COPY" watermark text that PyMuPDF often captures as
    fragmented/strike-split spans (e.g. ``~~W~~ ORKING COPY``).
    """
    if not text:
        return text
    t = text
    n = 0
    patterns = (
        r"~+\s*W\s*~+\s*ORKING\s+COPY",
        r"(?<![A-Za-z])W\s*ORKING\s+COPY(?![A-Za-z])(?:\s*[-—]+)?",
        r"(?<![A-Za-z])WORKING\s+COPY(?![A-Za-z])(?:\s*[-—]+)?",
    )
    for pat in patterns:
        t, k = re.subn(pat, " ", t, flags=re.I)
        n += k
    # Drop lines that are only watermark noise (after removing markers)
    lines_out: List[str] = []
    dropped = 0
    for line in t.splitlines():
        collapsed = re.sub(r"[*~]+", "", line)
        collapsed = re.sub(r"\s+", " ", collapsed).strip()
        if re.fullmatch(r"W\s*ORKING\s+COPY(?:\s*[-—]+)?", collapsed, re.I):
            dropped += 1
            continue
        if re.fullmatch(r"WORKING\s+COPY(?:\s*[-—]+)?", collapsed, re.I):
            dropped += 1
            continue
        lines_out.append(line)
    t = "\n".join(lines_out)
    if n or dropped:
        logger.info(
            "[extractor] Stripped WORKING COPY watermark (%d inline + %d line hits)",
            n,
            dropped,
        )
    t = "\n".join(re.sub(r" {2,}", " ", ln) for ln in t.splitlines())
    return t


def _sanitize_charter_party_pdf_text(full_text: str) -> str:
    """
    Post-process PDF full text for negotiated charter / working-copy layouts:
    working-copy footers, watermark, inline right-margin numerals, then
    row/column numbers.
    """
    if not full_text:
        return full_text
    t = _strip_cp_pdf_repeating_footer_lines(full_text)
    t = _strip_working_copy_watermark(t)
    t = _strip_trailing_margin_line_numbers(t)
    t = _strip_text_line_numbers(t)
    return t


def _strip_text_line_numbers(text: str) -> str:
    """
    Detect and strip sequential margin line numbers from plain text output.

    Used by the pdfplumber extraction paths (``_extract_pdf_raw`` and
    ``_extract_pdf_pdfplumber``) where coordinate metadata is unavailable.

    pdfplumber preserves PDF line-breaks, so left-margin numbers appear as
    leading integers on each text line::

        '23 1. Duration'
        '24 The Owners agree to let...'
        '26'                               ← empty content line (bare number)
        '27'
        '28 within below mentioned...'

    Stripping produces::

        '1. Duration'
        'The Owners agree to let...'
        ''
        ''
        'within below mentioned...'

    **Right-margin** columns often appear as **bare** numeric lines interleaved
    with body text; a second pass removes those rows when they match the same
    sequential pattern (without requiring 25% of all lines to carry numbers).

    Returns the text unchanged if no pattern matches.
    """
    if not text:
        return text

    lines = text.splitlines()

    left_try = _try_strip_left_prefix_line_numbers(lines)
    if left_try is not None:
        lines = left_try

    bare_try = _strip_bare_sequential_line_rows(lines)
    if bare_try is not None:
        lines = bare_try

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# PDF strikethrough (PyMuPDF vector graphics)
# ─────────────────────────────────────────────────────────────────────────────

def _fitz_span_y_bounds_from_blocks(blocks: list) -> List[Tuple[float, float]]:
    """Collect (y0, y1) for every text span in a ``get_text`` blocks list."""
    bounds: List[Tuple[float, float]] = []
    for blk in blocks:
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            for sp in ln.get("spans", []):
                bb = sp.get("bbox")
                if bb and bb[3] > bb[1]:
                    bounds.append((bb[1], bb[3]))
    return bounds


def _fitz_span_bboxes_from_blocks(blocks: list) -> List[Tuple[float, float, float, float]]:
    """Collect (x0, y0, x1, y1) for every text span — underline vs strike filtering.

    Spans taller than _MAX_SPAN_H are excluded.  In BIMCO SmartCon PDFs the
    "WORKING COPY" watermark is rendered as a single span whose bbox covers
    almost the entire page height (~490 pt).  Every underline on the page falls
    within that bbox at a small rel-value (0.10–0.40) and is therefore
    falsely "confirmed" as a strikethrough candidate.  Normal text spans in
    9–12 pt fonts are 12–18 pt tall; 40 pt is a safe upper limit.
    """
    _MAX_SPAN_H = 40.0   # pt — exclude watermark / decoration spans
    out: List[Tuple[float, float, float, float]] = []
    for blk in blocks:
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            for sp in ln.get("spans", []):
                bb = sp.get("bbox")
                if bb and len(bb) >= 4 and bb[3] > bb[1] and bb[2] > bb[0]:
                    if (bb[3] - bb[1]) <= _MAX_SPAN_H:
                        out.append((bb[0], bb[1], bb[2], bb[3]))
    return out


# Shared strike-vs-underline vertical geometry (relative to a text box).
# Underlines sit just above the baseline → small (y1 - y_mid) vs box height.
# Strikethrough crosses the glyph body → more “air” below the rule.
_FITZ_STRIKE_REL_LO = 0.10   # minimum rel — bands at rel<0.10 are above the glyph
# Real strikethroughs in BIMCO SmartCon PDFs sit at rel≈0.55–0.65 (the glyph
# body takes up the lower portion of the bbox once ascenders/descenders are
# included); raise the ceiling from 0.50 → 0.65 to capture them.
_FITZ_STRIKE_REL_HI = 0.65
_FITZ_STRIKE_MIN_GAP_BELOW = 0.22   # fraction of box height; underlines are < ~0.18
# Underline bleed-up: the underline of span N sits at the very TOP of span N+1
# (due to line-spacing bbox overlap) giving rel≈0.14 AND gap_below≈0.86.
# A real strikethrough never has that much air below it (gap_below≈0.35–0.55).
# Capping gap_below at 0.70 rejects all bleed-up cases while leaving true
# strikes untouched — confirmed by diagnostic on Ocean7_Revolution_NEW_CP.pdf.
_FITZ_STRIKE_MAX_GAP_BELOW = 0.70


def _fitz_path_skip_for_strike_inference(path: dict) -> bool:
    """
    Omit vector paths that are very unlikely to be manuscript strikethrough.

    ``WORKING COPY``-style watermarks and other faint gray line-art otherwise
    yield horizontal bands that bisect glyph boxes (per-char ~~splitting~~).
    """
    fo = path.get("fill_opacity")
    so = path.get("stroke_opacity")
    # Very transparent strokes/fills — typical for background watermarks.
    if so is not None and so < 0.38:
        return True
    if fo is not None and fo < 0.38 and path.get("type") in (None, "f", "fs"):
        return True

    def _rgb_min(comp) -> Optional[float]:
        if comp is None:
            return None
        if isinstance(comp, (tuple, list)) and len(comp) >= 3:
            return min(float(comp[0]), float(comp[1]), float(comp[2]))
        return None

    # Light gray (or near-white) strokes/fills — not manuscript strike ink.
    _GRAY_CUT = 0.51
    for key in ("color", "fill"):
        m = _rgb_min(path.get(key))
        if m is not None and m >= _GRAY_CUT:
            return True
    return False


def _fitz_horizontal_rule_is_strikethrough_in_box(
    x0: float,
    x1: float,
    y_mid: float,
    bx0: float,
    by0: float,
    bx1: float,
    by1: float,
    *,
    min_overlap_pt: float,
    min_overlap_frac: Optional[float] = None,
) -> bool:
    """True if a near-horizontal drawing rule crosses the mid-body of *box*, not underline zone."""
    if not (by0 <= y_mid <= by1):
        return False
    overlap_w = min(x1, bx1) - max(x0, bx0)
    bw, bh = bx1 - bx0, by1 - by0
    if bw <= 0 or bh <= 0:
        return False
    if min_overlap_frac is not None:
        if overlap_w < min_overlap_frac * bw:
            return False
    elif overlap_w < min_overlap_pt:
        return False
    rel = (y_mid - by0) / bh
    gap_below = (by1 - y_mid) / bh
    if rel < _FITZ_STRIKE_REL_LO or rel > _FITZ_STRIKE_REL_HI:
        return False
    if gap_below < _FITZ_STRIKE_MIN_GAP_BELOW:
        return False
    # Reject underline bleed-up: the underline of span N sits at the very top
    # of span N+1, giving a very large gap_below (≈0.86).  Real strikethroughs
    # cross the glyph body and have moderate gap_below (≈0.35–0.55).
    if gap_below > _FITZ_STRIKE_MAX_GAP_BELOW:
        return False
    return True


def _fitz_page_span_ybounds(page) -> List[Tuple[float, float]]:
    """All text span vertical extents on a page — used to reject non-text horizontal rules."""
    return _fitz_span_y_bounds_from_blocks(page.get_text("rawdict")["blocks"])


def _fitz_collect_strike_bands(
    page,
    blocks: Optional[list] = None,
) -> List[Tuple[float, float, float]]:
    """
    Horizontal strike segments as (x0, x1, y_mid) in page coordinates.

    Producers vary: thin filled rects, stroke lines, or both.  PyMuPDF may
    represent a line as an ``l`` item and/or a thin bounding ``rect``.  We
    ingest both so stroke-only strikethroughs are not missed when ``rect`` is
    absent or unhelpful.

    Bands are filtered so y_mid lies in the **middle** vertical band of at
    least one overlapping text span (horizontal overlap required).  Lines that
    sit in the **bottom** band of a span (typical **underlines**) are rejected;
    only rules that cross the glyph body are kept as strikethrough candidates.

    Pass *blocks* from a cached ``page.get_text(\"rawdict\")[\"blocks\"]`` to
    avoid parsing the page twice when the caller already loads rawdict.
    """
    if blocks is None:
        blocks = page.get_text("rawdict")["blocks"]
    span_bboxes = _fitz_span_bboxes_from_blocks(blocks)
    if not span_bboxes:
        return []

    # Slightly looser than legacy (3 / 10): accommodates thicker hairlines and
    # short per-phrase strikes without admitting typical table borders.
    _MAX_BAND_H = 4.5
    _MIN_BAND_W = 8.0
    _MAX_LINE_DY = 2.8   # pt — nearly horizontal

    candidates: List[Tuple[float, float, float]] = []
    for d in page.get_drawings():
        if _fitz_path_skip_for_strike_inference(d):
            continue
        rect = d.get("rect")
        if rect is not None:
            band_h = rect.y1 - rect.y0
            band_w = rect.x1 - rect.x0
            if band_h <= _MAX_BAND_H and band_w >= _MIN_BAND_W:
                y_mid = (rect.y0 + rect.y1) / 2
                candidates.append((rect.x0, rect.x1, y_mid))

        for it in d.get("items") or ():
            if not it:
                continue
            op = it[0]
            if op == "l":
                p1, p2 = it[1], it[2]
                x0, x1 = sorted((p1.x, p2.x))
                y0, y1 = sorted((p1.y, p2.y))
                dx, dy = x1 - x0, y1 - y0
                if dx >= _MIN_BAND_W and dy <= _MAX_LINE_DY:
                    candidates.append((x0, x1, (y0 + y1) / 2))
            elif op == "re":
                r2 = it[1]
                band_h = r2.y1 - r2.y0
                band_w = r2.x1 - r2.x0
                if band_h <= _MAX_BAND_H and band_w >= _MIN_BAND_W:
                    y_mid = (r2.y0 + r2.y1) / 2
                    candidates.append((r2.x0, r2.x1, y_mid))

    strike_bands: List[Tuple[float, float, float]] = []
    for x0, x1, y_mid in candidates:
        ok = False
        for sx0, sy0, sx1, sy1 in span_bboxes:
            if _fitz_horizontal_rule_is_strikethrough_in_box(
                x0, x1, y_mid, sx0, sy0, sx1, sy1, min_overlap_pt=4.0
            ):
                ok = True
                break
        if ok:
            strike_bands.append((x0, x1, y_mid))
    return strike_bands


def _fitz_char_is_struck(bbox, strike_bands: List[Tuple[float, float, float]]) -> bool:
    """Glyph-level strike test — uses the same vertical band as span-level."""
    sx0, sy0, sx1, sy1 = bbox
    cw = sx1 - sx0
    ch = sy1 - sy0
    if cw <= 0 or ch <= 0:
        return False
    for bx0, bx1, by_mid in strike_bands:
        if _fitz_horizontal_rule_is_strikethrough_in_box(
            bx0, bx1, by_mid, sx0, sy0, sx1, sy1,
            min_overlap_pt=0.0,
            min_overlap_frac=0.52,
        ):
            return True
    return False


def _fitz_span_is_struck(bbox, strike_bands: List[Tuple[float, float, float]]) -> bool:
    """
    Whole-span fallback when ``rawdict`` glyph boxes are missing.

    A narrow strike that crosses only part of a long merged text run must
    **not** mark the entire span — that was the bug introduced by the
    overlap/min(span_w, band_w) heuristic.  We only flag the full span when
    the strike clearly covers most of the run *or* the decorative line is
    nearly as wide as the span (full-width clause strike).
    """
    sx0, sy0, sx1, sy1 = bbox
    span_w = sx1 - sx0
    span_h = sy1 - sy0
    if span_w <= 0 or span_h <= 0:
        return False

    for bx0, bx1, by_mid in strike_bands:
        overlap_w = min(sx1, bx1) - max(sx0, bx0)
        if overlap_w < 1.5:
            continue
        band_w = bx1 - bx0
        if band_w <= 0:
            continue
        if not _fitz_horizontal_rule_is_strikethrough_in_box(
            bx0, bx1, by_mid, sx0, sy0, sx1, sy1, min_overlap_pt=1.5
        ):
            continue
        # Strike sits across most of this span's width.
        if overlap_w >= 0.42 * span_w:
            return True
        # Nearly full-width rule / line drawn as wide as the text object.
        if band_w >= 0.85 * span_w and overlap_w >= 0.82 * band_w:
            return True
    return False


def _fitz_span_plain_text(span: dict) -> str:
    """Plain text for a span from ``rawdict`` or ``dict``."""
    t = span.get("text")
    if t is not None:
        return t
    return "".join(ch.get("c", "") for ch in span.get("chars") or ())


def _fitz_span_strike_segments(
    span: dict,
    strike_bands: List[Tuple[float, float, float]],
) -> List[Tuple[str, bool]]:
    """
    Split a span into (substring, is_struck) pieces using per-glyph boxes when
    present (``rawdict``), otherwise one piece with whole-span heuristics.
    """
    chars = span.get("chars") or []
    if not chars:
        raw = _fitz_span_plain_text(span)
        if not raw.strip():
            return []
        return [(raw, _fitz_span_is_struck(span["bbox"], strike_bands))]

    segments: List[Tuple[str, bool]] = []
    cur: List[str] = []
    cur_st: Optional[bool] = None

    for ch in chars:
        c = ch.get("c") or ""
        bb = ch.get("bbox")
        if bb and bb[2] > bb[0] and bb[3] > bb[1]:
            st = _fitz_char_is_struck(bb, strike_bands)
        else:
            st = False
        if cur_st is None:
            cur_st = st
            cur.append(c)
        elif st == cur_st:
            cur.append(c)
        else:
            segments.append(("".join(cur), cur_st))
            cur = [c]
            cur_st = st
    if cur and cur_st is not None:
        segments.append(("".join(cur), cur_st))

    if not segments:
        raw = _fitz_span_plain_text(span)
        if raw.strip():
            return [(raw, _fitz_span_is_struck(span["bbox"], strike_bands))]
        return []

    return [(t, st) for t, st in segments if t]


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — PyMuPDF (primary)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_fitz(raw_bytes: bytes) -> List[Dict]:
    """
    Extract structured chunks from a PDF using PyMuPDF font metadata.

    Strategy:
    - Read every page using get_text("dict") to capture per-span font size,
      bold flag, and bounding-box coordinates.
    - Filter out tiny text (< 7pt) which is typically cargo diagram labels or
      other graphical annotation noise.
    - For each page, use bin-density analysis to detect two-column layouts.
    - CRITICAL — cross-page clause joining: elements are accumulated across
      ALL pages with a global y-offset rather than being chunked per-page.
      This ensures that clause bodies which span multiple pages are kept
      together under their heading and not split into titleless orphan chunks.
      (e.g. NYPE 2015 clause 1 "Duration/Trip Description" has its header on
      page 1 and body text that continues onto page 2.)
    - Two-column pages are handled by placing the left column's elements before
      the right column's elements in global y-order (right column base offset =
      y_global + page_height), so column reading order is preserved globally.
    - A single `_column_to_chunks` call processes all pages' elements as one
      continuous document, maintaining header context across page boundaries.
    """
    import fitz

    MIN_FONT = 7.0   # pt — skip cargo diagram / slot-plan label text

    all_elements: List[Dict] = []   # global accumulator across all pages
    all_plain_lines: List[str] = []  # fallback accumulator

    with fitz.open(stream=raw_bytes, filetype="pdf") as pdf:
        y_global: float = 0.0   # running y-offset into global coordinate space

        for page in pdf:
            page_height = page.rect.height
            page_width  = page.rect.width
            elements: List[Dict] = []

            blocks = page.get_text("rawdict")["blocks"]
            strike_bands = _fitz_collect_strike_bands(page, blocks=blocks)

            for block in blocks:
                if block.get("type") != 0:   # skip image blocks
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text_spans = [s for s in spans if _fitz_span_plain_text(s).strip()]
                    if not text_spans:
                        continue

                    # Skip sub-threshold font sizes (noise / diagram labels)
                    max_size = max(s["size"] for s in text_spans)
                    if max_size < MIN_FONT:
                        continue

                    # Build text with per-span strikethrough markers.
                    # Consecutive struck spans are GROUPED before calling
                    # _clean_pdf_text() so that ligature fragments like
                    # "par" + "ti" + "es" are merged into "parties" inside a
                    # single ~~…~~ block rather than ~~par~~ ~~ti~~ ~~es~~.
                    parts: List[str] = []
                    struck_buf: List[str] = []

                    def _flush_struck() -> None:
                        if not struck_buf:
                            return
                        merged = _clean_pdf_text(" ".join(struck_buf)).strip()
                        if merged:
                            parts.append(f"~~{merged}~~")
                        struck_buf.clear()

                    for s in text_spans:
                        sub_struck: List[str] = []
                        for piece, is_struck in _fitz_span_strike_segments(s, strike_bands):
                            if piece == "":
                                continue
                            if is_struck:
                                if not piece.strip():
                                    continue
                                sub_struck.append(piece)
                            else:
                                if sub_struck:
                                    struck_buf.append("".join(sub_struck))
                                    sub_struck = []
                                _flush_struck()
                                parts.append(_clean_pdf_text(piece))
                        if sub_struck:
                            struck_buf.append("".join(sub_struck))
                    _flush_struck()

                    combined = " ".join(parts).strip()
                    if not combined:
                        continue

                    # Bold = flag bit 4 (value 16) set on ALL content spans
                    is_bold = all(bool(s["flags"] & 16) for s in text_spans)
                    x_pos = min(s["bbox"][0] for s in text_spans)
                    y_pos = min(s["bbox"][1] for s in text_spans)

                    elements.append({
                        "text": combined,
                        "size": max_size,
                        "is_bold": is_bold,
                        "x": x_pos,
                        "y": y_pos,   # page-local y; globalised below
                    })
                    all_plain_lines.append(combined)

            if not elements:
                y_global += page_height + 100
                continue

            # Per-page column detection (layout may differ between pages)
            col_split = _detect_column_split(elements, page_width)

            if col_split is not None:
                # Two-column page: sort each column by (y, x), then
                # place the entire left column before the right in global order.
                left_els  = sorted(
                    [e for e in elements if e["x"] <  col_split],
                    key=lambda e: (e["y"], e["x"]),
                )
                right_els = sorted(
                    [e for e in elements if e["x"] >= col_split],
                    key=lambda e: (e["y"], e["x"]),
                )
                # Left column: y_global + local_y (naturally before right)
                for e in left_els:
                    all_elements.append({**e, "y": y_global + e["y"]})
                # Right column: y_global + page_height + local_y
                # (ensures right-col elements sort AFTER left-col elements)
                right_base = y_global + page_height + 10
                for e in right_els:
                    all_elements.append({**e, "y": right_base + e["y"]})
            else:
                # Single-column page: simply offset local y into global space
                for e in elements:
                    all_elements.append({**e, "y": y_global + e["y"]})

            # Advance global offset past this page (+ small inter-page gap)
            y_global += page_height * 2 + 200

    # ── Margin line-number stripping (left and/or right column) ─────────────
    # Charter-party PDFs print sequential integers in a margin for cross-ref.
    # PyMuPDF yields separate elements per column; strip both sides before chunking.
    #
    # Two-pass strategy per detected column:
    #   Pass 1 — remove elements whose entire text is a bare line number
    #   Pass 2 — strip leading "N " prefix via _LN_PREFIX_RE (merged left margin)
    strip_ids: set = set()
    for _side in ("left", "right"):
        ln_info = _detect_margin_line_numbers(all_elements, margin_side=_side)
        if ln_info:
            strip_ids |= ln_info["strip_ids"]
    if strip_ids:
        all_elements = [e for e in all_elements if id(e) not in strip_ids]
        for e in all_elements:
            e['text'] = _LN_PREFIX_RE.sub('', e['text'])
    # ──────────────────────────────────────────────────────────────────────────

    # Single global pass — header context is maintained across page boundaries
    all_chunks = _column_to_chunks(all_elements) if all_elements else []

    if not all_chunks:
        all_chunks = _fallback_fixed_chunks(all_plain_lines)

    return all_chunks


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — pdfplumber (fallback, no font metadata)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_pdfplumber(raw_bytes: bytes) -> List[Dict]:
    """
    Extract chunks from PDF via pdfplumber (clause-regex only, no font info).
    Used when PyMuPDF is not installed.
    """
    import io, pdfplumber

    full_text_lines: List[str] = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text = _clean_pdf_text(text)
                full_text_lines.extend(text.splitlines())

    if not full_text_lines:
        return []

    # Working-copy footers, margin numerals, then row/column line numbers (CP PDFs).
    joined = _sanitize_charter_party_pdf_text("\n".join(full_text_lines))
    clause_try = _presplit_on_clauses(joined)
    if clause_try:
        return clause_try

    full_text_lines = joined.splitlines()

    chunks: List[Dict] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []

    def _flush():
        body = "\n".join(current_lines).strip()
        if body and not _is_junk_body(body):
            chunks.append({"title": current_title, "body": body})

    for line in full_text_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _is_clause_header(stripped):
            _flush()
            current_title = stripped
            current_lines = []
        else:
            current_lines.append(stripped)

    _flush()

    return chunks or _fallback_fixed_chunks(full_text_lines)


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — raw mode (full-text, AI-first)
# ─────────────────────────────────────────────────────────────────────────────

# Document categories where heuristic font-based chunking causes more harm
# than good.  For these, we extract the entire document as a single blob (or
# one chunk per page for very long docs) and let the AI enrichment pass do all
# the semantic splitting.  This is the right call when:
#   - The PDF uses bold/larger text for data labels (e.g. "Flag", "GT/NT")
#     rather than section headers, confusing the header detector.
#   - The document is a fixture recap where numbered items like "1. Vessel /
#     Owners" have dozens of sub-items that belong in the same chunk.
#   - The document is a narrative addendum where every paragraph is connected.
_RAW_EXTRACTION_CATEGORIES = {
    "fixture_recap",
    "addendum",
    "delivery_details",
    "vessel_specifications",
    "vessel_owners_details",
    "speed_consumption",
}

# Document categories that require BOTH PyMuPDF's graphical strikethrough
# detection AND numbered-clause splitting (instead of font-based headers).
#
# Charter parties are the canonical case: they are heavily negotiated, so
# struck-out text carries legal significance and must be preserved as
# ~~markers~~.  At the same time, BIMCO/NYPE PDFs use bold body text for
# legal emphasis, which causes the font-based header detector to misfire and
# produce hundreds of spurious chunks.
#
# The fitz-presplit path solves both problems:
#   1. Extracts text via PyMuPDF (graphical strikethrough preserved)
#   2. Splits on numbered top-level clauses, not font properties
#   3. Falls back to page batches + AI enrichment when no clause structure
_FITZ_PRESPLIT_CATEGORIES = {
    "charter_party",
}

# Word count threshold below which the whole document is returned as ONE chunk
# rather than split by page.  7-page fixture recaps are ~3 000 words — well
# within gpt-4o-mini's 128k-token context.
_RAW_SINGLE_CHUNK_THRESHOLD = 6_000  # words


def _presplit_line_core_for_heading_match(stripped: str) -> str:
    """
    Peel outer ~~…~~ wrappers (PyMuPDF strikethrough markdown) and leading /
    trailing ``**`` bold markers so lines like ``~~APPENDIX A~~`` or
    ``** Clause 46 - Title **`` can still match heading detection.
    """
    core = (stripped or "").strip()
    while core.startswith("~~") and core.endswith("~~") and len(core) >= 4:
        core = core[2:-2].strip()
    while True:
        t = re.sub(r"^\*+\s*", "", core)
        t = re.sub(r"\s*\*+$", "", t).strip()
        if t == core:
            break
        core = t
    return core


def _presplit_parse_clause_heading_line(core: str) -> Optional[Tuple[int, str]]:
    """
    If *core* opens a top-level clause heading, return (clause_num, title_rest).

    Supports:
      - ``46. BIMCO …``                             (main form)
      - ``Clause 46 - …`` / ``Clause 46 — …``       (rider / additional-clauses)
      - ``Clause 58`` / ``Clause 58.``              (bare rider heading; the
                                                     real title or body
                                                     follows on the next line)
    """
    if not core:
        return None
    m = re.match(r"^(\d{1,3})\.\s+(.+)$", core)
    if m:
        return int(m.group(1)), (m.group(2) or "").strip()
    m2 = re.match(r"(?i)^Clause\s+(\d{1,3})\s*[—–\-]\s*(.+)$", core)
    if m2:
        return int(m2.group(1)), (m2.group(2) or "").strip()
    # e.g. "47 Charter Party - SEE APPENDIX B FOR VESSEL DESCRIPTION" (owner riders)
    m3 = re.match(
        r"(?i)^(\d{1,3})\s+Charter Party\s*[—–\-]\s*(.+)$",
        core,
    )
    if m3:
        n = int(m3.group(1))
        r = (m3.group(2) or "").strip()
        tail = f"Charter Party - {r}" if r else "Charter Party"
        return n, tail
    # Bare "Clause N" / "Clause N." line — common in additional/rider clauses
    # where the body text is a self-contained paragraph and there is no
    # inline title (e.g. "Clause 59\nThe Charterers have the liberty…").
    m4 = re.match(r"(?i)^Clause\s+(\d{1,3})\s*\.?\s*$", core)
    if m4:
        return int(m4.group(1)), ""
    return None


def _title_is_bare_clause_marker(title: str) -> bool:
    """
    True when chunk title is only the numeric / chapter marker so the real
    BIMCO rider name is likely on the first line of the body (separate PDF line).
    """
    if not (title or "").strip():
        return False
    core = _presplit_line_core_for_heading_match(title.strip())
    if re.fullmatch(r"\d{1,3}\.?", core):
        return True
    if re.fullmatch(r"CHAPTER\s+\d{1,3}\.?", core, re.I):
        return True
    if re.fullmatch(r"Clause\s+\d{1,3}\.?", core, re.I):
        return True
    return False


def _looks_like_rider_clause_subtitle_line(line: str) -> bool:
    """
    First body line that names the rider clause (BIMCO / NYPE …), not (a)/(1).
    """
    s = (line or "").strip()
    if len(s) < 12:
        return False
    if re.match(r"^\(\s*[a-z]\s*\)\s", s, re.I):
        return False
    if re.match(r"^\(\s*\d+\s*\)\s", s):
        return False
    if re.match(r"^[a-z]\)\s", s, re.I):
        return False
    if re.match(r"^\d+\.\d+\s", s):
        return False
    head = s[:120]
    if re.search(r"(?i)\bBIMCO\b", head):
        return True
    if re.search(r"(?i)\bNYPE\b", head):
        return True
    if re.search(r"(?i)Clause\s+for\s+(?:Time\s+)?Charter", head):
        return True
    if re.search(r"(?i)\b(?:INTERTANKO|BARECON|SUPPLYTIME)\b", head):
        return True
    return False


def _promote_rider_subtitle_line_into_chunk_titles(chunks: List[Dict]) -> None:
    """Merge ``46.`` + first body line ``BIMCO …`` into one title."""
    for c in chunks:
        title = (c.get("title") or "").strip()
        body = (c.get("body") or "").strip()
        if not body or not title:
            continue
        if not _title_is_bare_clause_marker(title):
            continue
        if "\n" in body:
            first_line, _, rest = body.partition("\n")
        else:
            first_line, rest = body, ""
        fl = first_line.strip()
        if not _looks_like_rider_clause_subtitle_line(fl):
            continue
        joiner = " " if title.endswith((".", ":")) else " - "
        c["title"] = f"{title}{joiner}{fl}".strip()
        c["body"] = rest.strip()


def _presplit_on_clauses(full_text: str) -> List[Dict]:
    """
    Split a large text blob on numbered top-level clause headers.

    Matches lines like:
        "1. Vessel / Owners"
        "14. C/P Details"
        "** Clause 46 - BIMCO Infectious… **"   ← additional/rider pages
        "~~46. Title~~"   ← strikethrough wrapper peeled before matching
        "1.1 Owners confirm ..."   ← sub-clause: kept inside parent chunk

    Also handles the BIMCO SmartCon split-line format where PyMuPDF renders
    the clause number and title as consecutive lines:
        "1."
        "Duration/Trip Description"
    These are joined into "1. Duration/Trip Description" in a pre-pass.

    Only TOP-level integers (no dot after the digit group) trigger a new
    chunk boundary, so sub-clauses like "1.1", "1.2" stay within their
    parent section.

    Returns a list of {"title": str | None, "body": str} dicts where each
    dict corresponds to one numbered clause (or a preamble if text precedes
    clause 1).  Chunks are further split at MAX_CHUNK_WORDS if a single
    clause is very long.
    """
    # Matches a bare clause number on its own line: "1." or "89."
    _BARE_NUM_RE = re.compile(r"^(\d{1,3})\.$")

    # Appendix / annex / schedule / additional-clauses header detection
    # (best-effort). Covers two common patterns:
    #   1. Line starts with keyword + identifier:
    #        "APPENDIX A", "Annex B", "Schedule I", "Exhibit C",
    #        "ADDENDUM No. 1"
    #   2. Short all-caps line containing keyword anywhere:
    #        "NYPE 2015 APPENDIX A (VESSEL DESCRIPTION)"
    #        "ADDITIONAL CLAUSES TO CHARTER PARTY DATED 19TH MAY 2021"
    # Custom appendix titles that use neither pattern are not caught — they
    # remain in the last clause body, which is acceptable.
    _APPENDIX_START_RE = re.compile(
        r"^(?:APPENDIX|ANNEX|SCHEDULE|EXHIBIT|ADDENDUM|ADDENDA)\b",
        re.IGNORECASE,
    )
    # Two flavours of structural divider:
    #   "restart"   — appendix/annex bodies that contain their own internal
    #                 numbering (clause 1, 2, 3 …) so we MUST reset
    #                 max_clause_seen or those numbers are swallowed by the
    #                 monotonic guard.
    #   "continue"  — additional / rider / addendum clauses whose numbers
    #                 continue from the main body (e.g. main body ends at
    #                 clause 57, rider starts at Clause 58). Resetting
    #                 max_clause_seen would let a stray "Clause 1" inside the
    #                 rider falsely register as a new top-level clause.
    _APPENDIX_RESTART_KEYWORDS = ("APPENDIX", "ANNEX", "SCHEDULE", "EXHIBIT")
    _APPENDIX_CONTINUE_KEYWORDS = (
        "ADDITIONAL CLAUSE",  # matches "ADDITIONAL CLAUSES" too
        "RIDER CLAUSE",
        "ADDENDUM",
        "ADDENDA",
    )
    _APPENDIX_ALL_KEYWORDS = (
        _APPENDIX_RESTART_KEYWORDS + _APPENDIX_CONTINUE_KEYWORDS
    )

    def _is_appendix_header(text: str) -> bool:
        if not text or len(text) > 80:
            return False
        if _APPENDIX_START_RE.match(text):
            return True
        # Short all-caps line containing one of the keywords.
        alpha = [c for c in text if c.isalpha()]
        if alpha and all(c.isupper() for c in alpha):
            up = text.upper()
            return any(kw in up for kw in _APPENDIX_ALL_KEYWORDS)
        return False

    def _appendix_restarts_numbering(text: str) -> bool:
        """True for APPENDIX/ANNEX/SCHEDULE/EXHIBIT (their own clause space)."""
        up = text.upper()
        # Continuation keywords win when both could match — e.g. an
        # "ADDENDUM No. 1 - APPENDIX A" line is logically a continuation
        # of the main numbering.
        if any(kw in up for kw in _APPENDIX_CONTINUE_KEYWORDS):
            return False
        return any(kw in up for kw in _APPENDIX_RESTART_KEYWORDS)

    raw_lines = full_text.splitlines()

    # Pre-pass: join bare "N." lines with the following content line.
    # BIMCO SmartCon PDFs often render the clause number and title as
    # separate text spans/lines in PyMuPDF, e.g.:
    #   "1."                        ← fitz line 1
    #   "Duration/Trip Description" ← fitz line 2
    # We merge them so TOP_CLAUSE_RE can match "1. Duration/Trip Description".
    lines: List[str] = []
    i = 0
    while i < len(raw_lines):
        stripped = raw_lines[i].strip()
        if _BARE_NUM_RE.match(stripped) and i + 1 < len(raw_lines):
            next_stripped = raw_lines[i + 1].strip()
            # Only join if the next line is non-empty content (not another
            # bare number or a known non-title pattern)
            if next_stripped and not _BARE_NUM_RE.match(next_stripped):
                lines.append(stripped + " " + next_stripped)
                i += 2
                continue
        lines.append(raw_lines[i])
        i += 1
    chunks: List[Dict] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []
    max_clause_seen: int = 0   # monotonic guard — same logic as _column_to_chunks

    def _flush():
        body = "\n".join(current_lines).strip()
        if body:
            chunks.append({"title": current_title, "body": body})

    for line in lines:
        stripped = line.strip()
        core = _presplit_line_core_for_heading_match(stripped)
        heading = _presplit_parse_clause_heading_line(core)
        if heading:
            num, title_rest = heading
            synthetic = f"{num}. {title_rest}" if title_rest else f"{num}."
            # Monotonic guard: a numbered line whose number is LOWER than the
            # highest clause seen so far is a sub-clause list item inside the
            # current clause body, not a new top-level clause.
            # e.g. "1. the Vessel shall..." inside clause 9 → body text.
            if num <= max_clause_seen:
                current_lines.append(line)
            elif _looks_like_merged_margin_clause_number(
                num, max_clause_seen, title_rest
            ):
                # e.g. right-margin "47" merged with "The Owners reserve…"
                current_lines.append(line)
            else:
                _flush()
                max_clause_seen = num
                # One PDF text line can hold an entire lettered sub-paragraph
                # ("1. Owners will arrange for the armed guards…"). Using that
                # as the chunk title blows up the title field and reads as a
                # mid-sentence split in the UI.
                #
                # Heuristic: real NYPE/BIMCO numbered clause titles are terse
                # (≤ 10 words after the number, e.g. "10. Rate of Hire/
                # Redelivery Areas and Notices").  Lines with > 10 words after
                # the number are almost always prose sentences inside an
                # additional-clauses block, not top-level headings.
                # The "> 220 chars" guard is kept as a hard upper limit.
                _title_word_count = len(title_rest.split())
                if len(synthetic) > 220 or _title_word_count > 10:
                    current_title = f"{num}."
                    current_lines = (
                        [title_rest.strip()] if title_rest.strip() else []
                    )
                else:
                    current_title = synthetic
                    current_lines = []
        elif _is_appendix_header(core):
            # Non-numbered structural boundary: Appendix A, Annex B, etc.
            # For APPENDIX/ANNEX/SCHEDULE/EXHIBIT we reset max_clause_seen so
            # sub-items inside the appendix (whose numbering restarts at 1)
            # are not swallowed by the monotonic guard. For ADDITIONAL
            # CLAUSES / RIDER / ADDENDUM the numbering CONTINUES from the
            # main body (e.g. main ends at 57, rider starts at Clause 58),
            # so we keep max_clause_seen intact — otherwise stray small
            # numerals inside the rider get misread as top-level clauses.
            #
            # Exception: a fully struck-through appendix marker (e.g.
            # "~~APPENDIX "A"~~") is deleted content — the appendix was
            # removed in negotiation and is NOT a real structural boundary.
            if stripped.startswith("~~"):
                # Struck-through appendix — treat as body text, keep counter.
                current_lines.append(line)
            else:
                _flush()
                if _appendix_restarts_numbering(core):
                    max_clause_seen = 0
                # else: continuation rider — preserve max_clause_seen so the
                #       parser correctly accepts Clause 58, Clause 59, …
                current_title = stripped
                current_lines = []
        else:
            current_lines.append(line)

    _flush()

    # If no clause markers were found, return empty so caller falls back
    if not any(c["title"] for c in chunks):
        return []

    _promote_rider_subtitle_line_into_chunk_titles(chunks)

    # Use a generous per-clause ceiling: clause integrity matters more than
    # chunk size for contract documents.  5 000 words ≈ 6 500 tokens — safely
    # below text-embedding-3-small's 8 191-token hard limit.
    return _split_into_chunks(chunks, max_words=5_000)


def _extract_pdf_raw(raw_bytes: bytes) -> List[Dict]:
    """
    Extract the full text of a PDF using pdfplumber, then split it into
    per-clause chunks without relying on font metadata.

    Strategy:
    1. Extract text page-by-page (pdfplumber handles text-layer PDFs well).
    2. Try to split on numbered top-level clauses ("1. Vessel / Owners",
       "14. C/P Details" …).  This produces correctly-sized chunks and
       preserves clause context that font-based detection misses.
    3. If no numbered structure is found, return the whole document as one
       or a few page-batched chunks for AI enrichment to handle.

    All extracted content is passed intact — no lines are discarded.
    The AI enrichment pass then assigns titles and may further refine splits.
    """
    import io
    import pdfplumber

    page_texts: List[str] = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text = _clean_pdf_text(text).strip()
                if text:
                    page_texts.append(text)

    if not page_texts:
        return []

    full_text = "\n\n".join(page_texts)

    # Sanitize working-copy footers, inline line indices, and row/column numbers
    # so clause headers like "23 1. Duration" become "1. Duration".
    full_text = _sanitize_charter_party_pdf_text(full_text)

    # Try numbered-clause splitting first
    clause_chunks = _presplit_on_clauses(full_text)
    if clause_chunks:
        logger.info(
            f"[extractor] Raw PDF: presplit into {len(clause_chunks)} clause chunks"
        )
        return clause_chunks

    # No numbered structure — return as one or a few word-budget batches
    # so AI enrichment still sees the full context
    total_words = len(full_text.split())
    logger.info(
        f"[extractor] Raw PDF: no clause markers found, "
        f"returning {total_words} words as page batches"
    )

    if total_words <= _RAW_SINGLE_CHUNK_THRESHOLD:
        return [{"title": None, "body": full_text}]

    chunks: List[Dict] = []
    batch_lines: List[str] = []
    batch_words = 0

    for page_text in page_texts:
        page_words = len(page_text.split())
        if batch_words + page_words > _RAW_SINGLE_CHUNK_THRESHOLD and batch_lines:
            chunks.append({"title": None, "body": "\n\n".join(batch_lines)})
            batch_lines = []
            batch_words = 0
        batch_lines.append(page_text)
        batch_words += page_words

    if batch_lines:
        chunks.append({"title": None, "body": "\n\n".join(batch_lines)})

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — PyMuPDF + clause presplit (hybrid, for negotiated contracts)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_fitz_presplit(raw_bytes: bytes) -> List[Dict]:
    """
    Hybrid extraction for negotiated contract PDFs (e.g. charter parties).

    Combines the two things each path does best:
      - PyMuPDF  : graphical strikethrough detection via drawing-rectangle
                   analysis.  Struck spans are wrapped in ~~markers~~ exactly
                   as _extract_pdf_fitz does, preserving legally significant
                   deleted text that pdfplumber cannot see at all.
      - Raw path : _presplit_on_clauses splits on numbered top-level clause
                   headers (1. Duration, 2. Delivery …) instead of relying on
                   font size / bold flags, which misfire on BIMCO/NYPE PDFs
                   where bold is used for legal emphasis inside clause bodies.

    Falls back to fixed page batches (for AI enrichment) when the document
    has no numbered clause structure.

    MIN_FONT is set to 8.0pt (higher than the 7.0pt used in _extract_pdf_fitz)
    to filter out the BIMCO SmartCon copyright footer, which is consistently
    printed at 7.0pt on every page.  All substantive clause text in NYPE/BIMCO
    forms is 9pt or larger, so no legal content is lost.
    """
    import fitz

    MIN_FONT = 8.0
    page_texts: List[str] = []

    # Footer-band drop: anything whose vertical midpoint falls in the bottom
    # FOOTER_BAND_FRAC of the page is treated as page-furniture (copyright
    # disclaimer, vendor ID, page X of N) and discarded. On scanned/OCR'd
    # PDFs the disclaimer band is randomized garbage that no regex can catch
    # reliably, but it ALWAYS sits in this band, so position-based filtering
    # is the safe net. The fraction is conservative (10%): real clause text
    # may extend close to the bottom margin on dense pages, but never into
    # the printed footer area.
    FOOTER_BAND_FRAC = 0.10

    with fitz.open(stream=raw_bytes, filetype="pdf") as pdf:
        for page in pdf:
            blocks = page.get_text("rawdict")["blocks"]
            strike_bands = _fitz_collect_strike_bands(page, blocks=blocks)
            footer_y_threshold = page.rect.height * (1.0 - FOOTER_BAND_FRAC)

            page_lines: List[str] = []
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    # Drop lines positioned in the printed footer band —
                    # vendor copyright disclaimer, "Chinsay ID", "Page N of M",
                    # and any OCR garbage layered on top of them.
                    line_bbox = line.get("bbox")
                    if line_bbox and len(line_bbox) >= 4:
                        line_y_mid = (line_bbox[1] + line_bbox[3]) / 2.0
                        if line_y_mid >= footer_y_threshold:
                            continue
                    # MIN_FONT gate: at least one substantive (non-whitespace)
                    # span on the line must clear the font threshold. Pure
                    # whitespace spans don't have a meaningful size.
                    sized_spans = [
                        s for s in spans if _fitz_span_plain_text(s).strip()
                    ]
                    if not sized_spans:
                        continue
                    max_size = max(s["size"] for s in sized_spans)
                    if max_size < MIN_FONT:
                        continue

                    # Build line text by concatenating EVERY span's text in
                    # source order — including whitespace-only spans, which
                    # PyMuPDF often emits as their own glyph cluster between
                    # words. Strikethrough state is tracked per character
                    # segment; consecutive struck pieces are wrapped in a
                    # single ~~…~~ marker.
                    #
                    # We deliberately do NOT use " ".join(parts): on heavily
                    # span-fragmented OCR'd PDFs (e.g. Toshiba MFP scans), a
                    # single visual line can yield 25+ spans and joining with
                    # spaces produces "Th e Vesse l dur i ng" instead of
                    # "The Vessel during". The character data inside each
                    # span already carries its own leading/trailing whitespace.
                    raw_pieces: List[Tuple[str, bool]] = []
                    for s in spans:
                        for piece, is_struck in _fitz_span_strike_segments(
                            s, strike_bands
                        ):
                            if piece:
                                raw_pieces.append((piece, is_struck))

                    # Coalesce adjacent (struck/non-struck) pieces, then emit
                    # ~~…~~ wrappers around struck runs only.
                    line_parts: List[str] = []
                    cur_text: List[str] = []
                    cur_struck: Optional[bool] = None

                    def _flush_run() -> None:
                        nonlocal cur_text, cur_struck
                        if not cur_text:
                            return
                        merged = _clean_pdf_text("".join(cur_text))
                        if cur_struck and merged.strip():
                            line_parts.append(f"~~{merged.strip()}~~")
                        elif merged:
                            line_parts.append(merged)
                        cur_text = []
                        cur_struck = None

                    for piece, is_struck in raw_pieces:
                        # Whitespace-only pieces stay attached to the
                        # surrounding non-struck run so that spaces between
                        # words are preserved without being marked struck.
                        if not piece.strip():
                            if cur_struck is True:
                                _flush_run()
                            cur_struck = False
                            cur_text.append(piece)
                            continue
                        if cur_struck is None or cur_struck == is_struck:
                            cur_struck = is_struck
                            cur_text.append(piece)
                        else:
                            _flush_run()
                            cur_struck = is_struck
                            cur_text.append(piece)
                    _flush_run()

                    combined = "".join(line_parts).strip()
                    if combined:
                        page_lines.append(combined)

            if page_lines:
                page_texts.append("\n".join(page_lines))

    if not page_texts:
        return []

    full_text = "\n\n".join(page_texts)

    full_text = _sanitize_charter_party_pdf_text(full_text)

    # Try numbered-clause splitting first
    clause_chunks = _presplit_on_clauses(full_text)
    if clause_chunks:
        logger.info(
            "[extractor] Fitz-presplit: %d clause chunks (strikethrough preserved)",
            len(clause_chunks),
        )
        return clause_chunks

    # No numbered structure — return as page batches so AI enrichment still
    # receives the full context (strikethrough markers intact throughout)
    total_words = len(full_text.split())
    logger.info(
        "[extractor] Fitz-presplit: no clause markers found, "
        "%d words as page batches (strikethrough preserved)",
        total_words,
    )

    if total_words <= _RAW_SINGLE_CHUNK_THRESHOLD:
        return [{"title": None, "body": full_text}]

    chunks: List[Dict] = []
    batch_lines: List[str] = []
    batch_words = 0
    for pt in page_texts:
        pw = len(pt.split())
        if batch_words + pw > _RAW_SINGLE_CHUNK_THRESHOLD and batch_lines:
            chunks.append({"title": None, "body": "\n\n".join(batch_lines)})
            batch_lines = []
            batch_words = 0
        batch_lines.append(pt)
        batch_words += pw
    if batch_lines:
        chunks.append({"title": None, "body": "\n\n".join(batch_lines)})
    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction — public entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf(file_obj: IO[bytes], document_category: Optional[str] = None) -> List[Dict]:
    """
    Extract chunks from a .pdf file.

    Dispatch order by document_category:

      1. _FITZ_PRESPLIT_CATEGORIES (e.g. charter_party)
             → _extract_pdf_fitz_presplit
             PyMuPDF strikethrough detection + numbered-clause splitting.
             Use when struck-out text is legally significant AND bold body
             text would confuse the font-based header detector.

      2. _RAW_EXTRACTION_CATEGORIES (e.g. fixture_recap, addendum)
             → _extract_pdf_raw
             pdfplumber full-text + numbered-clause splitting, AI-first.
             Use when font heuristics misfire and strikethrough is not needed.

      3. Everything else
             → _extract_pdf_fitz  (font-aware header detection)
             →   pdfplumber fallback (clause-regex only, no font info)
             →   fixed-size word chunks (last resort)
    """
    raw_bytes = file_obj.read()

    # Hybrid: PyMuPDF strikethrough + clause-based splitting
    if document_category in _FITZ_PRESPLIT_CATEGORIES:
        logger.info(
            f"[extractor] Using fitz-presplit (strikethrough + clause) "
            f"for category '{document_category}'"
        )
        try:
            chunks = _extract_pdf_fitz_presplit(raw_bytes)
            if chunks:
                return chunks
            logger.warning(
                "[extractor] Fitz-presplit returned no content — "
                "falling back to font-based extraction"
            )
        except ImportError:
            logger.warning(
                "[extractor] PyMuPDF not available for fitz-presplit — "
                "falling back to raw extraction"
            )

    # Raw: pdfplumber full-text, AI-first (no strikethrough detection)
    if document_category in _RAW_EXTRACTION_CATEGORIES:
        logger.info(
            f"[extractor] Using raw (AI-first) PDF extraction "
            f"for category '{document_category}'"
        )
        try:
            chunks = _extract_pdf_raw(raw_bytes)
            if chunks:
                return chunks
            logger.warning(
                "[extractor] Raw extraction returned no content — "
                "falling back to font-based extraction"
            )
        except ImportError:
            logger.warning(
                "[extractor] pdfplumber not available for raw extraction — "
                "falling back to font-based extraction"
            )

    try:
        chunks = _extract_pdf_fitz(raw_bytes)
    except ImportError:
        try:
            chunks = _extract_pdf_pdfplumber(raw_bytes)
        except ImportError:
            raise ImportError(
                "No PDF library found. Run: pip install PyMuPDF "
                "(or pip install pdfplumber as fallback)"
            )

    if not chunks:
        return [{"title": None, "body": ""}]

    return _split_into_chunks(chunks)


# ─────────────────────────────────────────────────────────────────────────────
# XLSX extraction — helpers
# ─────────────────────────────────────────────────────────────────────────────

# Footer row labels (lower-cased, with colon) — terminate the data section.
_XLSX_FOOTER_KEYS = frozenset({
    'date:', 'vessel:', "vessel's name:", 'master:', "master's name:",
    'chief mate:', "ch. mate's name:", 'chief officer:', 'voyage no.:',
    'date of inventory:',
})
_XLSX_REMARKS_KEYS = frozenset({'remarks:', 'remarks'})
_XLSX_PAGE_RE = re.compile(r'^page\s+\d+\s*(of|/)\s*\d+', re.IGNORECASE)


def _xlsx_get_raw_grid(ws):
    """
    Return a 1-indexed 2D list of cell values WITHOUT merge expansion, plus
    (max_row, max_col). grid[row][col], both 1-based; grid[0] is unused.
    """
    max_row = ws.max_row or 0
    max_col = ws.max_column or 0
    if not max_row or not max_col:
        return [], 0, 0
    grid = [[None] * (max_col + 1) for _ in range(max_row + 1)]
    for row in ws.iter_rows():
        for cell in row:
            grid[cell.row][cell.column] = cell.value
    return grid, max_row, max_col


def _xlsx_expand_merged(ws, raw_grid, max_row, max_col):
    """
    Return a copy of raw_grid with every merged-cell region filled with the
    anchor cell's value.  Used for composite header building and section-label
    detection (where fully-merged rows have one value repeated across all cols).
    """
    import copy
    grid = copy.deepcopy(raw_grid)
    for mr in ws.merged_cells.ranges:
        anchor = grid[mr.min_row][mr.min_col]
        for r in range(mr.min_row, mr.max_row + 1):
            for c in range(mr.min_col, mr.max_col + 1):
                grid[r][c] = anchor
    return grid


def _xlsx_fmt(v) -> str:
    """Format a cell value to a clean, newline-free string."""
    if v is None:
        return ""
    from datetime import datetime as _dt
    if isinstance(v, _dt):
        return v.strftime("%Y-%m-%d")
    return str(v).strip().replace('\n', ' ')


def _xlsx_find_header_rows(raw_grid, max_row, max_col):
    """
    Use the RAW (unexpanded) grid to locate the main header row.

    Scans rows 4–14 for the first row that has:
    - ≥4 non-None string cells AND ≥3 DISTINCT values
      (rules out merged single-value title rows like "Inventory of …")

    Then checks whether the immediately following row is a sub-header:
    - 2 ≤ cells < main row cells AND ≥2 distinct values

    Returns (main_row_idx, sub_row_idx_or_None).
    """
    for r in range(4, min(max_row + 1, 15)):
        row = raw_grid[r][1:max_col + 1]
        str_cells = [v for v in row if v is not None and isinstance(v, str) and v.strip()]
        if len(str_cells) >= 4 and len(set(str_cells)) >= 3:
            sub_row = None
            if r + 1 <= max_row:
                nxt = raw_grid[r + 1][1:max_col + 1]
                nxt_str = [v for v in nxt if v is not None and isinstance(v, str) and v.strip()]
                if 2 <= len(nxt_str) < len(str_cells) and len(set(nxt_str)) >= 2:
                    sub_row = r + 1
            return r, sub_row
    return None, None


def _xlsx_build_headers(exp_grid, main_row, sub_row, max_col) -> dict:
    """
    Build {col_idx: header_name} using the EXPANDED grid so that
    merge-spanned column values are available in all relevant columns.

    When a sub-row value differs from the main-row value → "Main (Sub)".
    E.g. main="Total quantity", sub="Good" → "Total quantity (Good)".
    """
    headers: dict = {}
    for c in range(1, max_col + 1):
        mv = exp_grid[main_row][c] if main_row else None
        sv = exp_grid[sub_row][c] if sub_row else None
        ms = str(mv).strip().replace('\n', ' ') if mv is not None else None
        ss = str(sv).strip().replace('\n', ' ') if sv is not None else None
        if ms and ss and ms != ss:
            headers[c] = f"{ms} ({ss})"
        elif ms:
            headers[c] = ms
        elif ss:
            headers[c] = ss
    return headers


def _xlsx_is_section_label(grid_row, max_col) -> bool:
    """
    True when the row represents an internal section/subsection header.

    Handles both:
    - A single non-None cell in column 1 (unmerged section label)
    - A fully-merged row where all non-None values are the same short string
      (e.g. the A8:P8 merged "Lashings" label row)

    Excludes footer keys, remarks labels, and "Page N of N" references.
    """
    non_none = [grid_row[c] for c in range(1, max_col + 1) if grid_row[c] is not None]
    if not non_none:
        return False
    distinct = set(str(v).strip() for v in non_none)
    if len(distinct) != 1:
        return False
    val_str = next(iter(distinct))
    if not val_str:
        return False
    lower = val_str.lower()
    if lower in _XLSX_FOOTER_KEYS:
        return False
    if lower.rstrip(':') in _XLSX_REMARKS_KEYS:
        return False
    if _XLSX_PAGE_RE.match(val_str):
        return False
    if len(val_str) > 80:
        return False
    return True


def _xlsx_get_vessel_name(wb) -> Optional[str]:
    """
    Scan all sheets for a "Vessel:" or "Vessel's name:" label row and return
    the first subsequent non-label cell value on the same row.
    """
    vessel_keys = {"vessel:", "vessel's name:"}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        raw_grid, max_row, max_col = _xlsx_get_raw_grid(ws)
        if max_row == 0:
            continue
        for r in range(1, max_row + 1):
            row = raw_grid[r]
            for c in range(1, max_col + 1):
                v = row[c]
                if v is None:
                    continue
                if str(v).strip().lower() in vessel_keys:
                    for nc in range(c + 1, max_col + 1):
                        nv = row[nc]
                        if nv is None:
                            continue
                        nv_str = str(nv).strip()
                        if nv_str and nv_str.lower() not in vessel_keys and nv_str.lower() not in _XLSX_FOOTER_KEYS:
                            return nv_str
    return None


def _xlsx_extract_front_page(
    sheet_name: str, exp_grid: list, max_row: int, max_col: int,
    vessel_name: Optional[str], chapter: Optional[str],
) -> Optional[Dict]:
    """
    Extract the Front Page tab as a labelled prose dump.
    Merged cell values are de-duplicated within each row.
    """
    lines: List[str] = []
    for r in range(1, max_row + 1):
        seen: set = set()
        vals: List[str] = []
        for c in range(1, max_col + 1):
            v = exp_grid[r][c]
            if v is None:
                continue
            vs = _xlsx_fmt(v)
            if vs and vs not in seen:
                seen.add(vs)
                vals.append(vs)
        if vals:
            lines.append("  ".join(vals))
    body = "\n".join(lines).strip()
    if not body:
        return None
    parts = [p for p in [vessel_name, chapter, sheet_name] if p]
    return {"title": " - ".join(parts) if parts else sheet_name, "body": body}


def _xlsx_extract_regular_sheet(
    sheet_name: str, raw_grid: list, exp_grid: list,
    max_row: int, max_col: int,
    vessel_name: Optional[str], chapter: Optional[str],
) -> List[Dict]:
    """
    Extract a regular data sheet into one chunk.

    Title: [Vessel Name] - [Chapter] - [Subchapter (tab name)]
    Body:
      [Section: X]                   ← subsection headers (fully-merged rows)
      Header: value  |  Header: value  ← data rows with composite column names
      ...
      Remarks: ...                   ← remarks block (if present)
      ---
      Date: ... | Vessel: ... | Master: ... | Chief Mate: ...
    """
    main_row, sub_row = _xlsx_find_header_rows(raw_grid, max_row, max_col)
    if main_row is None:
        # Fallback: dump all non-empty rows
        lines: List[str] = []
        for r in range(1, max_row + 1):
            seen: set = set()
            vals: List[str] = []
            for c in range(1, max_col + 1):
                v = exp_grid[r][c]
                if v is None:
                    continue
                vs = _xlsx_fmt(v)
                if vs and vs not in seen:
                    seen.add(vs)
                    vals.append(vs)
            if vals:
                lines.append("  ".join(vals))
        body = "\n".join(lines).strip()
        parts = [p for p in [vessel_name, chapter, sheet_name] if p]
        return [{"title": " - ".join(parts), "body": body}] if body else []

    headers = _xlsx_build_headers(exp_grid, main_row, sub_row, max_col)
    if not headers:
        return []

    data_start = (sub_row or main_row) + 1
    body_lines: List[str] = []
    footer_parts: List[str] = []
    remarks_lines: List[str] = []
    in_remarks = False

    for r in range(data_start, max_row + 1):
        row = exp_grid[r]

        if all(row[c] is None for c in range(1, max_col + 1)):
            continue

        # First non-None value in this row
        first_val: Optional[str] = None
        first_col: Optional[int] = None
        for c in range(1, max_col + 1):
            if row[c] is not None:
                first_val = str(row[c]).strip()
                first_col = c
                break
        if first_val is None:
            continue

        # ── Footer row ────────────────────────────────────────────────────────
        if first_val.lower() in _XLSX_FOOTER_KEYS:
            in_remarks = False
            kvs = [first_val]
            for c in range(first_col + 1, max_col + 1):
                nv = row[c]
                if nv is None:
                    continue
                nv_str = str(nv).strip()
                if nv_str and nv_str.lower() not in _XLSX_FOOTER_KEYS and nv_str.lower() != first_val.lower():
                    kvs.append(_xlsx_fmt(nv))
                    break
            footer_parts.append(" ".join(kvs))
            continue

        # ── Remarks row ───────────────────────────────────────────────────────
        if first_val.lower().rstrip(':') in _XLSX_REMARKS_KEYS:
            in_remarks = True
            inline: List[str] = []
            for c in range(first_col + 1, max_col + 1):
                nv = row[c]
                if nv is not None:
                    vs = _xlsx_fmt(nv)
                    if vs:
                        inline.append(vs)
            remarks_lines.append(
                "Remarks: " + " | ".join(inline) if inline else "Remarks:"
            )
            continue

        if in_remarks:
            seen2: set = set()
            vals2: List[str] = []
            for c in range(1, max_col + 1):
                v = row[c]
                if v is None:
                    continue
                vs = _xlsx_fmt(v)
                if vs and vs not in seen2:
                    seen2.add(vs)
                    vals2.append(vs)
            if vals2:
                remarks_lines.append("  ".join(vals2))
            continue

        # ── Section label (fully merged or single-cell row) ───────────────────
        if _xlsx_is_section_label(row, max_col):
            body_lines.append(f"\n[Section: {first_val}]")
            continue

        # ── Data row → composite header: value pairs ──────────────────────────
        pairs: List[str] = []
        for c in range(1, max_col + 1):
            v = row[c]
            if v is None:
                continue
            vs = _xlsx_fmt(v)
            if vs and vs != '-':
                pairs.append(f"{headers.get(c, f'Col{c}')}: {vs}")
        if pairs:
            body_lines.append("  |  ".join(pairs))

    # ── Assemble ──────────────────────────────────────────────────────────────
    body_parts: List[str] = []
    if body_lines:
        body_parts.append("\n".join(body_lines).strip("\n"))
    if remarks_lines:
        body_parts.append("\n".join(remarks_lines))
    if footer_parts:
        body_parts.append("---\n" + " | ".join(footer_parts))

    body = "\n\n".join(body_parts).strip()
    if not body:
        return []

    title_parts = [p for p in [vessel_name, chapter, sheet_name] if p]
    return [{"title": " - ".join(title_parts), "body": body}]


# ─────────────────────────────────────────────────────────────────────────────
# XLSX extraction — public entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_xlsx(file_obj: IO[bytes]) -> List[Dict]:
    """
    Extract chunks from an .xlsx/.xls maritime inventory or equipment file.

    Structure detected per sheet:
    - Front Page tab: dumped as labelled prose (vessel info, master, instructions).
    - Regular tabs:   one chunk each.
        Title: [Vessel Name] - [Chapter] - [Subchapter (tab name)]
        Body:  [Section: X] labels  +  structured data rows (Header: value |…)
               +  Remarks block  +  '---' footer line

    Composite column headers are built by merging a two-row header band:
    e.g. "Total quantity" / "Good" → "Total quantity (Good)".
    Fully-merged section label rows (e.g. A8:P8 "Lashings") are preserved
    as [Section: Lashings] markers in the body.

    Falls back to xlrd for legacy .xls files if openpyxl cannot load the file.
    """
    import io as _io

    raw = file_obj.read()

    try:
        import openpyxl
        wb = openpyxl.load_workbook(_io.BytesIO(raw), data_only=True)
    except Exception as e:
        try:
            return _extract_xlsx_xlrd_fallback(raw)
        except ImportError:
            raise ImportError(
                "No spreadsheet library found. Run: pip install openpyxl"
            ) from e

    vessel_name = _xlsx_get_vessel_name(wb)
    chunks: List[Dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        raw_grid, max_row, max_col = _xlsx_get_raw_grid(ws)
        if max_row == 0 or max_col == 0:
            continue

        exp_grid = _xlsx_expand_merged(ws, raw_grid, max_row, max_col)

        # Chapter title: first non-None value in row 1 (expansion covers merged title)
        chapter: Optional[str] = None
        for c in range(1, max_col + 1):
            if exp_grid[1][c] is not None:
                chapter = _xlsx_fmt(exp_grid[1][c])
                break

        norm = sheet_name.lower().replace(" ", "").replace("_", "")
        if norm in ("frontpage", "coverpage", "titlepage"):
            chunk = _xlsx_extract_front_page(
                sheet_name, exp_grid, max_row, max_col, vessel_name, chapter
            )
            if chunk:
                chunks.append(chunk)
        else:
            chunks.extend(
                _xlsx_extract_regular_sheet(
                    sheet_name, raw_grid, exp_grid, max_row, max_col,
                    vessel_name, chapter,
                )
            )

    if not chunks:
        chunks = [{"title": None, "body": "No data found in spreadsheet."}]

    return _split_into_chunks(chunks)


def _extract_xlsx_xlrd_fallback(raw: bytes) -> List[Dict]:
    """
    Basic .xls extraction via xlrd (legacy format, no merged-cell awareness).
    Returns one chunk per sheet as pipe-separated rows.
    """
    import xlrd
    wb = xlrd.open_workbook(file_contents=raw)
    chunks: List[Dict] = []
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        lines: List[str] = []
        for rx in range(ws.nrows):
            row_vals = [
                str(ws.cell(rx, cx).value).strip()
                for cx in range(ws.ncols)
                if str(ws.cell(rx, cx).value).strip()
            ]
            if row_vals:
                lines.append("  |  ".join(row_vals))
        body = "\n".join(lines).strip()
        if body:
            chunks.append({"title": sheet_name, "body": body})
    return _split_into_chunks(chunks or [{"title": None, "body": "No data found in spreadsheet."}])


# ─────────────────────────────────────────────────────────────────────────────
# Public interface
# ─────────────────────────────────────────────────────────────────────────────

def extract(
    file_obj: IO[bytes],
    filename: str,
    document_category: Optional[str] = None,
) -> List[Dict]:
    """
    Route to the correct extractor based on file extension.

    Args:
        file_obj:           Readable binary stream of the uploaded file.
        filename:           Original filename (used for extension detection).
        document_category:  Dossier section slug (e.g. "fixture_recap").
                            When provided, PDF extraction uses the appropriate
                            strategy for that document type — raw full-text
                            extraction for narrative/spec docs so that AI
                            enrichment receives complete context.

    Returns:
        List of {"title": str|None, "body": str} dicts, ready for DB insertion.
    Raises:
        ValueError: unsupported file type
        ImportError: missing optional dependency (pdfplumber)
    """
    ext = filename.rsplit(".", 1)[-1].lower()
    logger.info(
        f"[extractor] Extracting {filename} (type={ext}, category={document_category})"
    )

    if ext == "docx":
        return extract_docx(file_obj)
    elif ext == "pdf":
        return extract_pdf(file_obj, document_category=document_category)
    elif ext in ("xlsx", "xls"):
        return extract_xlsx(file_obj)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
