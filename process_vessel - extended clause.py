#!/usr/bin/env python3
"""
Vessel Data Extractor - Updated for new folder structure
Extracts vessel information from DOCX and PDF files in source/vessels folder
and creates separate output files for each vessel.
"""

import os
import json
import csv
import logging
import re
from pathlib import Path
from docx import Document
import PyPDF2
import io
from typing import Dict, List, Any
import requests
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

class VesselDataExtractor:
    def __init__(self, source_dir: str = 'source/vessels'):
        """
        Initialize the VesselDataExtractor.
        
        Args:
            source_dir: Directory containing vessel documents
        """
        self.source_dir = source_dir
        self.extracted_data = {}
        self.upload_enabled = True
        # Load existing attachment ID mappings
        self.attachment_mappings = self.load_attachment_mappings()
        # Cache for strikethrough detection and text processing
        self._strikethrough_cache = {}
        self._text_cache = {}
        
    def _remove_strikethrough_text(self, text: str) -> str:
        """Process strikethrough text patterns from the given text by wrapping them in [STRIKETHROUGH] tags."""
        # Common strike-through patterns found in DOCX files
        import re
        
        # Pattern 1: [[STRIKE ... STRIKE]] - common DOCX strike-through format
        text = re.sub(r'\[\[STRIKE\s+(.*?)\s+STRIKE\]\]', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 2: QUOTE STRIKE-THROUGH ... UNQUOTE STRIKE-THROUGH
        text = re.sub(r'QUOTE STRIKE-THROUGH\s+(.*?)\s+UNQUOTE STRIKE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 3: Simple strike-through markers
        text = re.sub(r'STRIKE\s+(.*?)\s+STRIKE', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 4: Standalone STRIKE markers
        text = re.sub(r'\bSTRIKE\b', '', text, flags=re.IGNORECASE)
        
        # Pattern 5: QUOTE-THROUGH ... UNQUOTE-THROUGH
        text = re.sub(r'QUOTE-THROUGH\s+(.*?)\s+UNQUOTE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 6: Any remaining incomplete strike-through blocks
        text = re.sub(r'\[\[STRIKE.*?\]\]', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'QUOTE STRIKE-THROUGH.*?UNQUOTE STRIKE-THROUGH', '', text, flags=re.DOTALL | re.IGNORECASE)
        
        return text.strip()

    def _inspect_docx_for_strikethrough(self, file_path: str):
        """Inspect the DOCX file's XML structure to find strikethrough text."""
        try:
            from docx import Document
            import zipfile
            import xml.etree.ElementTree as ET
            
            # Method 1: Use python-docx to inspect
            doc = Document(file_path)
            strike_count = 0
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                        rPr = run._element.rPr
                        if rPr is not None:
                            # Check for strike elements
                            strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                            dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                            
                            # Check for strike attributes
                            strike_attr = rPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                            dstrike_attr = rPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                            
                            if (strike_elem is not None or dstrike_elem is not None or 
                                strike_attr in ['true', '1', 'on', 'single', 'double'] or 
                                dstrike_attr in ['true', '1', 'on', 'single', 'double']):
                                strike_count += 1
            
        except Exception as e:
            logging.error(f"Error inspecting DOCX for strikethrough: {e}")

    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to group consecutive strikethrough runs and wrap them in [STRIKETHROUGH]...[/STRIKETHROUGH] tags."""
        parts = []
        in_strike = False
        buffer = []
        last_run_ended_with_space = False
        
        for run in paragraph.runs:
            # Get the run's text
            run_text = run.text
            
            # Skip empty runs
            if not run_text:
                continue
            
            # Check for strikethrough in multiple ways to ensure we catch all formats
            is_strike = False
            
            # Method 1: Check XML structure (most reliable for MS Word)
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Check for strike element
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    # Check for strike attribute
                    strike_attr = rPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_attr = rPr.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if (strike_elem is not None or dstrike_elem is not None or 
                        strike_attr in ['true', '1', 'on', 'single', 'double'] or 
                        dstrike_attr in ['true', '1', 'on', 'single', 'double']):
                        is_strike = True
            
            # Method 2: Check font property (backup method)
            if not is_strike and hasattr(run, 'font') and run.font:
                if hasattr(run.font, 'strike') and run.font.strike:
                    is_strike = True
                elif hasattr(run.font, '_element') and run.font._element is not None:
                    strike_val = run.font._element.get('strike')
                    dstrike_val = run.font._element.get('dstrike')
                    if (strike_val in ['true', '1', 'on', 'single', 'double'] or 
                        dstrike_val in ['true', '1', 'on', 'single', 'double']):
                        is_strike = True
            
            # Method 3: Check run style (some documents use styles for strikethrough)
            if not is_strike and hasattr(run, 'style') and run.style:
                style_id = run.style.style_id.lower() if hasattr(run.style, 'style_id') else ''
                style_name = run.style.name.lower() if hasattr(run.style, 'name') else ''
                if ('strikethrough' in style_id or 'strike' in style_id or 
                    'strikethrough' in style_name or 'strike' in style_name):
                    is_strike = True
            
            # Method 4: Check paragraph style for strikethrough
            if not is_strike and hasattr(paragraph, 'style') and paragraph.style:
                para_style_id = paragraph.style.style_id.lower() if hasattr(paragraph.style, 'style_id') else ''
                para_style_name = paragraph.style.name.lower() if hasattr(paragraph.style, 'name') else ''
                if ('strikethrough' in para_style_id or 'strike' in para_style_id or 
                    'strikethrough' in para_style_name or 'strike' in para_style_name):
                    is_strike = True
            
            # Add space between runs if needed
            if parts and not last_run_ended_with_space and not run_text.startswith(' '):
                if in_strike:
                    buffer.append(' ')
                else:
                    parts.append(' ')
            
            # Handle the run based on strikethrough status
            if is_strike:
                if not in_strike:
                    # Starting a new strikethrough block
                    if buffer:
                        parts.append(''.join(buffer))
                        buffer = []
                    in_strike = True
                buffer.append(run_text)
            else:
                if in_strike:
                    # Ending a strikethrough block
                    if buffer:
                        strikethrough_text = ''.join(buffer)
                        parts.append(f"[STRIKETHROUGH]{strikethrough_text}[/STRIKETHROUGH]")
                        buffer = []
                    in_strike = False
                parts.append(run_text)
            
            # Track if this run ends with a space
            last_run_ended_with_space = run_text.endswith(' ')
        
        # Handle any remaining text
        if buffer:
            if in_strike:
                strikethrough_text = ''.join(buffer)
                parts.append(f"[STRIKETHROUGH]{strikethrough_text}[/STRIKETHROUGH]")
            else:
                parts.append(''.join(buffer))
        
        # Join all parts and clean up any artifacts
        result = ''.join(parts).strip()
        return result
    
    def _process_document_once(self, file_path: str) -> tuple:
        """
        PERFORMANCE OPTIMIZATION: Process document only once to extract both text and style information.
        Returns: (full_text, contract_section_text, paragraph_styles)
        """
        # Check cache first
        cache_key = f"{file_path}_{os.path.getmtime(file_path)}"
        if cache_key in self._text_cache:
            return self._text_cache[cache_key]
            
        try:
            doc = Document(file_path)
            text_parts = []
            paragraph_styles = {}
            contract_section_text_parts = []
            
            # Single pass through all paragraphs
            in_contract_section = False
            contract_started = False
            
            for paragraph in doc.paragraphs:
                # Process strikethrough only once per paragraph
                text = self._process_paragraph_for_strikethrough_optimized(paragraph)
                if text:
                    text_parts.append(text)
                    
                    # Check for contract details section
                    text_stripped = text.strip()
                    if re.match(r'^2\.?\s*contract details', text_stripped, re.IGNORECASE):
                        contract_started = True
                        in_contract_section = True
                        continue
                    elif contract_started and re.match(r'^3\.?\s*delivery figures', text_stripped, re.IGNORECASE):
                        in_contract_section = False
                        
                    # Store style information and contract text in one pass
                    if contract_started and in_contract_section:
                        contract_section_text_parts.append(text)
                        # Store style information for subheader detection
                        is_heading1 = paragraph.style and paragraph.style.name and paragraph.style.name.lower().startswith('heading 1')
                        # Only store True values, don't overwrite existing True with False
                        if is_heading1 or text_stripped not in paragraph_styles:
                            paragraph_styles[text_stripped] = is_heading1
            
            # Process tables only once
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- TABLE {i+1} ---")
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = []
                        for p in cell.paragraphs:
                            processed_text = self._process_paragraph_for_strikethrough_optimized(p)
                            if processed_text:
                                cell_text.append(processed_text)
                        row_cells.append(' '.join(cell_text).strip())
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("--- END TABLE ---\n")
            
            # Cache results
            full_text = '\n\n'.join(text_parts)
            contract_text = '\n'.join(contract_section_text_parts)
            result = (full_text, contract_text, paragraph_styles)
            self._text_cache[cache_key] = result
            
            return result
            
        except Exception as e:
            logging.error(f"Error processing document {file_path}: {e}")
            return "", "", {}
    
    def _process_paragraph_for_strikethrough_optimized(self, paragraph) -> str:
        """
        PERFORMANCE OPTIMIZATION: Optimized strikethrough processing with caching.
        """
        # Create a cache key based on paragraph content and style
        para_key = f"{paragraph.text}_{id(paragraph)}"
        if para_key in self._strikethrough_cache:
            return self._strikethrough_cache[para_key]
            
        parts = []
        in_strike = False
        buffer = []
        
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
            
            # Optimized strikethrough detection - check only essential properties
            is_strike = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Fast check for strike elements
                    if (rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike') is not None or
                        rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike') is not None):
                        is_strike = True
            
            if is_strike:
                if not in_strike:
                    in_strike = True
                buffer.append(run_text)
            else:
                if in_strike:
                    if buffer:
                        parts.append(f"[STRIKETHROUGH]{' '.join(buffer)}[/STRIKETHROUGH]")
                        buffer = []
                    in_strike = False
                parts.append(run_text)
        
        # Handle remaining strikethrough buffer
        if in_strike and buffer:
            parts.append(f"[STRIKETHROUGH]{' '.join(buffer)}[/STRIKETHROUGH]")
        
        result = ''.join(parts)
        self._strikethrough_cache[para_key] = result
        return result

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file, preserving strikethrough formatting."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs, using the strikethrough-aware function
            for paragraph in doc.paragraphs:
                text = self._process_paragraph_for_strikethrough(paragraph)
                if text:
                    text_parts.append(text)
            
            # Extract tables and add them to the text
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- TABLE {i+1} ---")
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # For each cell, process all paragraphs for strikethrough
                        cell_text = []
                        for p in cell.paragraphs:
                            processed_text = self._process_paragraph_for_strikethrough(p)
                            if processed_text:
                                cell_text.append(processed_text)
                        row_cells.append(' '.join(cell_text).strip())
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("--- END TABLE ---\n")
            
            # Join all parts with double newlines for better section separation
            return '\n\n'.join(text_parts)
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logging.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return ""

    def _extract_vessel_name(self, text: str, file_path: str = "") -> str:
        """Extract vessel name from the text content or filename."""
        patterns = [
            r'Vessel name:\s*([^\n]+)',
            r'Name:\s*([^\n]+)',
            r'MV\s+([^\s\n]+)',
            r'([A-Z]{2,}\s+[A-Z]{2,})',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                vessel_name = match.group(1).strip()
                vessel_name = re.sub(r'[^\w\s-]', '', vessel_name)
                vessel_name = vessel_name.replace(' ', '_').upper()
                if not vessel_name.startswith('MV_'):
                    vessel_name = f'MV_{vessel_name}'
                return vessel_name
        
        # If no vessel name found in text, try to extract from filename
        if file_path:
            filename = os.path.basename(file_path)
            # Remove file extension
            filename_without_ext = os.path.splitext(filename)[0]
            
            # Check for known vessel names in filename
            for name in ['ARA_ROTTERDAM', 'FRANZISKA', 'ALASKABORG', 'HC_EVA_MARIE']:
                if name in filename.upper():
                    return f'MV_{name}'
            
            # If no known vessel name found, create a vessel name from the filename
            # Clean the filename and convert to vessel name format
            vessel_name = re.sub(r'[^\w\s-]', '', filename_without_ext)
            vessel_name = vessel_name.replace(' ', '_').upper()
            if vessel_name and not vessel_name.startswith('MV_'):
                vessel_name = f'MV_{vessel_name}'
                return vessel_name
        
        return ""

    def _extract_vessel_details(self, text: str, doc: Document = None) -> Dict[str, Any]:
        import re
        details = {
            'vessel_details': {'content': ''},
            'contract_details': {'structured': []},  # Initialize with structured list
            'delivery_figures': {'content': ''},
            'speed_and_consumption': {'content': ''},
            'lifting_equipment': {'content': ''},
            'hseq_documentation': {'content': ''}
        }
        # Regex patterns for section headers
        section_patterns = [
            ('vessel_details', re.compile(r'^1\.?\s*vessel details', re.IGNORECASE)),
            ('contract_details', re.compile(r'^2\.?\s*contract details', re.IGNORECASE)),
            ('delivery_figures', re.compile(r'^3\.?\s*delivery figures', re.IGNORECASE)),
            ('speed_and_consumption', re.compile(r'^4\.?\s*speed and consumption', re.IGNORECASE)),
            ('lifting_equipment', re.compile(r'^5\.?\s*lifting equipment', re.IGNORECASE)),
            ('hseq_documentation', re.compile(r'^6\.?\s*hseq documentation', re.IGNORECASE)),
        ]
        # First, find all strikethrough blocks and replace them with unique markers
        strikethrough_blocks = []
        def save_strikethrough(match):
            block_id = f"__STRIKETHROUGH_{len(strikethrough_blocks)}__"
            strikethrough_blocks.append(match.group(1))
            return block_id
        # Replace strikethrough blocks with markers
        processed_text = re.sub(r'\[STRIKETHROUGH\](.*?)\[/STRIKETHROUGH\]', save_strikethrough, text, flags=re.DOTALL)
        # Split into lines and find section boundaries
        lines = processed_text.split('\n')
        section_indices = {}
        for i, line in enumerate(lines):
            for key, pattern in section_patterns:
                if key not in section_indices and pattern.match(line.strip()):
                    section_indices[key] = i
        # Sort by appearance
        sorted_sections = sorted(section_indices.items(), key=lambda x: x[1])
        # Extract content for each section
        for idx, (key, start) in enumerate(sorted_sections):
            end = sorted_sections[idx + 1][1] if idx + 1 < len(sorted_sections) else len(lines)
            section_lines = lines[start + 1:end]
            section_text = '\n'.join(section_lines).strip()
            # Restore strikethrough blocks
            for i, block in enumerate(strikethrough_blocks):
                section_text = section_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            if key == 'contract_details' and doc is not None:
                details[key]['structured'] = self._extract_contract_details_structured(doc, section_text)
            else:
                details[key]['content'] = section_text
        # Vessel Specifications extraction (simple: from 'Vessel Specifications' to 'Contract Details')
        vessel_spec_content = []
        in_vessel_spec = False
        for line in lines:
            if not in_vessel_spec and re.search(r'vessel\s*specifications?', line, re.IGNORECASE):
                in_vessel_spec = True
            if in_vessel_spec:
                # Stop collecting BEFORE adding the Contract Details line
                if re.search(r'contract\s*details', line, re.IGNORECASE):
                    break
                vessel_spec_content.append(line)
        if vessel_spec_content:
            vessel_spec_text = '\n'.join(vessel_spec_content).strip()
            for i, block in enumerate(strikethrough_blocks):
                vessel_spec_text = vessel_spec_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            details['vessel_specifications'] = {'content': vessel_spec_text}
        # Extract the speed_and_consumption table (all rows)
        table_lines = []
        in_table = False
        for line in lines:
            if '--- TABLE' in line and '---' in line:
                in_table = True
                continue
            elif '--- END TABLE ---' in line:
                in_table = False
                break
            elif in_table and 'speed (knots)' in line.lower() and '|' in line:
                table_lines.append(line.strip())
            elif in_table and '|' in line:
                table_lines.append(line.strip())
        if table_lines:
            table_text = '\n'.join(table_lines)
            # Restore strikethrough blocks in table
            for i, block in enumerate(strikethrough_blocks):
                table_text = table_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            details['speed_and_consumption']['content'] = table_text
        return details

    def _extract_vessel_details_optimized(self, text: str, contract_text: str, paragraph_styles: dict) -> Dict[str, Any]:
        """
        PERFORMANCE OPTIMIZATION: Extract vessel details using pre-processed data.
        """
        import re
        details = {
            'vessel_details': {'content': ''},
            'contract_details': {'structured': []},
            'delivery_figures': {'content': ''},
            'speed_and_consumption': {'content': ''},
            'lifting_equipment': {'content': ''},
            'hseq_documentation': {'content': ''}
        }
        
        # Pre-compiled regex patterns for better performance
        section_patterns = [
            ('vessel_details', re.compile(r'^1\.?\s*vessel details', re.IGNORECASE)),
            ('contract_details', re.compile(r'^2\.?\s*contract details', re.IGNORECASE)),
            ('delivery_figures', re.compile(r'^3\.?\s*delivery figures', re.IGNORECASE)),
            ('speed_and_consumption', re.compile(r'^4\.?\s*speed and consumption', re.IGNORECASE)),
            ('lifting_equipment', re.compile(r'^5\.?\s*lifting equipment', re.IGNORECASE)),
            ('hseq_documentation', re.compile(r'^6\.?\s*hseq documentation', re.IGNORECASE)),
        ]
        
        # First, find all strikethrough blocks and replace them with unique markers
        strikethrough_blocks = []
        def save_strikethrough(match):
            block_id = f"__STRIKETHROUGH_{len(strikethrough_blocks)}__"
            strikethrough_blocks.append(match.group(1))
            return block_id
        
        # Replace strikethrough blocks with markers
        processed_text = re.sub(r'\[STRIKETHROUGH\](.*?)\[/STRIKETHROUGH\]', save_strikethrough, text, flags=re.DOTALL)
        
        # Split into lines and find section boundaries
        lines = processed_text.split('\n')
        section_indices = {}
        for i, line in enumerate(lines):
            for key, pattern in section_patterns:
                if key not in section_indices and pattern.match(line.strip()):
                    section_indices[key] = i
        
        # Sort by appearance
        sorted_sections = sorted(section_indices.items(), key=lambda x: x[1])
        
        # Extract content for each section
        for idx, (key, start) in enumerate(sorted_sections):
            end = sorted_sections[idx + 1][1] if idx + 1 < len(sorted_sections) else len(lines)
            section_lines = lines[start + 1:end]
            section_text = '\n'.join(section_lines).strip()
            
            # Restore strikethrough blocks
            for i, block in enumerate(strikethrough_blocks):
                section_text = section_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            
            if key == 'contract_details':
                # Use pre-processed contract data
                details[key]['structured'] = self._extract_contract_details_structured(None, contract_text, paragraph_styles)
            else:
                details[key]['content'] = section_text
        
        # Extract vessel specifications
        vessel_spec_content = []
        in_vessel_spec = False
        for line in lines:
            if not in_vessel_spec and re.search(r'vessel\s*specifications?', line, re.IGNORECASE):
                in_vessel_spec = True
            if in_vessel_spec:
                if re.search(r'contract\s*details', line, re.IGNORECASE):
                    break
                vessel_spec_content.append(line)
        
        if vessel_spec_content:
            vessel_spec_text = '\n'.join(vessel_spec_content).strip()
            for i, block in enumerate(strikethrough_blocks):
                vessel_spec_text = vessel_spec_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            details['vessel_specifications'] = {'content': vessel_spec_text}
        
        # Extract speed and consumption table
        table_lines = []
        in_table = False
        for line in lines:
            if '--- TABLE' in line and '---' in line:
                in_table = True
                continue
            elif '--- END TABLE ---' in line:
                in_table = False
                break
            elif in_table and 'speed (knots)' in line.lower() and '|' in line:
                table_lines.append(line.strip())
            elif in_table and '|' in line:
                table_lines.append(line.strip())
        
        if table_lines:
            table_text = '\n'.join(table_lines)
            for i, block in enumerate(strikethrough_blocks):
                table_text = table_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            details['speed_and_consumption']['content'] = table_text
        
        return details

    def _extract_contract_details_structured(self, doc: Document, preprocessed_text: str = None, paragraph_styles: dict = None):
        """Extract Contract Details as a structured list (numbered, lettered, bulleted), and tag Addendums, Fixture Recap, and Vessel Specifications."""
        import re
        in_contract_section = False
        items = []
        current_content = []
        special_section = None
        special_content = []
        special_type = None

        def is_heading(para, levels=(1,2,3,4)):
            if para.style and para.style.name:
                for lvl in levels:
                    if para.style.name.lower().startswith(f'heading {lvl}'):
                        return True
            return False

        def is_addendum_section(text):
            return re.search(r'2\.?\s*1\.?\s*addendum', text, re.IGNORECASE) or (re.search(r'addendum', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_fixture_recap_section(text):
            return re.search(r'2\.?\s*2\.?\s*fixture\s*recap', text, re.IGNORECASE) or (re.search(r'fixture\s*recap', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_charter_party_section(text):
            return re.search(r'2\.?\s*3\.?\s*charter\s*party', text, re.IGNORECASE) or (re.search(r'charter\s*party', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        # PERFORMANCE OPTIMIZATION: Use preprocessed text and styles when available
        if preprocessed_text and paragraph_styles:
            lines = preprocessed_text.split('\n')
            in_contract_section = True
            current_subheader = None  # Track current subheader for contextual tagging
            
            for line in lines:
                text = line.strip()
                if not text:
                    continue
                    
                # Stop at next major section
                if re.match(r'^3\.?\s*delivery figures', text, re.IGNORECASE):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                        special_section = None
                        special_content = []
                        special_type = None
                    break

                # Check for section transitions
                if is_addendum_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                    special_section = 'addendums'
                    special_type = 'addendums'
                    special_content = [text]
                    current_subheader = None  # Reset subheader context in special sections
                    continue
                elif is_fixture_recap_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                    special_section = 'fixture_recap'
                    special_type = 'fixture_recap'
                    special_content = [text]
                    current_subheader = None  # Reset subheader context in special sections
                    continue
                elif is_charter_party_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                        special_section = None
                        special_content = []
                        special_type = None
                    items.append({
                        "type": "text",
                        "content": text,
                        "tag": "CLAUSE"
                    })
                    current_subheader = None  # Reset subheader context
                    continue

                # If in a special section, collect content
                if special_section:
                    special_content.append(text)
                    continue

                # Check for subheaders using pre-processed style information
                is_heading1_style = paragraph_styles.get(text, False)
                is_numbered_heading = re.match(r'^\d+\.\s*[A-Z]', text)
                
                if is_heading1_style or is_numbered_heading:
                    if current_content:
                        items.extend(current_content)
                        current_content = []
                    
                    # Update current subheader for contextual tagging
                    current_subheader = text
                    
                    # Create contextual tag for the subheader
                    clean_heading = re.sub(r'^\d+\.\s*', '', current_subheader)  # Remove leading numbers
                    clean_heading = re.sub(r'[^\w\s-]', '', clean_heading)  # Remove special chars except dash
                    clean_heading = clean_heading.strip()
                    contextual_tag = f"CLAUSE - {clean_heading}" if clean_heading else "CLAUSE"
                    
                    items.append({
                        "type": "subheader",
                        "content": text,
                        "tag": contextual_tag  # Use contextual tag for subheaders
                    })
                else:
                    # Regular text - clean up spacing and apply contextual tag
                    text = re.sub(r'\s+', ' ', text)
                    text = re.sub(r'\s+([.,;:])', r'\1', text)
                    
                    # Create contextual tag with subheader
                    tag = "CLAUSE"
                    if current_subheader:
                        # Clean the subheader for tag use (remove numbers, special chars)
                        clean_subheader = re.sub(r'^\d+\.\s*', '', current_subheader)  # Remove leading numbers
                        clean_subheader = re.sub(r'[^\w\s-]', '', clean_subheader)     # Remove special chars except dash
                        clean_subheader = clean_subheader.strip()
                        if clean_subheader:
                            tag = f"CLAUSE - {clean_subheader}"
                    
                    current_content.append({
                        "type": "text",
                        "content": text,
                        "tag": tag
                    })

            # Flush any remaining content
            if current_content:
                items.extend(current_content)
            if special_section:
                items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})

        else:
            # Fallback to original document processing (should rarely be used now)
            current_subheader = None  # Track current subheader for contextual tagging
            
            for para in doc.paragraphs:
                text = self._process_paragraph_for_strikethrough(para).strip()
                if not text:
                    continue
                if re.match(r'^2\.?\s*contract details', text, re.IGNORECASE):
                    in_contract_section = True
                    continue
                if not in_contract_section:
                    continue
                if re.match(r'^3\.?\s*delivery figures', text, re.IGNORECASE):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                        special_section = None
                        special_content = []
                        special_type = None
                    break

                # Apply same logic as preprocessed text path
                if is_addendum_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                    special_section = 'addendums'
                    special_type = 'addendums'
                    special_content = [text]
                    current_subheader = None  # Reset subheader context in special sections
                    continue
                elif is_fixture_recap_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                    special_section = 'fixture_recap'
                    special_type = 'fixture_recap'
                    special_content = [text]
                    current_subheader = None  # Reset subheader context in special sections
                    continue
                elif is_charter_party_section(text):
                    if special_section:
                        items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})
                        special_section = None
                        special_content = []
                        special_type = None
                    items.append({
                        "type": "text",
                        "content": text,
                        "tag": "CLAUSE"
                    })
                    current_subheader = None  # Reset subheader context
                    continue

                if special_section:
                    special_content.append(text)
                    continue

                is_heading1 = para.style and para.style.name and para.style.name.lower().startswith('heading 1')
                if is_heading1:
                    if current_content:
                        items.extend(current_content)
                        current_content = []
                    
                    # Update current subheader for contextual tagging
                    current_subheader = text
                    
                    # Create contextual tag for the subheader
                    clean_heading = re.sub(r'^\d+\.\s*', '', current_subheader)  # Remove leading numbers
                    clean_heading = re.sub(r'[^\w\s-]', '', clean_heading)  # Remove special chars except dash
                    clean_heading = clean_heading.strip()
                    contextual_tag = f"CLAUSE - {clean_heading}" if clean_heading else "CLAUSE"
                    
                    items.append({
                        "type": "subheader",
                        "content": text,
                        "tag": contextual_tag  # Use contextual tag for subheaders
                    })
                else:
                    text = re.sub(r'\s+', ' ', text)
                    text = re.sub(r'\s+([.,;:])', r'\1', text)
                    
                    # Create contextual tag with subheader
                    tag = "CLAUSE"
                    if current_subheader:
                        # Clean the subheader for tag use (remove numbers, special chars)
                        clean_subheader = re.sub(r'^\d+\.\s*', '', current_subheader)  # Remove leading numbers
                        clean_subheader = re.sub(r'[^\w\s-]', '', clean_subheader)     # Remove special chars except dash
                        clean_subheader = clean_subheader.strip()
                        if clean_subheader:
                            tag = f"CLAUSE - {clean_subheader}"
                    
                    current_content.append({
                        "type": "text",
                        "content": text,
                        "tag": tag
                    })

            if current_content:
                items.extend(current_content)
            if special_section:
                items.append({"type": special_type, "content": '\n'.join(special_content), "tag": special_type.upper()})

        return items

    def process_files(self) -> Dict[str, Dict[str, Any]]:
        """Process all vessel files and extract data."""
        vessels_data = {}
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return vessels_data
        
        # Filter out Microsoft Word temporary files (starting with ~$) and only include .docx files
        docx_files = [f for f in os.listdir(self.source_dir) 
                     if f.lower().endswith('.docx') and not f.startswith('~$')]
        logging.info(f"Found {len(docx_files)} DOCX files to process: {docx_files}")
        
        for filename in docx_files:
            file_path = os.path.join(self.source_dir, filename)
            logging.info(f"Processing {filename}")
            
            try:
                # Use the optimized document processing function
                full_text, contract_text, paragraph_styles = self._process_document_once(file_path)
                
                vessel_name = self._extract_vessel_name(full_text, filename)
                
                if vessel_name:
                    vessel_details = self._extract_vessel_details_optimized(full_text, contract_text, paragraph_styles)
                    vessels_data[vessel_name] = vessel_details
                    logging.info(f"Successfully extracted data for {vessel_name}")
                else:
                    logging.warning(f"Could not extract vessel name from {filename}")
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
        
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels: {list(vessels_data.keys())}")
        return vessels_data

    def save_outputs(self, vessels_data: Dict[str, Dict[str, Any]], output_dir: str = 'output/vessels'):
        """Save vessel data to JSON and Markdown files."""
        os.makedirs(output_dir, exist_ok=True)
        logging.info(f"Saving outputs for {len(vessels_data)} vessels to {output_dir}")
        
        def convert_strikethrough_to_md(text):
            """Convert [STRIKETHROUGH] tags to markdown strikethrough syntax."""
            if not text or not isinstance(text, str):
                return text
            # Replace [STRIKETHROUGH]...[/STRIKETHROUGH] with ~~...~~
            text = re.sub(r'\[STRIKETHROUGH\](.*?)\[/STRIKETHROUGH\]', r'~~\1~~', text)
            return text
        
        def format_item_content(items, indent_level=0):
            """Format items content with proper indentation and clause tags, including subheader name in the clause tag and as the first line inside the clause. Also handle special tags for Addendums, Fixture Recap, and Vessel Specifications."""
            if not isinstance(items, list):
                items = [items]
            indent = '    ' * indent_level
            content = []
            in_clause = False
            current_clause_name = None
            for i, item in enumerate(items):
                if isinstance(item, dict):
                    if item.get('type') == 'subheader':
                        if in_clause:
                            content.append(f"{indent}[/CLAUSE]")
                            content.append("")
                        current_clause_name = item['content']
                        content.append("")
                        content.append(f"{indent}[CLAUSE]")
                        # content.append(f"{indent}[CLAUSE - {current_clause_name}]")
                        content.append(f"{indent}{current_clause_name}")
                        in_clause = True
                    elif item.get('type') == 'addendums':
                        content.append("")
                        content.append(f"{indent}[ADDENDUMS]")
                        content.append(f"{indent}{item['content']}")
                        content.append(f"{indent}[/ADDENDUMS]")
                        content.append("")
                    elif item.get('type') == 'fixture_recap':
                        content.append("")
                        content.append(f"{indent}[FIXTURE RECAP]")
                        content.append(f"{indent}{item['content']}")
                        content.append(f"{indent}[/FIXTURE RECAP]")
                        content.append("")
                    elif item.get('type') == 'vessel_specifications':
                        content.append("")
                        content.append(f"{indent}[VESSEL SPECIFICATIONS]")
                        content.append(f"{indent}{item['content']}")
                        content.append(f"{indent}[/VESSEL SPECIFICATIONS]")
                        content.append("")
                    elif item.get('type') == 'text':
                        content.append(f"{indent}{item['content']}")
                        is_last = i == len(items) - 1
                        next_is_subheader = not is_last and items[i + 1].get('type') == 'subheader'
                        if in_clause and (is_last or next_is_subheader):
                            content.append(f"{indent}[/CLAUSE]")
                            content.append("")
                            in_clause = False
                            current_clause_name = None
                elif isinstance(item, list):
                    formatted = format_item_content(item, indent_level)
                    content.extend(formatted)
                elif isinstance(item, str):
                    content.append(f"{indent}{item}")
            return content
        
        def insert_vessel_spec_tagged(vessel_details_text, vessel_spec_text):
            import re
            lines = vessel_details_text.split('\n')
            out_lines = []
            in_spec = False
            spec_lines = []
            for line in lines:
                if not in_spec and re.search(r'vessel\s*specifications?', line, re.IGNORECASE):
                    in_spec = True
                    spec_lines = [line]
                    continue
                if in_spec:
                    # End block if we see a section header (e.g., numbered section or contract details)
                    if re.search(r'^(\d+\.|contract\s*details)', line, re.IGNORECASE):
                        # Insert tagged block
                        out_lines.append('[VESSEL SPECIFICATIONS]')
                        out_lines.append(convert_strikethrough_to_md(vessel_spec_text))
                        out_lines.append('[/VESSEL SPECIFICATIONS]')
                        in_spec = False
                        out_lines.append(line)
                        continue
                    else:
                        spec_lines.append(line)
                        continue
                out_lines.append(line)
            # If the block was at the end
            if in_spec:
                out_lines.append('[VESSEL SPECIFICATIONS]')
                out_lines.append(convert_strikethrough_to_md(vessel_spec_text))
                out_lines.append('[/VESSEL SPECIFICATIONS]')
            return '\n'.join(out_lines)

        for vessel_name, vessel_data in vessels_data.items():
            vessel_output_dir = os.path.join(output_dir, vessel_name)
            os.makedirs(vessel_output_dir, exist_ok=True)
            
            # Save JSON file as intermediary format (will be deleted after markdown is generated)
            json_file = os.path.join(vessel_output_dir, f"{vessel_name}_data.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(vessel_data, f, indent=2)
            logging.info(f"Saved intermediary JSON file: {json_file}")
            
            # Generate markdown content
            md_content = []
            md_content.append(f"# {vessel_name} - Vessel Data\n")
            
            # Add vessel details section
            if vessel_data.get('vessel_details'):
                md_content.append("## Vessel Details\n")
                vessel_details_text = ''
                if isinstance(vessel_data['vessel_details'], dict) and vessel_data['vessel_details'].get('content'):
                    vessel_details_text = vessel_data['vessel_details']['content']
                else:
                    vessel_details_text = str(vessel_data['vessel_details'])
                # If vessel_specifications exists, replace the block in vessel_details with the tagged version
                if vessel_data.get('vessel_specifications') and vessel_data['vessel_specifications'].get('content'):
                    vessel_details_text = insert_vessel_spec_tagged(
                        vessel_details_text,
                        vessel_data['vessel_specifications']['content']
                    )
                md_content.append(convert_strikethrough_to_md(vessel_details_text))
                md_content.append("")
            
            # Add contract details section
            if vessel_data.get('contract_details'):
                md_content.append("## Contract Details\n")
                if isinstance(vessel_data['contract_details'], dict):
                    if vessel_data['contract_details'].get('structured'):
                        formatted_content = format_item_content(vessel_data['contract_details']['structured'])
                        md_content.extend(formatted_content)  # Extend with the list of formatted lines
                    else:
                        formatted_content = format_item_content(vessel_data['contract_details'])
                        md_content.extend(formatted_content)  # Extend with the list of formatted lines
                else:
                    formatted_content = format_item_content(vessel_data['contract_details'])
                    md_content.extend(formatted_content)  # Extend with the list of formatted lines
                md_content.append("")
            
            # Add delivery figures section
            if vessel_data.get('delivery_figures'):
                md_content.append("## Delivery Figures\n")
                if isinstance(vessel_data['delivery_figures'], dict) and vessel_data['delivery_figures'].get('content'):
                    md_content.append(convert_strikethrough_to_md(vessel_data['delivery_figures']['content']))
                else:
                    md_content.append(convert_strikethrough_to_md(str(vessel_data['delivery_figures'])))
                md_content.append("")
            
            # Add speed and consumption section
            if vessel_data.get('speed_and_consumption'):
                md_content.append("## Speed and Consumption\n")
                if isinstance(vessel_data['speed_and_consumption'], dict) and vessel_data['speed_and_consumption'].get('content'):
                    md_content.append(convert_strikethrough_to_md(vessel_data['speed_and_consumption']['content']))
                else:
                    md_content.append(convert_strikethrough_to_md(str(vessel_data['speed_and_consumption'])))
                md_content.append("")
            
            # Add lifting equipment section
            if vessel_data.get('lifting_equipment'):
                md_content.append("## Lifting Equipment\n")
                if isinstance(vessel_data['lifting_equipment'], dict) and vessel_data['lifting_equipment'].get('content'):
                    md_content.append(convert_strikethrough_to_md(vessel_data['lifting_equipment']['content']))
                else:
                    md_content.append(convert_strikethrough_to_md(str(vessel_data['lifting_equipment'])))
                md_content.append("")
            
            # Add HSEQ documentation section
            if vessel_data.get('hseq_documentation'):
                md_content.append("## HSEQ Documentation\n")
                if isinstance(vessel_data['hseq_documentation'], dict) and vessel_data['hseq_documentation'].get('content'):
                    md_content.append(convert_strikethrough_to_md(vessel_data['hseq_documentation']['content']))
                else:
                    md_content.append(convert_strikethrough_to_md(str(vessel_data['hseq_documentation'])))
                md_content.append("")

            # Remove the block that appends [VESSEL SPECIFICATIONS] at the end
            # (No further action needed here)
            
            # Save markdown file
            md_file = os.path.join(vessel_output_dir, f"{vessel_name}_data.md")
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join([line for line in md_content if isinstance(line, str)]))
            logging.info(f"Saved Markdown file: {md_file}")
            
            # Upload markdown file to Langdock
            if self.upload_enabled:
                api_key = os.getenv('LANGDOCK_API_KEY')
                folder_id = os.getenv('LANGDOCK_FOLDER_ID')
                if api_key and folder_id:
                    if self.upload_or_update_to_langdock(md_file, folder_id, api_key):
                        logging.info(f"Successfully updated {md_file} to Langdock Knowledge Folder")
                    else:
                        logging.error(f"Failed to update {md_file} to Langdock Knowledge Folder")
                else:
                    logging.warning("Langdock API key or folder ID not found in environment variables")
            
            # Do NOT delete the intermediary JSON file, so embedding_uploader.py can use it
            # try:
            #     os.remove(json_file)
            #     logging.info(f"Removed intermediary JSON file: {json_file}")
            # except Exception as e:
            #     logging.warning(f"Could not remove JSON file {json_file}: {e}")
        
        logging.info(f"Successfully saved outputs for {len(vessels_data)} vessels")

    def load_attachment_mappings(self) -> Dict[str, str]:
        """Load existing attachment ID mappings from JSON file."""
        mapping_file = 'attachment_mappings.json'
        if os.path.exists(mapping_file):
            try:
                with open(mapping_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logging.warning(f"Could not load attachment mappings: {e}")
        return {}

    def save_attachment_mappings(self):
        """Save current attachment ID mappings to JSON file."""
        mapping_file = 'attachment_mappings.json'
        try:
            with open(mapping_file, 'w') as f:
                json.dump(self.attachment_mappings, f, indent=2)
            logging.info(f"Saved attachment mappings to {os.path.abspath(mapping_file)}")
            logging.info(f"[DEBUG] Mapping file written: {os.path.abspath(mapping_file)}")
        except Exception as e:
            logging.error(f"Could not save attachment mappings: {e}")
            logging.error(f"[ERROR] Could not save attachment mappings: {e}")

    def _normalize_filename(self, filename: str) -> str:
        """Normalize filename for consistent mapping and upload (lowercase, underscores)."""
        return filename.strip().replace(' ', '_').lower()

    def upload_or_update_to_langdock(self, file_path: str, folder_id: str, api_key: str) -> bool:
        """Upload or update a file in Langdock Knowledge Folder using attachmentId if available."""
        filename = os.path.basename(file_path)
        norm_filename = self._normalize_filename(filename)
        attachment_id = self.attachment_mappings.get(norm_filename)
        
        try:
            headers = {'Authorization': f'Bearer {api_key}'}
            with open(file_path, 'rb') as f:
                files = {'file': (filename, f, 'application/octet-stream')}
                
                if attachment_id:
                    # Update existing file
                    url = f"https://api.langdock.com/knowledge/{folder_id}"
                    data = {'attachmentId': attachment_id}
                    response = requests.patch(url, headers=headers, files=files, data=data)
                    action = 'updated'
                else:
                    # Upload new file
                    url = f"https://api.langdock.com/knowledge/{folder_id}"
                    response = requests.post(url, headers=headers, files=files)
                    action = 'uploaded'
                
            if response.status_code == 200:
                logging.info(f"Successfully {action} {file_path} to Langdock Knowledge Folder")
                # If this was a new upload, try to extract attachmentId from response
                mapping_updated = False
                if not attachment_id and action == 'uploaded':
                    try:
                        response_data = response.json()
                        new_attachment_id = None
                        
                        # Check different possible response formats
                        if 'attachmentId' in response_data:
                            new_attachment_id = response_data['attachmentId']
                        elif 'result' in response_data and 'id' in response_data['result']:
                            new_attachment_id = response_data['result']['id']
                        elif 'id' in response_data:
                            new_attachment_id = response_data['id']
                        
                        if new_attachment_id:
                            self.attachment_mappings[norm_filename] = new_attachment_id
                            logging.info(f"Stored new attachmentId for {norm_filename}: {new_attachment_id}")
                        else:
                            logging.warning(f"No attachmentId/id found in response for {norm_filename}")
                            logging.info(f"[DEBUG] Full API response for {norm_filename}: {response_data}")
                    except Exception as e:
                        logging.warning(f"Could not extract attachmentId from response for {norm_filename}: {e}")
                        logging.error(f"[ERROR] Could not extract attachmentId from response for {norm_filename}: {e}")
                # Always save mapping after upload/update
                try:
                    self.save_attachment_mappings()
                except Exception as e:
                    logging.error(f"Failed to save attachment mappings after {action} for {norm_filename}: {e}")
                    logging.error(f"[ERROR] Failed to save attachment mappings after {action} for {norm_filename}: {e}")
                # Log mapping for debugging
                logging.info(f"[DEBUG] Current attachment_mappings: {json.dumps(self.attachment_mappings, indent=2)}")
                if not mapping_updated and action == 'uploaded':
                    logging.warning(f"[WARNING] Mapping was not updated for {norm_filename} after upload!")
                return True
            else:
                logging.error(f"Failed to {action} {file_path}. Status code: {response.status_code}, Response: {response.text}")
                logging.error(f"[ERROR] Failed to {action} {file_path}. Status code: {response.status_code}, Response: {response.text}")
                return False
                
        except Exception as e:
            logging.error(f"Error uploading/updating {file_path} to Langdock: {str(e)}")
            logging.error(f"[ERROR] Exception during upload/update: {e}")
            return False

    def main(self):
        # Enable debug logging
        logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
        vessels_data = self.process_files()
        self.save_outputs(vessels_data)
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels.")
        for vessel_name, data in vessels_data.items():
            print(f"\nVessel: {vessel_name}")
            for section, content in data.items():
                if section == 'contract_details' and 'structured' in content:
                    content_length = len(content.get('structured', []))
                    print(f"  {section}: {content_length} structured items")
                else:
                    content_length = len(content.get('content', ''))
                    print(f"  {section}: {content_length} characters")
        # Upload files to Langdock Knowledge Folder (optional)
        try:
            folder_id = "69c0327b-24fe-4cd8-97d8-92688b526bac"
            api_key = os.getenv('LANGDOCK_API_KEY')
            if api_key:
                files_to_upload = []
                for vessel_name in vessels_data.keys():
                    vessel_dir = os.path.join('output/vessels', vessel_name)
                    # Only upload .md files (per user request)
                    files_to_upload.append(os.path.join(vessel_dir, f'{vessel_name}_data.md'))
                for file_path in files_to_upload:
                    if os.path.exists(file_path):
                        self.upload_or_update_to_langdock(file_path, folder_id, api_key)
                    else:
                        logging.warning(f"File {file_path} not found, skipping upload")
                # Save updated mappings after all uploads
                self.save_attachment_mappings()
                logging.info(f"[DEBUG] Final attachment_mappings: {json.dumps(self.attachment_mappings, indent=2)}")
            else:
                logging.warning("LANGDOCK_API_KEY environment variable not set. Skipping upload to Langdock.")
        except Exception as e:
            logging.error(f"Error uploading files to Langdock: {str(e)}")
            logging.error(f"[ERROR] Error uploading files to Langdock: {e}")

if __name__ == "__main__":
    extractor = VesselDataExtractor()
    extractor.main() 