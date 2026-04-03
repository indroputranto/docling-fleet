#!/usr/bin/env python3
"""
Optimized Vessel Data Extractor - Enhanced for AI Search
Extracts vessel information with semantic enhancement for better AI chatbot performance
"""

import os
import json
import logging
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class SemanticClauseEnhancer:
    """Semantic enhancement engine for maritime contract clauses and vessel specifications"""
    
    def __init__(self):
        # Vessel specification patterns for semantic chunking
        self.vessel_spec_patterns = {
            'basic_info': {
                'patterns': ['vessel name', 'imo number', 'flag', 'call sign', 'built', 'yard', 'class'],
                'keywords': ['vessel name', 'IMO', 'flag state', 'call sign', 'built year', 'shipyard', 'classification'],
                'description': 'basic vessel identification information including name, IMO, flag, and build details',
                'priority': 'high'
            },
            'dimensions': {
                'patterns': ['loa', 'length overall', 'beam', 'draft', 'depth', 'height', 'dimensions'],
                'keywords': ['LOA', 'length overall', 'beam', 'draft', 'depth', 'air draft', 'dimensions'],
                'description': 'vessel dimensions including length, beam, draft, and height specifications',
                'priority': 'high'
            },
            'capacities': {
                'patterns': ['dwt', 'deadweight', 'cargo', 'holds', 'tanks', 'capacity', 'volume'],
                'keywords': ['DWT', 'deadweight', 'cargo capacity', 'holds', 'fuel tanks', 'ballast tanks', 'volume'],
                'description': 'vessel cargo and tank capacities including DWT and hold specifications',
                'priority': 'high'
            },
            'equipment': {
                'patterns': ['crane', 'winch', 'gear', 'equipment', 'handling', 'machinery'],
                'keywords': ['cranes', 'cargo gear', 'winches', 'handling equipment', 'deck machinery'],
                'description': 'vessel cargo handling equipment and deck machinery specifications',
                'priority': 'medium'
            },
            'propulsion': {
                'patterns': ['engine', 'main engine', 'auxiliary', 'propeller', 'power', 'speed'],
                'keywords': ['main engine', 'auxiliary engine', 'propeller', 'power', 'service speed', 'propulsion'],
                'description': 'vessel propulsion system including engines and performance specifications',
                'priority': 'medium'
            },
            'navigation': {
                'patterns': ['navigation', 'gps', 'radar', 'communication', 'bridge', 'instruments'],
                'keywords': ['navigation equipment', 'GPS', 'radar', 'communication', 'bridge equipment'],
                'description': 'vessel navigation and communication equipment specifications',
                'priority': 'medium'
            },
            'accommodation': {
                'patterns': ['accommodation', 'crew', 'cabins', 'berths', 'mess', 'officers'],
                'keywords': ['crew accommodation', 'cabins', 'berths', 'mess room', 'officers quarters'],
                'description': 'vessel accommodation facilities and crew living spaces',
                'priority': 'low'
            },
            'certificates': {
                'patterns': ['certificate', 'survey', 'inspection', 'valid', 'expires', 'classification'],
                'keywords': ['certificates', 'class surveys', 'inspections', 'validity', 'expiry dates'],
                'description': 'vessel certificates, surveys, and inspection records',
                'priority': 'medium'
            }
        }
        
        self.clause_patterns = {
            'bunkers': {
                'patterns': ['bunkers', 'bunker', 'fuel', 'delivery bunkers', 'redelivery bunkers'],
                'keywords': ['bunkers', 'fuel', 'VLSFO', 'MGO', 'fuel oil', 'marine gas oil', 'bunker quantities', 'bunker prices', 'fuel specifications'],
                'description': 'fuel delivery, redelivery, specifications, and bunker management requirements',
                'priority': 'high'
            },
            'redelivery': {
                'patterns': ['redelivery', 'redeliver', 'redelivery port', 'redelivery notice'],
                'keywords': ['redelivery', 'vessel redelivery', 'redelivery port', 'redelivery notice', 'redelivery condition'],
                'description': 'vessel redelivery procedures, location, timing, and conditions',
                'priority': 'high'
            },
            'delivery': {
                'patterns': ['delivery', 'deliver', 'on-hire', 'delivery port', 'delivery notice'],
                'keywords': ['delivery', 'vessel delivery', 'delivery port', 'delivery notice', 'delivery condition', 'on-hire'],
                'description': 'vessel delivery procedures, location, timing, and conditions',
                'priority': 'high'
            },
            'hire': {
                'patterns': ['hire', 'rate of hire', 'hire payment', 'daily hire'],
                'keywords': ['hire', 'rate of hire', 'hire payment', 'daily hire', 'hire calculation', 'payment terms'],
                'description': 'daily hire rates, payment terms, and hire calculation methods',
                'priority': 'high'
            },
            'performance': {
                'patterns': ['performance', 'voyages', 'navigation', 'speed'],
                'keywords': ['performance', 'voyages', 'vessel performance', 'navigation', 'speed', 'master duties'],
                'description': 'vessel performance requirements, voyage execution, and master responsibilities',
                'priority': 'high'
            },
            'off_hire': {
                'patterns': ['off-hire', 'off hire', 'oﬀ-hire', 'off hire', 'breakdown', 'deficiency'],
                'keywords': ['off-hire', 'off hire', 'hire suspension', 'breakdown', 'vessel deficiency', 'time lost'],
                'description': 'circumstances for hire suspension and off-hire procedures',
                'priority': 'high'
            },
            'drydocking': {
                'patterns': ['drydocking', 'drydock', 'dry dock', 'docking', 'dry-docking'],
                'keywords': ['drydocking', 'drydock', 'dry dock', 'vessel maintenance', 'docking period', 'drydock survey'],
                'description': 'drydocking requirements, procedures, and vessel maintenance scheduling',
                'priority': 'high'
            },
            'ice': {
                'patterns': ['ice', 'ice clause', 'icebound', 'ice conditions', 'force ice'],
                'keywords': ['ice', 'ice conditions', 'icebound port', 'force ice', 'ice navigation', 'winter conditions'],
                'description': 'ice navigation restrictions and winter weather procedures',
                'priority': 'high'
            },
            'laydays': {
                'patterns': ['laydays', 'laytime', 'cancelling', 'laydays/cancelling', 'notice time'],
                'keywords': ['laydays', 'laytime', 'cancelling date', 'notice time', 'charter commencement'],
                'description': 'laydays, cancelling dates, and charter commencement timing',
                'priority': 'high'
            },
            'cargo_claims': {
                'patterns': ['cargo claims', 'cargo damage', 'cargo liability', 'claims'],
                'keywords': ['cargo claims', 'cargo damage', 'cargo liability', 'cargo insurance', 'damage claims'],
                'description': 'cargo claims procedures and liability allocation',
                'priority': 'medium'
            },
            'liens': {
                'patterns': ['liens', 'lien', 'maritime lien', 'security'],
                'keywords': ['liens', 'maritime lien', 'security interest', 'lien rights', 'cargo lien'],
                'description': 'maritime liens and security interests in cargo and freight',
                'priority': 'medium'
            },
            'requisition': {
                'patterns': ['requisition', 'government requisition', 'vessel requisition'],
                'keywords': ['requisition', 'government requisition', 'vessel seizure', 'requisition compensation'],
                'description': 'government requisition procedures and compensation',
                'priority': 'medium'
            },
            'total_loss': {
                'patterns': ['total loss', 'vessel loss', 'constructive total loss'],
                'keywords': ['total loss', 'vessel loss', 'constructive total loss', 'loss compensation'],
                'description': 'total loss procedures and compensation arrangements',
                'priority': 'medium'
            },
            'salvage': {
                'patterns': ['salvage', 'salvage services', 'general average'],
                'keywords': ['salvage', 'salvage services', 'general average', 'salvage compensation'],
                'description': 'salvage operations and general average procedures',
                'priority': 'medium'
            },
            'pollution': {
                'patterns': ['pollution', 'oil pollution', 'environmental', 'marpol'],
                'keywords': ['pollution', 'oil pollution', 'environmental compliance', 'MARPOL', 'pollution prevention'],
                'description': 'pollution prevention and environmental compliance requirements',
                'priority': 'medium'
            },
            'taxes': {
                'patterns': ['taxes', 'duties', 'customs', 'port taxes'],
                'keywords': ['taxes', 'port taxes', 'customs duties', 'local taxes', 'government charges'],
                'description': 'taxes, duties, and government charges allocation',
                'priority': 'medium'
            },
            'liberties': {
                'patterns': ['liberties', 'vessel liberties', 'navigation liberties'],
                'keywords': ['liberties', 'vessel liberties', 'navigation rights', 'routing freedom'],
                'description': 'vessel navigation liberties and routing flexibility',
                'priority': 'medium'
            },
            'notices': {
                'patterns': ['notices', 'notification', 'notice requirements'],
                'keywords': ['notices', 'notification', 'notice requirements', 'communication procedures'],
                'description': 'notice and communication requirements between parties',
                'priority': 'medium'
            },
            'commissions': {
                'patterns': ['commissions', 'brokerage', 'commission payment'],
                'keywords': ['commissions', 'brokerage', 'commission payment', 'broker fees'],
                'description': 'commission and brokerage fee arrangements',
                'priority': 'medium'
            },
            'exceptions': {
                'patterns': ['exceptions', 'excepted perils', 'force majeure'],
                'keywords': ['exceptions', 'excepted perils', 'force majeure', 'excluded liabilities'],
                'description': 'exceptions from liability and excepted perils',
                'priority': 'medium'
            },
            'cargo_handling': {
                'patterns': ['cargo handling', 'cargo gear', 'loading', 'discharging'],
                'keywords': ['cargo handling', 'cargo gear', 'loading equipment', 'discharging', 'stevedoring'],
                'description': 'cargo handling equipment and loading/discharging procedures',
                'priority': 'medium'
            },
            'supercargo': {
                'patterns': ['supercargo', 'cargo superintendent', 'cargo supervision'],
                'keywords': ['supercargo', 'cargo superintendent', 'cargo supervision', 'cargo inspection'],
                'description': 'supercargo appointment and cargo supervision duties',
                'priority': 'medium'
            },
            'stowaways': {
                'patterns': ['stowaways', 'stowaway', 'illegal passengers'],
                'keywords': ['stowaways', 'stowaway', 'illegal passengers', 'unauthorized persons'],
                'description': 'stowaway procedures and cost allocation',
                'priority': 'medium'
            },
            'smuggling': {
                'patterns': ['smuggling', 'contraband', 'illegal cargo'],
                'keywords': ['smuggling', 'contraband', 'illegal cargo', 'customs violations'],
                'description': 'smuggling prevention and liability procedures',
                'priority': 'medium'
            },
            'insurance': {
                'patterns': ['insurance', 'P&I', 'hull insurance', 'war risks'],
                'keywords': ['insurance', 'P&I insurance', 'vessel insurance', 'hull insurance', 'war risks'],
                'description': 'insurance requirements, coverage, and risk management',
                'priority': 'medium'
            },
            'stevedore_damage': {
                'patterns': ['stevedore', 'damage', 'loading damage'],
                'keywords': ['stevedore damage', 'cargo damage', 'loading damage', 'stevedore liability'],
                'description': 'stevedore damage liability and compensation procedures',
                'priority': 'medium'
            },
            'bills_of_lading': {
                'patterns': ['bills of lading', 'bill of lading', 'B/L', 'cargo documents'],
                'keywords': ['bills of lading', 'bill of lading', 'B/L', 'cargo documents', 'shipping documents'],
                'description': 'bill of lading procedures, responsibilities, and requirements',
                'priority': 'medium'
            },
            'law_arbitration': {
                'patterns': ['law', 'arbitration', 'disputes', 'governing law'],
                'keywords': ['law', 'arbitration', 'disputes', 'governing law', 'legal jurisdiction'],
                'description': 'governing law, dispute resolution, and arbitration procedures',
                'priority': 'medium'
            },
            'war_risks': {
                'patterns': ['war', 'piracy', 'conflict', 'war zones'],
                'keywords': ['war risks', 'war', 'piracy', 'conflict', 'war zones', 'security'],
                'description': 'war risks, security concerns, and conflict zone procedures',
                'priority': 'medium'
            },
            'protective': {
                'patterns': ['protective', 'protection', 'liability', 'indemnity'],
                'keywords': ['protective clauses', 'protection', 'liability protection', 'indemnity'],
                'description': 'protective clauses and liability limitations',
                'priority': 'medium'
            }
        }
    
    def identify_clause_type(self, heading: str, content: str) -> str:
        """Identify the clause type based on heading and content"""
        heading_lower = heading.lower()
        content_lower = content.lower()[:500]  # Check first 500 chars
        
        # Check for specific patterns first (most specific to least specific)
        clause_order = [
            'redelivery', 'off_hire', 'delivery', 'bunkers', 'hire', 'performance', 
            'drydocking', 'laydays', 'cargo_claims', 'liens', 'requisition',
            'total_loss', 'salvage', 'pollution', 'taxes', 'liberties', 'notices',
            'commissions', 'exceptions', 'cargo_handling', 'supercargo', 'stowaways',
            'smuggling', 'ice', 'insurance', 'stevedore_damage', 'bills_of_lading', 
            'law_arbitration', 'war_risks', 'protective'
        ]
        
        for clause_type in clause_order:
            config = self.clause_patterns[clause_type]
            patterns = config['patterns']
            # Check heading first (more reliable)
            if any(pattern in heading_lower for pattern in patterns):
                return clause_type
        
        # Then check content if no heading match
        for clause_type in clause_order:
            config = self.clause_patterns[clause_type]
            patterns = config['patterns']
            if any(pattern in content_lower for pattern in patterns):
                return clause_type
        
        return 'general'
    
    def identify_vessel_spec_type(self, content: str) -> str:
        """Identify the vessel specification type based on content"""
        content_lower = content.lower()
        
        # Check for specific patterns (most specific to least specific)
        spec_order = ['basic_info', 'dimensions', 'capacities', 'propulsion', 'equipment', 'navigation', 'certificates', 'accommodation']
        
        for spec_type in spec_order:
            config = self.vessel_spec_patterns[spec_type]
            patterns = config['patterns']
            if any(pattern in content_lower for pattern in patterns):
                return spec_type
        
        return 'general'
    
    def enhance_vessel_spec_data(self, vessel_name: str, content: str, spec_type: str) -> Dict[str, Any]:
        """Enhance vessel specification data with semantic information"""
        enhanced_data = {
            "type": "vessel_spec",
            "content": content,
            "spec_type": spec_type,
            "vessel_name": vessel_name
        }
        
        if spec_type != 'general' and spec_type in self.vessel_spec_patterns:
            config = self.vessel_spec_patterns[spec_type]
            
            # Add semantic enhancement
            enhanced_data.update({
                "semantic_keywords": config['keywords'],
                "description": f"{spec_type.replace('_', ' ').title()} specifications for {vessel_name} containing {config['description']}",
                "priority": config['priority'],
                "is_primary_spec": True,
                "enhanced_for_search": True
            })
            
            # Add specific flags for important spec types
            if spec_type == 'basic_info':
                enhanced_data['is_vessel_basic_info'] = True
            elif spec_type == 'dimensions':
                enhanced_data['is_vessel_dimensions'] = True
            elif spec_type == 'capacities':
                enhanced_data['is_vessel_capacities'] = True
        else:
            enhanced_data.update({
                "priority": "low",
                "is_primary_spec": False,
                "enhanced_for_search": False
            })
        
        return enhanced_data
    
    def enhance_clause_data(self, vessel_name: str, heading: str, content: str, tag: str) -> Dict[str, Any]:
        """Enhance clause data with semantic information"""
        clause_type = self.identify_clause_type(heading, content)
        
        enhanced_data = {
            "type": "text",
            "content": content,
            "tag": tag,
            "clause_type": clause_type,
            "vessel_name": vessel_name
        }
        
        if clause_type != 'general':
            config = self.clause_patterns[clause_type]
            
            # Add semantic enhancement
            enhanced_data.update({
                "semantic_keywords": config['keywords'],
                "description": f"{clause_type.replace('_', ' ').title()} clause for {vessel_name} containing {config['description']}",
                "priority": config['priority'],
                "is_primary_clause": True,
                "enhanced_for_search": True
            })
            
            # Add specific flags for important clause types
            if clause_type == 'bunkers':
                enhanced_data['is_primary_bunkers_clause'] = True
            elif clause_type == 'delivery':
                enhanced_data['is_primary_delivery_clause'] = True
            elif clause_type == 'redelivery':
                enhanced_data['is_primary_redelivery_clause'] = True
        else:
            enhanced_data.update({
                "priority": "low",
                "is_primary_clause": False,
                "enhanced_for_search": False
            })
        
        return enhanced_data

class OptimizedVesselDataExtractor:
    def __init__(self, source_dir: str = 'source/vessels'):
        """Initialize the optimized vessel data extractor"""
        self.source_dir = source_dir
        self.semantic_enhancer = SemanticClauseEnhancer()
        self.extracted_data = {}
        self._strikethrough_cache = {}
        self._text_cache = {}
    
    def _remove_strikethrough_text(self, text: str) -> str:
        """Process strikethrough text patterns from the given text by wrapping them in [STRIKETHROUGH] tags."""
        import re
        
        # Pattern 1: [[STRIKE ... STRIKE]] - common DOCX strike-through format
        text = re.sub(r'\[\[STRIKE\s+(.*?)\s+STRIKE\]\]', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 2: QUOTE STRIKE-THROUGH ... UNQUOTE STRIKE-THROUGH
        text = re.sub(r'QUOTE STRIKE-THROUGH\s+(.*?)\s+UNQUOTE STRIKE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 3: Simple strike-through markers
        text = re.sub(r'STRIKE\s+(.*?)\s+STRIKE', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 4: Standalone STRIKE markers
        text = re.sub(r'\bSTRIKE\b', '', text, flags=re.IGNORECASE)
        
        # Pattern 5: QUOTE-THROUGH ... UNQUOTE-THROUGH
        text = re.sub(r'QUOTE-THROUGH\s+(.*?)\s+UNQUOTE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        return text.strip()

    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to group consecutive strikethrough runs and wrap them in [STRIKETHROUGH]...[/STRIKETHROUGH] tags."""
        if not hasattr(paragraph, 'runs'):
            return paragraph.text if hasattr(paragraph, 'text') else str(paragraph)
        
        result_parts = []
        current_strikethrough_parts = []
        in_strikethrough = False
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Check if this run has strikethrough formatting
            is_strikethrough = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Check for strike elements
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if strike_elem is not None or dstrike_elem is not None:
                        is_strikethrough = True
            
            if is_strikethrough:
                if not in_strikethrough:
                    # Starting a new strikethrough section
                    in_strikethrough = True
                current_strikethrough_parts.append(text)
            else:
                if in_strikethrough:
                    # Ending strikethrough section
                    if current_strikethrough_parts:
                        result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
                        current_strikethrough_parts = []
                    in_strikethrough = False
                result_parts.append(text)
        
        # Handle any remaining strikethrough text
        if current_strikethrough_parts:
            result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
        
        return ''.join(result_parts)

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file, preserving strikethrough formatting."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs, using the strikethrough-aware function
            for paragraph in doc.paragraphs:
                text = self._process_paragraph_for_strikethrough(paragraph)
                if text:
                    text_parts.append(text)
            
            # Extract tables and add them to the text
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- TABLE {i+1} ---")
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # For each cell, process all paragraphs for strikethrough
                        cell_text = []
                        for p in cell.paragraphs:
                            processed_text = self._process_paragraph_for_strikethrough(p)
                            if processed_text:
                                cell_text.append(processed_text)
                        row_cells.append(' '.join(cell_text).strip())
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("--- END TABLE ---\n")
            
            # Join all parts with double newlines for better section separation
            return '\n\n'.join(text_parts)
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_vessel_name(self, text: str, filename: str) -> str:
        """Extract vessel name from text or filename."""
        # Try to find vessel name in text first
        lines = text.split('\n')[:20]  # Check first 20 lines
        for line in lines:
            # Look for patterns like "Vessel Name: MV SOMETHING"
            match = re.search(r'vessel\s+name:?\s*(.+)', line, re.IGNORECASE)
            if match:
                vessel_name = match.group(1).strip()
                if vessel_name and not vessel_name.lower().startswith('insert'):
                    return vessel_name.replace('"', '').replace("'", "")
        
        # Fallback: extract from filename
        base_name = os.path.splitext(filename)[0]
        # Remove common prefixes/suffixes
        vessel_name = re.sub(r'^(mv|m/v|vessel)[\s_-]*', '', base_name, flags=re.IGNORECASE)
        vessel_name = re.sub(r'[\s_-]*(final|ns|done).*$', '', vessel_name, flags=re.IGNORECASE)
        vessel_name = vessel_name.replace('_', ' ').replace('-', ' ').strip()
        
        # Add MV prefix if not present
        if vessel_name and not vessel_name.upper().startswith('MV'):
            vessel_name = f"MV {vessel_name.upper()}"
        
        return vessel_name

    def _extract_vessel_details_structured(self, text: str, vessel_name: str) -> List[Dict[str, Any]]:
        """Extract vessel details as structured, semantic chunks"""
        items = []
        
        # Split text into logical chunks by line breaks and common separators
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        current_chunk = []
        current_chunk_type = None
        
        for line in lines:
            # Skip headers and section titles
            if re.match(r'^\d+\.?\s*vessel\s*details', line, re.IGNORECASE):
                continue
            
            # Check if this line starts a new logical group
            line_lower = line.lower()
            
            # Group related information together
            if any(keyword in line_lower for keyword in ['vessel name', 'imo', 'flag', 'call sign', 'built', 'yard']):
                # Flush previous chunk
                if current_chunk:
                    chunk_content = '\n'.join(current_chunk)
                    spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                    enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                    items.append(enhanced_data)
                
                current_chunk = [line]
                current_chunk_type = 'basic_info'
                
            elif any(keyword in line_lower for keyword in ['loa', 'beam', 'draft', 'depth', 'length', 'height', 'dimensions']):
                if current_chunk_type != 'dimensions':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'dimensions'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['dwt', 'deadweight', 'cargo', 'holds', 'tanks', 'capacity']):
                if current_chunk_type != 'capacities':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'capacities'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['engine', 'propeller', 'power', 'speed', 'auxiliary']):
                if current_chunk_type != 'propulsion':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'propulsion'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['crane', 'gear', 'equipment', 'winch', 'handling']):
                if current_chunk_type != 'equipment':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'equipment'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['navigation', 'gps', 'radar', 'communication', 'bridge']):
                if current_chunk_type != 'navigation':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'navigation'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['certificate', 'survey', 'class', 'inspection', 'valid']):
                if current_chunk_type != 'certificates':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'certificates'
                else:
                    current_chunk.append(line)
                    
            elif any(keyword in line_lower for keyword in ['accommodation', 'crew', 'cabins', 'berths']):
                if current_chunk_type != 'accommodation':
                    # Flush previous chunk
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk)
                        spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
                        enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
                        items.append(enhanced_data)
                    
                    current_chunk = [line]
                    current_chunk_type = 'accommodation'
                else:
                    current_chunk.append(line)
            else:
                # Add to current chunk or start a general chunk
                if current_chunk:
                    current_chunk.append(line)
                else:
                    current_chunk = [line]
                    current_chunk_type = 'general'
        
        # Flush any remaining chunk
        if current_chunk:
            chunk_content = '\n'.join(current_chunk)
            spec_type = self.semantic_enhancer.identify_vessel_spec_type(chunk_content)
            enhanced_data = self.semantic_enhancer.enhance_vessel_spec_data(vessel_name, chunk_content, spec_type)
            items.append(enhanced_data)
        
        return items

    def _extract_contract_details_enhanced(self, doc: Document, vessel_name: str) -> List[Dict[str, Any]]:
        """Extract contract details with semantic enhancement"""
        items = []
        in_contract_section = False
        current_content = []
        current_subheader = None
        special_section = None
        special_content = []
        special_type = None
        in_charter_party = False

        def is_heading(para, levels=(1,2,3,4)):
            if para.style and para.style.name:
                for lvl in levels:
                    if para.style.name.lower().startswith(f'heading {lvl}'):
                        return True
            return False

        def is_addendum_section(text):
            return re.search(r'2\.?\s*1\.?\s*addendum', text, re.IGNORECASE) or (re.search(r'addendum', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_fixture_recap_section(text):
            return re.search(r'2\.?\s*2\.?\s*fixture\s*recap', text, re.IGNORECASE) or (re.search(r'fixture\s*recap', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_charter_party_section(text):
            return re.search(r'2\.?\s*3\.?\s*charter\s*party', text, re.IGNORECASE) or (re.search(r'charter\s*party', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            text = self._process_paragraph_for_strikethrough(paragraph)
            
            if re.search(r'contract\s*details', text, re.IGNORECASE):
                in_contract_section = True
                continue
            
            if not in_contract_section:
                continue

            # Check for special sections
            if is_addendum_section(text):
                if current_content:
                    items.extend(current_content)
                    current_content = []
                if special_section:
                    items.append({
                        "type": special_type,
                        "content": '\n'.join(special_content),
                        "tag": special_type.upper(),
                        "vessel_name": vessel_name,
                        "priority": "medium",  # Changed from "low" to "medium"
                        "enhanced_for_search": True  # Changed from False to True
                    })
                special_section = "addendums"
                special_type = "addendums"
                special_content = [text]
                continue

            if is_fixture_recap_section(text):
                if current_content:
                    items.extend(current_content)
                    current_content = []
                if special_section:
                    items.append({
                        "type": special_type,
                        "content": '\n'.join(special_content),
                        "tag": special_type.upper(),
                        "vessel_name": vessel_name,
                        "priority": "medium",  # Changed from "low" to "medium"
                        "enhanced_for_search": True  # Changed from False to True
                    })
                special_section = "fixture_recap"
                special_type = "fixture_recap"
                special_content = [text]
                continue

            if is_charter_party_section(text):
                if special_section:
                    items.append({
                        "type": special_type,
                        "content": '\n'.join(special_content),
                        "tag": special_type.upper(),
                        "vessel_name": vessel_name,
                        "priority": "medium",  # Changed from "low" to "medium"
                        "enhanced_for_search": True  # Changed from False to True
                    })
                    special_section = None
                in_charter_party = True
                continue

            if special_section:
                special_content.append(text)
                continue

            # Process charter party content with semantic enhancement
            if in_charter_party:
                is_heading1 = is_heading(paragraph, levels=(1,))
                
                if is_heading1:
                    if current_content:
                        items.extend(current_content)
                        current_content = []
                    
                    current_subheader = text
                    
                    # Create enhanced subheader with semantic information
                    clean_heading = re.sub(r'^\d+\.\s*', '', current_subheader)
                    clean_heading = re.sub(r'[^\w\s-]', '', clean_heading).strip()
                    contextual_tag = f"CLAUSE - {clean_heading}" if clean_heading else "CLAUSE"
                    
                    # Add semantic enhancement to subheader
                    clause_type = self.semantic_enhancer.identify_clause_type(current_subheader, "")
                    enhanced_subheader = {
                        "type": "subheader",
                        "content": text,
                        "tag": contextual_tag,
                        "clause_type": clause_type,
                        "vessel_name": vessel_name,
                        "priority": "high",
                        "enhanced_for_search": True
                    }
                    
                    items.append(enhanced_subheader)
                else:
                    # Regular text with semantic enhancement
                    text = re.sub(r'\s+', ' ', text)
                    text = re.sub(r'\s+([.,;:])', r'\1', text)
                    
                    # Create contextual tag with subheader
                    tag = "CLAUSE"
                    if current_subheader:
                        clean_subheader = re.sub(r'^\d+\.\s*', '', current_subheader)
                        clean_subheader = re.sub(r'[^\w\s-]', '', clean_subheader).strip()
                        if clean_subheader:
                            tag = f"CLAUSE - {clean_subheader}"
                    
                    # Apply semantic enhancement
                    enhanced_item = self.semantic_enhancer.enhance_clause_data(
                        vessel_name, current_subheader or "", text, tag
                    )
                    current_content.append(enhanced_item)

        # Flush any remaining content
        if current_content:
            items.extend(current_content)
        if special_section:
            items.append({
                "type": special_type,
                "content": '\n'.join(special_content),
                "tag": special_type.upper(),
                "vessel_name": vessel_name,
                "priority": "medium",  # Changed from "low" to "medium"
                "enhanced_for_search": True  # Changed from False to True
            })

        return items

    def _extract_vessel_details(self, text: str, vessel_name: str, doc: Document = None) -> Dict[str, Any]:
        """Extract vessel details with semantic enhancement"""
        import re
        details = {
            'vessel_details': {'structured': []},  # Changed to structured like contract_details
            'contract_details': {'structured': []},
            'delivery_figures': {'content': ''},
            'speed_and_consumption': {'content': ''},
            'lifting_equipment': {'content': ''},
            'hseq_documentation': {'content': ''}
        }
        
        # Section patterns
        section_patterns = [
            ('vessel_details', re.compile(r'^1\.?\s*vessel details', re.IGNORECASE)),
            ('contract_details', re.compile(r'^2\.?\s*contract details', re.IGNORECASE)),
            ('delivery_figures', re.compile(r'^3\.?\s*delivery figures', re.IGNORECASE)),
            ('speed_and_consumption', re.compile(r'^4\.?\s*speed and consumption', re.IGNORECASE)),
            ('lifting_equipment', re.compile(r'^5\.?\s*lifting equipment', re.IGNORECASE)),
            ('hseq_documentation', re.compile(r'^6\.?\s*hseq documentation', re.IGNORECASE)),
        ]
        
        # Process strikethrough markers
        strikethrough_blocks = []
        def save_strikethrough(match):
            block_id = f"__STRIKETHROUGH_{len(strikethrough_blocks)}__"
            strikethrough_blocks.append(match.group(1))
            return block_id
        
        processed_text = re.sub(r'\[STRIKETHROUGH\](.*?)\[/STRIKETHROUGH\]', save_strikethrough, text, flags=re.DOTALL)
        
        # Find section boundaries
        lines = processed_text.split('\n')
        section_indices = {}
        for i, line in enumerate(lines):
            for key, pattern in section_patterns:
                if key not in section_indices and pattern.match(line.strip()):
                    section_indices[key] = i
        
        # Extract content for each section
        sorted_sections = sorted(section_indices.items(), key=lambda x: x[1])
        for idx, (key, start) in enumerate(sorted_sections):
            end = sorted_sections[idx + 1][1] if idx + 1 < len(sorted_sections) else len(lines)
            section_lines = lines[start + 1:end]
            section_text = '\n'.join(section_lines).strip()
            
            # Restore strikethrough blocks
            for i, block in enumerate(strikethrough_blocks):
                section_text = section_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            
            if key == 'contract_details' and doc is not None:
                details[key]['structured'] = self._extract_contract_details_enhanced(doc, vessel_name)
            elif key == 'vessel_details':
                details[key]['structured'] = self._extract_vessel_details_structured(section_text, vessel_name)
            else:
                details[key]['content'] = section_text
        
        return details

    def process_files(self) -> Dict[str, Dict[str, Any]]:
        """Process all vessel files with semantic enhancement"""
        vessels_data = {}
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return vessels_data
        
        # Filter for DOCX files
        docx_files = [f for f in os.listdir(self.source_dir) 
                     if f.lower().endswith('.docx') and not f.startswith('~$')]
        logging.info(f"Found {len(docx_files)} DOCX files to process: {docx_files}")
        
        for filename in docx_files:
            file_path = os.path.join(self.source_dir, filename)
            logging.info(f"Processing {filename}")
            
            try:
                # Extract text and document
                full_text = self._extract_text_from_docx(file_path)
                doc = Document(file_path)
                
                vessel_name = self._extract_vessel_name(full_text, filename)
                
                if vessel_name:
                    vessel_details = self._extract_vessel_details(full_text, vessel_name, doc)
                    vessels_data[vessel_name] = vessel_details
                    logging.info(f"Successfully extracted data for {vessel_name} with semantic enhancement")
                else:
                    logging.warning(f"Could not extract vessel name from {filename}")
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
        
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels: {list(vessels_data.keys())}")
        return vessels_data

    def save_outputs(self, vessels_data: Dict[str, Dict[str, Any]], output_dir: str = 'output/vessels'):
        """Save processed vessel data to JSON and markdown files"""
        os.makedirs(output_dir, exist_ok=True)
        
        for vessel_name, vessel_data in vessels_data.items():
            vessel_dir = os.path.join(output_dir, vessel_name)
            os.makedirs(vessel_dir, exist_ok=True)
            
            # Save JSON with semantic enhancements
            json_file = os.path.join(vessel_dir, f"{vessel_name}_data.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(vessel_data, f, indent=2, ensure_ascii=False)
            
            # Save markdown
            md_file = os.path.join(vessel_dir, f"{vessel_name}_data.md")
            self._save_markdown(vessel_data, md_file)
            
            logging.info(f"Saved enhanced data for {vessel_name}")

    def _save_markdown(self, vessel_data: Dict[str, Any], md_file: str):
        """Save vessel data as markdown with semantic structure"""
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(f"# {vessel_data.get('vessel_details', {}).get('content', '').split('Vessel Name:')[-1].split()[0] if 'Vessel Name:' in vessel_data.get('vessel_details', {}).get('content', '') else 'Vessel'} Data\n\n")
            
            # Add structured vessel details
            if vessel_data.get('vessel_details', {}).get('structured'):
                f.write("## Vessel Details\n\n")
                for item in vessel_data['vessel_details']['structured']:
                    if isinstance(item, dict):
                        spec_type = item.get('spec_type', 'general')
                        content = item.get('content', '')
                        if content:
                            f.write(f"### {spec_type.replace('_', ' ').title()}\n\n")
                            if item.get('enhanced_for_search'):
                                f.write(f"*Specification Type: {spec_type}*\n\n")
                            f.write(f"{content}\n\n")
            
            # Add enhanced contract details
            if vessel_data.get('contract_details', {}).get('structured'):
                f.write("## Contract Details\n\n")
                for item in vessel_data['contract_details']['structured']:
                    if isinstance(item, dict):
                        if item.get('type') == 'subheader':
                            f.write(f"### {item['content']}\n\n")
                            if item.get('enhanced_for_search'):
                                f.write(f"*Clause Type: {item.get('clause_type', 'general')}*\n\n")
                        elif item.get('content'):
                            f.write(f"{item['content']}\n\n")
            
            # Add other sections
            for section in ['delivery_figures', 'speed_and_consumption', 'lifting_equipment', 'hseq_documentation']:
                if vessel_data.get(section, {}).get('content'):
                    f.write(f"## {section.replace('_', ' ').title()}\n\n")
                    f.write(vessel_data[section]['content'])
                    f.write("\n\n")

def main():
    """Main function to run the optimized vessel data extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Optimized Vessel Data Extractor with Semantic Enhancement')
    parser.add_argument('--source', default='source/vessels', help='Source directory containing vessel DOCX files')
    parser.add_argument('--output', default='output/vessels', help='Output directory for processed data')
    args = parser.parse_args()
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('vessel_extraction_optimized.log'),
            logging.StreamHandler()
        ]
    )
    
    # Create extractor and process files
    extractor = OptimizedVesselDataExtractor(args.source)
    vessels_data = extractor.process_files()
    
    if vessels_data:
        extractor.save_outputs(vessels_data, args.output)
        print(f"\n✅ Successfully processed {len(vessels_data)} vessels with semantic enhancement!")
        print(f"📁 Output saved to: {args.output}")
        print(f"🚀 Ready for optimized AI search!")
    else:
        print("❌ No vessels were processed")

if __name__ == "__main__":
    main() 