#!/usr/bin/env python3
"""
Latest Vessel Data Extractor - Heading-Based Chunking
Focus: Use Word heading styles for intelligent content chunking
"""

import os
import json
import logging
import re
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.styles.style import _ParagraphStyle
from typing import Dict, List, Any, Tuple, Optional

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

class HeadingBasedVesselExtractor:
    """Vessel data extraction using Word heading styles for chunking"""
    
    def __init__(self, source_dir: str = 'source/vessels'):
        self.source_dir = source_dir
        self.extracted_data = {}

    def _get_heading_level(self, paragraph: Paragraph) -> Optional[int]:
        """Extract heading level from Word paragraph"""
        if not hasattr(paragraph, 'style') or not paragraph.style:
            return None
            
        style_name = paragraph.style.name
        if not style_name or not isinstance(style_name, str):
            return None
            
        # Check for Word heading styles
        if style_name.startswith('Heading '):
            try:
                return int(style_name.split(' ')[1])
            except (IndexError, ValueError):
                return None
                
        return None

    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to handle strikethrough formatting"""
        if not hasattr(paragraph, 'runs'):
            return paragraph.text if hasattr(paragraph, 'text') else str(paragraph)
        
        result_parts = []
        current_strikethrough_parts = []
        in_strikethrough = False
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Check if this run has strikethrough formatting
            is_strikethrough = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if strike_elem is not None or dstrike_elem is not None:
                        is_strikethrough = True
            
            if is_strikethrough:
                if not in_strikethrough:
                    in_strikethrough = True
                current_strikethrough_parts.append(text)
            else:
                if in_strikethrough:
                    if current_strikethrough_parts:
                        result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
                        current_strikethrough_parts = []
                    in_strikethrough = False
                result_parts.append(text)
        
        # Handle any remaining strikethrough text
        if current_strikethrough_parts:
            result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
        
        return ''.join(result_parts)

    def _identify_content_type(self, heading: str) -> str:
        """Identify content type from heading with enhanced semantic understanding"""
        heading_lower = heading.lower()
        
        # Charter Party Clauses - More specific identification
        if "charter party" in heading_lower:
            if "cargo exclusions" in heading_lower or "excluded cargo" in heading_lower:
                return "charter_party_cargo_exclusions"
            elif "trading exclusions" in heading_lower or "trading limits" in heading_lower:
                return "charter_party_trading_exclusions"
            elif "duration" in heading_lower or "description" in heading_lower:
                return "charter_party_duration_description"
            elif "delivery" in heading_lower:
                return "charter_party_delivery"
            elif "redelivery" in heading_lower:
                return "charter_party_redelivery"
            elif "hire" in heading_lower or "payment" in heading_lower:
                return "charter_party_hire_payment"
            elif "bunkers" in heading_lower:
                return "charter_party_bunkers"
            return "charter_party_clause"
        
        # Fixture Recap Sections - Enhanced categorization
        if "fixture recap" in heading_lower:
            if "trading" in heading_lower or "exclusions" in heading_lower:
                return "fixture_recap_trading_exclusions"
            elif "cargo" in heading_lower:
                return "fixture_recap_cargo_details"
            elif "account" in heading_lower:
                return "fixture_recap_account"
            elif "period" in heading_lower:
                return "fixture_recap_period"
            elif "laycan" in heading_lower:
                return "fixture_recap_laycan"
            elif "delivery" in heading_lower:
                return "fixture_recap_delivery"
            elif "redelivery" in heading_lower:
                return "fixture_recap_redelivery"
            elif "hire" in heading_lower:
                return "fixture_recap_hire"
            elif "bunkers" in heading_lower:
                return "fixture_recap_bunkers"
            return "fixture_recap_general"
        
        # Vessel Specifications - Detailed categorization
        if any(term in heading_lower for term in ["vessel", "specification"]):
            if "name" in heading_lower:
                return "vessel_name"
            elif "registration" in heading_lower:
                return "vessel_registration"
            elif "tonnage" in heading_lower:
                return "vessel_tonnage"
            elif "dimensions" in heading_lower:
                return "vessel_dimensions"
            elif "equipment" in heading_lower:
                return "vessel_equipment"
            return "vessel_specification"
        
        # Cargo Related Content - Specific identification
        if "cargo" in heading_lower:
            if "exclusions" in heading_lower:
                return "cargo_exclusions_primary"  # Mark as primary source
            elif "restrictions" in heading_lower:
                return "cargo_restrictions"
            elif "requirements" in heading_lower:
                return "cargo_requirements"
            return "cargo_related"
        
        # Trading Related Content
        if "trading" in heading_lower:
            if "exclusions" in heading_lower or "restrictions" in heading_lower:
                return "trading_exclusions_primary"  # Mark as primary source
            elif "limits" in heading_lower:
                return "trading_limits"
            return "trading_related"
        
        return "general"

    def _extract_keywords(self, chunk: Dict[str, Any]) -> List[str]:
        """Extract enhanced keywords from heading and content"""
        keywords = set()
        text = f"{chunk['heading']} {chunk['content']}".lower()
        
        # Add heading terms
        keywords.update(word.strip() for word in chunk['heading'].lower().split() 
                       if len(word.strip()) > 3)
        
        # Add content type as keyword
        content_type = chunk['content_type'].replace('_', ' ')
        keywords.update(content_type.split())
        
        # Maritime-specific term detection
        maritime_terms = {
            'charter_party': ['charter', 'chartering', 'charterer', 'charter party', 'nype'],
            'vessel': ['vessel', 'ship', 'maritime', 'mv', 'motor vessel'],
            'cargo': ['cargo', 'goods', 'merchandise', 'commodities', 'shipment'],
            'cargo_exclusions': ['exclusions', 'excluded', 'prohibited', 'restrictions', 'not allowed'],
            'dangerous_cargo': ['dangerous', 'hazardous', 'flammable', 'explosive', 'radioactive'],
            'trading': ['trading', 'trade', 'voyage', 'route', 'destination'],
            'documentation': ['documentation', 'certificates', 'papers', 'requirements'],
            'operations': ['loading', 'discharge', 'stowage', 'handling', 'operation']
        }
        
        # Add relevant maritime terms based on content
        for category, terms in maritime_terms.items():
            if any(term in text for term in terms):
                keywords.update(terms)
        
        # Extract specific cargo types mentioned
        cargo_patterns = [
            r'(?:carrying|transporting)\s+([a-zA-Z\s]+(?:cargo|goods|materials))',
            r'excluding\s+([a-zA-Z\s,]+)(?=\.)',
            r'prohibited\s+([a-zA-Z\s,]+)(?=\.)',
            r'([A-Z][A-Z\s/]+)(?=\s+(?:is|are)\s+(?:allowed|prohibited|excluded))'
        ]
        
        for pattern in cargo_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                cargo_type = match.group(1).strip().lower()
                if len(cargo_type) > 3:
                    keywords.add(cargo_type)
        
        # Add semantic markers for primary clauses
        if chunk['content_type'].endswith('_primary'):
            keywords.update(['primary clause', 'authoritative source', 'main section'])
        
        # Remove duplicates and sort
        return sorted(list(keywords))

    def _create_enhanced_searchable_text(self, chunk: Dict[str, Any]) -> str:
        """Create semantically enhanced searchable text"""
        text_parts = []
        
        # Add document context
        text_parts.append(f"VESSEL: {chunk['vessel']}")
        
        # Add hierarchical context
        if chunk.get('parent_headings'):
            text_parts.append(f"SECTION PATH: {' > '.join(chunk['parent_headings'])}")
        
        # Add content type context
        text_parts.append(f"CONTENT TYPE: {chunk['content_type'].replace('_', ' ')}")
        
        # Add heading with semantic marker
        text_parts.append(f"HEADING: {chunk['heading']}")
        
        # Add content with semantic marker
        text_parts.append(f"CONTENT: {chunk['content']}")
        
        # Add keywords with semantic marker
        text_parts.append(f"KEYWORDS: {', '.join(chunk['keywords'])}")
        
        # Add semantic markers for primary clauses
        if chunk['content_type'].endswith('_primary'):
            text_parts.append("DOCUMENT ROLE: Primary authoritative source for this subject")
        
        return "\n\n".join(text_parts)

    def _extract_text_and_structure(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text and structure from DOCX using heading styles"""
        try:
            doc = Document(file_path)
            chunks = []
            current_chunk = None
            current_content = []
            
            # Track heading hierarchy
            heading_stack = []
            
            # First pass - extract vessel name from early sections
            vessel_name = None
            for i, paragraph in enumerate(doc.paragraphs[:20]):  # Check first 20 paragraphs
                text = self._process_paragraph_for_strikethrough(paragraph)
                if "vessel" in text.lower() and "name" in text.lower():
                    match = re.search(r'(?:vessel|ship)\s+name:?\s*(.+)', text, re.IGNORECASE)
                    if match:
                        vessel_name = match.group(1).strip()
                        break
                elif "MV" in text or "M/V" in text:
                    match = re.search(r'(?:MV|M/V)\s+([A-Z\s]+)', text)
                    if match:
                        vessel_name = f"MV {match.group(1).strip()}"
                        break
            
            # If no vessel name found, extract from filename
            if not vessel_name:
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                vessel_name = re.sub(r'^(mv|m/v|vessel)[\s_-]*', '', base_name, flags=re.IGNORECASE)
                vessel_name = re.sub(r'[\s_-]*(final|ns|done).*$', '', vessel_name, flags=re.IGNORECASE)
                vessel_name = vessel_name.replace('_', ' ').replace('-', ' ').strip()
                if vessel_name and not vessel_name.upper().startswith('MV'):
                    vessel_name = f"MV {vessel_name.upper()}"
            
            # Second pass - process content with known vessel name
            for paragraph in doc.paragraphs:
                heading_level = self._get_heading_level(paragraph)
                text = self._process_paragraph_for_strikethrough(paragraph)
                
                if heading_level is not None:
                    # Save previous chunk if exists
                    if current_chunk and current_content:
                        current_chunk["content"] = "\n\n".join(current_content)
                        current_chunk["vessel"] = vessel_name
                        # Add enhanced keywords and searchable text
                        current_chunk["keywords"] = self._extract_keywords(current_chunk)
                        current_chunk["searchable_text"] = self._create_enhanced_searchable_text(current_chunk)
                        chunks.append(current_chunk)
                    
                    # Update heading stack
                    while heading_stack and heading_stack[-1]["level"] >= heading_level:
                        heading_stack.pop()
                    
                    # Create new chunk
                    current_chunk = {
                        "heading": text.strip(),
                        "heading_level": heading_level,
                        "parent_headings": [h["heading"] for h in heading_stack],
                        "content_type": self._identify_content_type(text),
                        "vessel": vessel_name,  # Add vessel name here
                        "source_file": os.path.basename(file_path)
                    }
                    
                    # Update heading stack
                    heading_stack.append({"level": heading_level, "heading": text.strip()})
                    current_content = []
                elif text.strip() and current_chunk is not None:
                    current_content.append(text.strip())
            
            # Save last chunk
            if current_chunk and current_content:
                current_chunk["content"] = "\n\n".join(current_content)
                current_chunk["vessel"] = vessel_name
                # Add enhanced keywords and searchable text
                current_chunk["keywords"] = self._extract_keywords(current_chunk)
                current_chunk["searchable_text"] = self._create_enhanced_searchable_text(current_chunk)
                chunks.append(current_chunk)
            
            return chunks
            
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return []

    def _extract_vessel_name(self, chunks: List[Dict]) -> str:
        """Extract vessel name from chunks"""
        for chunk in chunks:
            if "vessel" in chunk["content_type"].lower() and "name" in chunk["content_type"].lower():
                return chunk["content"].strip()
                
            if "MV" in chunk["content"] or "M/V" in chunk["content"]:
                match = re.search(r"(?:MV|M/V)\s+([A-Z\s]+)", chunk["content"])
                if match:
                    return f"MV {match.group(1).strip()}"
        
        return "UNKNOWN_VESSEL"

    def process_files(self) -> Dict[str, List[Dict[str, Any]]]:
        """Process all vessel files with heading-based chunking"""
        vessels_data = {}
        
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return vessels_data
        
        docx_files = [f for f in os.listdir(self.source_dir) 
                     if f.lower().endswith('.docx') and not f.startswith('~$')]
        logging.info(f"Found {len(docx_files)} DOCX files to process")
        
        for filename in docx_files:
            file_path = os.path.join(self.source_dir, filename)
            logging.info(f"Processing {filename}")
            
            try:
                # Extract text with heading structure
                chunks = self._extract_text_and_structure(file_path)
                
                if chunks:
                    # Extract vessel name from chunks
                    vessel_name = self._extract_vessel_name(chunks)
                    
                    # Add vessel and source info to chunks
                    for chunk in chunks:
                        chunk["vessel"] = vessel_name
                        chunk["source_file"] = filename
                        
                        # Create searchable text combining all relevant fields
                        chunk["searchable_text"] = f"{vessel_name} {chunk['heading']} {chunk['content']} {' '.join(chunk['keywords'])}"
                    
                    vessels_data[vessel_name] = chunks
                    logging.info(f"Successfully extracted {len(chunks)} chunks for {vessel_name}")
                else:
                    logging.warning(f"No chunks extracted from {filename}")
                    
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
        
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels")
        return vessels_data

    def save_outputs(self, vessels_data: Dict[str, List[Dict[str, Any]]], output_dir: str = 'output/vessels'):
        """Save processed vessel data"""
        os.makedirs(output_dir, exist_ok=True)
        
        for vessel_name, chunks in vessels_data.items():
            vessel_dir = os.path.join(output_dir, vessel_name)
            os.makedirs(vessel_dir, exist_ok=True)
            
            # Save JSON
            json_file = os.path.join(vessel_dir, f"{vessel_name}_data.json")
            vessel_data = {
                "vessel_name": vessel_name,
                "total_chunks": len(chunks),
                "content_chunks": chunks
            }
            
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(vessel_data, f, indent=2, ensure_ascii=False)
            
            # Save markdown
            md_file = os.path.join(vessel_dir, f"{vessel_name}_data.md")
            self._save_markdown(vessel_data, md_file)
            
            logging.info(f"Saved data for {vessel_name}")

    def _save_markdown(self, vessel_data: Dict[str, Any], md_file: str):
        """Save markdown format with heading structure"""
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(f"# {vessel_data['vessel_name']}\n\n")
            f.write(f"**Total Sections:** {vessel_data['total_chunks']}\n\n")
            
            # Group chunks by heading level
            current_level = 0
            for chunk in vessel_data['content_chunks']:
                heading_level = chunk['heading_level']
                heading_prefix = '#' * (heading_level + 1)  # Add 1 to account for document title
                
                # Write heading with appropriate level
                f.write(f"\n{heading_prefix} {chunk['heading']}\n\n")
                
                # Write metadata
                f.write(f"**Content Type:** {chunk['content_type']}\n\n")
                f.write(f"**Keywords:** {', '.join(chunk['keywords'])}\n\n")
                if chunk['parent_headings']:
                    f.write(f"**Parent Sections:** {' > '.join(chunk['parent_headings'])}\n\n")
                
                # Write content
                f.write(f"{chunk['content']}\n\n")
                f.write("---\n\n")

def main():
    """Main function to run heading-based vessel extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Latest Vessel Data Extractor - Heading-Based Chunking')
    parser.add_argument('--source', default='source/vessels', help='Source directory')
    parser.add_argument('--output', default='output/vessels', help='Output directory')
    args = parser.parse_args()
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('vessel_extraction_latest.log'),
            logging.StreamHandler()
        ]
    )
    
    # Create extractor and process files
    extractor = HeadingBasedVesselExtractor(args.source)
    vessels_data = extractor.process_files()
    
    if vessels_data:
        extractor.save_outputs(vessels_data, args.output)
        print(f"\n✅ Successfully processed {len(vessels_data)} vessels with heading-based chunking!")
        print(f"📁 Output saved to: {args.output}")
        print(f"🔍 Ready for enhanced AI search!")
    else:
        print("❌ No vessels were processed")

if __name__ == "__main__":
    main() 