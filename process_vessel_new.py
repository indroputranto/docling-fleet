#!/usr/bin/env python3
"""
Dynamic Text Extraction System
Extracts all text from source files and outputs to JSON and MD formats
Uses dynamic file detection and processing without hardcoding
"""

import os
import json
import logging
from typing import List, Dict, Any, Optional, Tuple
import re
from pathlib import Path


class DynamicTextExtractor:
    """Dynamic text extraction system that handles multiple file formats"""
    
    def __init__(self, source_dir: str = 'source/vessels', output_dir: str = 'output/vessels'):
        """Initialize the dynamic text extractor"""
        self.source_dir = source_dir
        self.output_dir = output_dir
        self.supported_formats = {}
        self.extracted_data = {}
        
        # Set up logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        
        # Dynamically detect available text extraction libraries
        self._detect_available_libraries()
    
    def _detect_available_libraries(self):
        """Dynamically detect which text extraction libraries are available"""
        available_libs = {}
        
        # Check for DOCX support
        try:
            from docx import Document
            available_libs['docx'] = Document
            self.supported_formats['.docx'] = self._extract_from_docx
            logging.info("DOCX support available")
        except ImportError:
            logging.warning("python-docx not available - DOCX files will not be processed")
        
        # Check for PDF support
        try:
            import PyPDF2
            available_libs['pypdf2'] = PyPDF2
            self.supported_formats['.pdf'] = self._extract_from_pdf_pypdf2
            logging.info("PDF support available (PyPDF2)")
        except ImportError:
            try:
                import pdfminer
                available_libs['pdfminer'] = pdfminer
                self.supported_formats['.pdf'] = self._extract_from_pdf_pdfminer
                logging.info("PDF support available (pdfminer)")
            except ImportError:
                logging.warning("No PDF libraries available - PDF files will not be processed")
        
        # Check for RTF support
        try:
            from striprtf.striprtf import rtf_to_text
            available_libs['rtf'] = rtf_to_text
            self.supported_formats['.rtf'] = self._extract_from_rtf
            logging.info("RTF support available")
        except ImportError:
            logging.warning("striprtf not available - RTF files will not be processed")
        
        # Check for TXT support (always available)
        self.supported_formats['.txt'] = self._extract_from_txt
        logging.info("TXT support available")
        
        # Check for CSV support
        try:
            import csv
            available_libs['csv'] = csv
            self.supported_formats['.csv'] = self._extract_from_csv
            logging.info("CSV support available")
        except ImportError:
            pass
        
        self.available_libraries = available_libs
        logging.info(f"Supported file formats: {list(self.supported_formats.keys())}")
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files with dynamic numbering preservation using docx2python"""
        try:
            # Use docx2python for better structure preservation including numbering
            try:
                from docx2python import docx2python
                logging.info("Using docx2python for dynamic numbering extraction")
                
                # Extract with docx2python to preserve numbering structure
                docx_content = docx2python(file_path)
                
                # Extract paragraphs from the body content
                paragraphs = []
                
                def extract_text_recursively(content, level=0):
                    """Recursively extract text from nested content structure"""
                    extracted_items = []
                    
                    if isinstance(content, str):
                        # Clean and add non-empty strings
                        text = content.strip()
                        if text:
                            extracted_items.append(text)
                    elif isinstance(content, (list, tuple)):
                        # Process nested lists/tuples
                        for item in content:
                            extracted_items.extend(extract_text_recursively(item, level + 1))
                    
                    return extracted_items
                
                # Extract all text content from the document body
                all_text_items = extract_text_recursively(docx_content.body)
                
                # Extract tables first to identify table content that should be excluded from paragraphs
                table_cells_content = set()
                try:
                    from docx import Document
                    doc_for_tables = Document(file_path)
                
                    for table in doc_for_tables.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                # Collect all table cell content to exclude from paragraph extraction
                                for p in cell.paragraphs:
                                    if p.text.strip():
                                        processed_text = self._process_paragraph_for_strikethrough(p)
                                        table_cells_content.add(processed_text.strip())
                                        table_cells_content.add(p.text.strip())  # Also add original text
                                        
                except Exception as table_check_error:
                    logging.warning(f"Error identifying table content: {table_check_error}")
                
                # Filter and clean the extracted text, excluding table content
                for text_item in all_text_items:
                    clean_text = text_item.strip()
                    
                    # Remove "Working Copy" artifacts
                    if clean_text == "Working Copy":
                        continue
                    
                    # Skip table cell content that's already captured in tables
                    if clean_text in table_cells_content:
                        continue
                    
                    # Convert any [[STRIKE]] patterns to double tildes for consistency
                    clean_text = re.sub(r'\[\[STRIKE\s+(.*?)\s+STRIKE\]\]', r'~~\1~~', clean_text, flags=re.DOTALL | re.IGNORECASE)
                    
                    if clean_text and len(clean_text) > 1:  # Skip very short items
                        paragraphs.append(clean_text)
                        
                        # Log when we find numbered items (both formats: "XX." and "XX)")
                        if re.match(r'^\d+[\.\)]', clean_text):
                            logging.info(f"Dynamically extracted numbered item: {clean_text[:50]}...")
                
                logging.info(f"docx2python extracted {len(paragraphs)} text items")
                
                # ENHANCEMENT: Now also try to capture missing numbering using python-docx XML parsing
                # This handles cases where docx2python might miss Word's automatic numbering
                try:
                    from docx import Document
                    from docx.oxml.ns import qn
                    
                    doc = Document(file_path)
                    numbering_enhanced_paragraphs = []
                    
                    # Track numbering sequences for different numbering definitions
                    numbering_counters = {}
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text = para.text.strip()
                            
                            # FIRST: Process strikethrough formatting
                            text_with_strikethrough = self._process_paragraph_for_strikethrough(para)
                            final_text = text_with_strikethrough
                            
                            # SECOND: Check if this paragraph has numbering properties in XML
                            try:
                                p_xml = para._element
                                numPr = p_xml.find(qn('w:numPr'))
                                
                                if numPr is not None:
                                    # Extract numbering ID and level
                                    numId_elem = numPr.find(qn('w:numId'))
                                    ilvl_elem = numPr.find(qn('w:ilvl'))
                                    
                                    if numId_elem is not None and ilvl_elem is not None:
                                        numId = numId_elem.get(qn('w:val'))
                                        ilvl = int(ilvl_elem.get(qn('w:val')))
                                        
                                        # Create a unique key for this numbering sequence
                                        num_key = f"{numId}_{ilvl}"
                                        
                                        # Initialize or increment counter for this numbering sequence
                                        if num_key not in numbering_counters:
                                            numbering_counters[num_key] = 1
                                        else:
                                            numbering_counters[num_key] += 1
                                        
                                        current_number = numbering_counters[num_key]
                                        
                                        # Only add numbering if the text doesn't already have it (apply to strikethrough-processed text)
                                        if not re.match(r'^\d+\.', text):
                                            final_text = f"{current_number}. {text_with_strikethrough}"
                                            logging.info(f"XML-based numbering reconstruction: {current_number}. for {text[:50]}...")
                                        
                                        # Replace the text in paragraphs if it exists (search for original text, replace with enhanced)
                                        for i, existing_text in enumerate(paragraphs):
                                            if existing_text.strip() == text:
                                                paragraphs[i] = final_text
                                                break
                                        else:
                                            # Add if not found
                                            numbering_enhanced_paragraphs.append(final_text)
                                            
                            except Exception as xml_e:
                                logging.debug(f"XML numbering extraction error: {xml_e}")
                                continue
                    
                    # Add any newly discovered numbered items
                    if numbering_enhanced_paragraphs:
                        paragraphs.extend(numbering_enhanced_paragraphs)
                        logging.info(f"XML enhancement added {len(numbering_enhanced_paragraphs)} numbered items")
                        
                except Exception as enhance_error:
                    logging.warning(f"XML numbering enhancement failed: {enhance_error}")
                
            except ImportError:
                logging.warning("docx2python not available, falling back to python-docx")
                # Fallback to python-docx if docx2python is not available
                from docx import Document
                doc = Document(file_path)
                
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
            
            except Exception as e:
                logging.error(f"Error with docx2python, falling back to python-docx: {e}")
                # Fallback to python-docx
                from docx import Document
                doc = Document(file_path)
                
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
            
            # Extract tables using python-docx (docx2python table handling is more complex)
            tables = []
            try:
                from docx import Document
                doc = Document(file_path)
                
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            # Process each paragraph in the cell for strikethrough
                            cell_text_parts = []
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    processed_text = self._process_paragraph_for_strikethrough(p)
                                    cell_text_parts.append(processed_text)
                            cell_text = ' '.join(cell_text_parts)
                            row_data.append(cell_text)
                        if any(cell.strip() for cell in row_data):  # Only add non-empty rows
                            table_data.append(row_data)
                    
                    # Only include tables - let the chapter assignment logic filter what's actually a table
                    if table_data:
                        tables.append(table_data)
                        
            except Exception as table_error:
                logging.warning(f"Error extracting tables: {table_error}")
            
            # CRITICAL: Apply strikethrough processing to all paragraphs
            # This ensures that strikethrough formatting is preserved from docx2python extraction
            try:
                from docx import Document
                doc = Document(file_path)
                
                # Create a mapping of text content to paragraphs for strikethrough processing
                para_text_map = {}
                for para in doc.paragraphs:
                    if para.text.strip():
                        para_text_map[para.text.strip()] = para
                
                # Process each paragraph for strikethrough
                strikethrough_processed_paragraphs = []
                for text in paragraphs:
                    processed_text = text
                    
                    # Try to find matching paragraph object for strikethrough processing
                    text_clean = text.strip()
                    if text_clean in para_text_map:
                        processed_text = self._process_paragraph_for_strikethrough(para_text_map[text_clean])
                        if '~~' in processed_text:
                            logging.info(f"Applied strikethrough processing to: {text_clean[:50]}...")
                    else:
                        # Fallback: just check for any [[STRIKE]] patterns that might have been missed
                        processed_text = re.sub(r'\[\[STRIKE\s+(.*?)\s+STRIKE\]\]', r'~~\1~~', text, flags=re.DOTALL | re.IGNORECASE)
                    
                    strikethrough_processed_paragraphs.append(processed_text)
                
                paragraphs = strikethrough_processed_paragraphs
                logging.info(f"Applied strikethrough processing to all {len(paragraphs)} paragraphs")
                
            except Exception as strikethrough_error:
                logging.warning(f"Could not apply strikethrough processing: {strikethrough_error}")
            
            full_text = '\n\n'.join(paragraphs)
            
            return {
                'full_text': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }
        except Exception as e:
            logging.error(f"Error extracting from DOCX {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _looks_like_clause_header(self, text: str) -> bool:
        """Determine if a text line looks like a clause header that should be numbered"""
        import re
        
        # Clean text for analysis
        text_clean = text.strip()
        text_lower = text_clean.lower()
        
        # Skip very short texts
        if len(text_clean) < 10:
            return False
            
        # Skip text that already has numbering
        if re.match(r'^\d+\.', text_clean):
            return False
            
        # Indicators that this is likely a clause header
        clause_indicators = [
            # Common clause types
            r'\bclause\b',
            r'\bregulation\b',
            r'\bprovision\b',
            r'\bagreement\b',
            r'\bcontract\b',
            r'\bcharter\b',
            
            # Specific clause patterns
            r'\bbimco\b',
            r'\bcensus\s+bureau\b',
            r'\bexport\s+system\b',
            r'\bfuel\s+eu\b',
            r'\bship\s+sale\b',
            r'\bvessel.*description\b',
            r'\btime\s+charter\b',
            r'\bmandatory\b',
            r'\bautomated\b',
            
            # Text ending patterns that suggest clause headers
            r':\s*$',  # Ends with colon
            r'\.\s*$',  # Ends with period
        ]
        
        # Check if any indicator is present
        for pattern in clause_indicators:
            if re.search(pattern, text_lower):
                return True
                
        # Additional heuristics
        # Check if it's ALL CAPS (common for clause headers)
        if text_clean.isupper() and len(text_clean.split()) >= 2:
            return True
            
        # Check if it starts with certain keywords
        starts_with_keywords = [
            'bimco', 'the vessel', 'charterers', 'owners', 'master',
            'cargo', 'freight', 'demurrage', 'laytime', 'notice'
        ]
        
        for keyword in starts_with_keywords:
            if text_lower.startswith(keyword):
                return True
                
        return False
    
    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to detect strikethrough formatting and wrap in double tildes (~~ ~~)"""
        if not hasattr(paragraph, 'runs'):
            return paragraph.text if hasattr(paragraph, 'text') else str(paragraph)
        
        result_parts = []
        current_strikethrough_parts = []
        in_strikethrough = False
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Check if this run has strikethrough formatting
            is_strikethrough = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Check for strike elements
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if strike_elem is not None or dstrike_elem is not None:
                        is_strikethrough = True
            
            if is_strikethrough:
                if not in_strikethrough:
                    # Starting a new strikethrough section
                    in_strikethrough = True
                current_strikethrough_parts.append(text)
            else:
                if in_strikethrough:
                    # Ending strikethrough section
                    if current_strikethrough_parts:
                        result_parts.append(f"~~{''.join(current_strikethrough_parts)}~~")
                        current_strikethrough_parts = []
                    in_strikethrough = False
                result_parts.append(text)
        
        # Handle any remaining strikethrough text
        if current_strikethrough_parts:
            result_parts.append(f"~~{''.join(current_strikethrough_parts)}~~")
        
        return ''.join(result_parts)
    
    def _organize_charter_party_clauses(self, charter_party_content: List[str]) -> tuple[Dict[str, Dict[str, Any]], List[str]]:
        """Organize Charter Party content into individual clauses using semantic content-based identification"""
        import re
        
        clauses = {}
        non_clause_content = []  # Content that doesn't belong to any numbered clause
        current_clause_number = None
        current_clause_title = None
        current_clause_content = []
        clause_counter = 1
        skip_next_lines = 0  # Counter to skip lines that are document structure labels
        
        # Define section headers that should be treated as clause delimiters
        clause_headers = [
            'Duration', 'Trip Description', 'Trading Exclusions', 'Trading Limits',  # Added missing clause headers
            'Delivery', 'Redelivery', 'Laydays', 'Cancelling', 'Laycan', # Added Laydays/Cancelling variants
            'On/Oﬀ-Hire Survey', 'On-Hire Survey', 'Off-Hire Survey', 'Hire Survey', # Added survey variants
            'Owners to Provide', 'Charterers to Provide',
            'Performance of Voyages', 'Bunkers', 'Rate of Hire', 'Hold Cleaning', 'Communications',
            'Victualing and Expenses', 'Victualling and Expenses', 'Hire Payment', 'Speed and Consumption', 'Spaces Available',
            'Supercargo', 'Sailing Orders and Logs', 'Cargo Exclusions', 'OFF HIRE CLAUSE', 'Off-Hire',
            'Pollution', 'Drydocking', 'Total Loss', 'Exceptions', 'Liberties', 'Liens',
            'Salvage', 'General Average', 'Navigation', 'Cargo Claims', 'Cargo Handling Gear',
            'Solid Bulk Cargoes', 'BIMCO Hull Fouling', 'Bills of Lading', 'Protective Clauses',
            'BIMCO War Risks', 'Ice', 'Requisition', 'Stevedore Damage', 'Slow Steaming',
            'BIMCO Piracy', 'Taxes', 'Industrial Action', 'Stowaways', 'Smuggling',
            'International Safety Management', 'International Ship and Port Facility',
            'ISPS Code', 'MTSA', 'Maritime Transportation Security Act',  # Recurring missing clauses
            'War Cancellation', 'FuelEU Maritime', 'EU ETS', 'Victualling',  # Spelling variant
            'Dangerous Goods', 'Boycott', 'Blacklist', 'C-TPAT', 'Fumigation',
            'Reefer Container', 'Canadian Clause Paramount', 'CONWARTIME', 'High Risk Area',
            'Additional Information', 'Owners responsibility', 'Owners to provide',
            'Bill Of Lading Identifier', 'U.S. Trade', 'C-TPAT', 'CUSTOMS-TRADE',
            'Sanctions', 'BIMCO Designated Entities', 'BIMCO North American Advance',
            'P&C', 'Privacy', 'Confidential', 'Privacy and Confidentiality',
            # FWN Sea and similar: additional clause headers for full capture
            'Crew Overtime', 'United States Regulations', 'Arrest', 'Detained', 'Vessels hatch cover',
            'BIMCO ISPS', 'ISM Code', 'Oil Pollution', 'P&I Club', 'ITF', 'Vaccination',
            'Master supervision', 'Forklift', 'Lay up', 'Ballast', 'Supercargo', 'Port Captain',
            'Lashing equipment', 'Welding', 'Weather Routing', 'Discharge into custody', 'Bill of Lading',
            'Fuel Specifications', 'BIMCO Stowaways', 'Sunday work', 'Bunkering prior to delivery',
            'Bunkering operation', 'Bunker quality', 'Fuel testing programme',
            'BIMCO 2020 Marine Fuel Sulphur', 'Types and quantities of bunkers',
            'BIMCO Sanction Clause', 'BIMCO Infectious', 'BIMCO Sanctions Clause',
            'US Coast Guard', 'Vessel arrested', 'Income Collected', 'Management Fee',
            'BIMCO U.S. Census Bureau', 'BIMCO EU Advance Cargo', 'Ballast Water Exchange',
            'Period Applicable Clauses', 'Law and Arbitration', 'Notices', 'Headings',
            'Singular/Plural', 'BIMCO CII Clause', 'BIMCO ETS', 'FUEL EU Clause'
        ]
        
        for idx, line in enumerate(charter_party_content):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Skip lines that were marked as document structure labels
            if skip_next_lines > 0:
                skip_next_lines -= 1
                # Treat skipped lines as regular content
                if current_clause_number is not None:
                    current_clause_content.append(line_stripped)
                else:
                    non_clause_content.append(line_stripped)
                continue
            
            # Check if this line is a clause header (section delimiter)
            is_clause_header = False
            
            # First, check for "Clause X - Title" pattern 
            clause_dash_match = re.match(r'^Clause\s+(\d+)\s*[-–]\s*(.+)$', line_stripped, re.IGNORECASE)
            if clause_dash_match:
                clause_num = clause_dash_match.group(1)
                clause_title = clause_dash_match.group(2).strip()
                
                # Apply character limit filter - legitimate clause titles are typically under 85 characters
                # Apply colon filter - legitimate clause titles don't contain colons (which indicate content text)
                if len(clause_title) <= 150 and ':' not in clause_title:
                    is_clause_header = True
                    # Save the previous clause BEFORE updating variables
                    if current_clause_number and current_clause_title:
                        clause_key = f"clause_{current_clause_number}"
                        logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                        if clause_key not in clauses:
                            # Only save if we have content, OR if this is a different clause number
                            if current_clause_content or clause_num != current_clause_number:
                                clauses[clause_key] = {
                                    "title": f"Clause {current_clause_number} - {current_clause_title}",
                                    "content": current_clause_content
                                }
                                logging.info(f"Successfully saved {clause_key}")
                            else:
                                logging.warning(f"Skipping save of {clause_key} - no content and same clause number detected again")
                        else:
                            # If clause already exists, only update if we have content to add
                            if current_clause_content:
                                # Merge content if the existing clause has empty content
                                if not clauses[clause_key]["content"]:
                                    clauses[clause_key]["content"] = current_clause_content
                                    logging.info(f"Merged content into existing {clause_key} (was empty)")
                                else:
                                    logging.warning(f"Duplicate clause number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title'][:50]}...")
                            else:
                                # Don't overwrite existing clause with empty content
                                logging.warning(f"Duplicate clause number {current_clause_number} detected with no new content. Keeping existing clause: {clauses[clause_key]['title'][:50]}...")
                                # Don't reset - keep collecting content for the existing clause
                                # Process this line as content for the existing clause instead of as a new clause header
                                is_clause_header = False
                                if current_clause_number is not None:
                                    current_clause_content.append(line_stripped)
                                continue  # Move to next line - treat as content, don't start new clause
                        if is_clause_header:
                            # NOW update to new clause (only if we're not skipping due to duplicate)
                            current_clause_number = clause_num
                            current_clause_title = clause_title
                            current_clause_content = []  # Reset for new clause
                            logging.info(f"Found 'Clause X - Title' format Charter Party clause {current_clause_number}: {clause_title[:50]}...")
            
            # Second, check for numbered patterns like "2) Delivery" or "3) Redelivery" 
            else:
                numbered_match = re.match(r'^(\d+)[\.\)]\s*(.+)$', line_stripped)
                # Also check specifically for tab-separated clauses which are very common
                # Handle both regular tabs and tabs with special Unicode characters (like zero-width space)
                tab_clause_match = re.match(r'^(\d+)\)\s*\t+\s*(.+)$', line_stripped)
                
                # Process tab-separated clauses first (they're a special case of numbered clauses)
                if tab_clause_match:
                    clause_num = tab_clause_match.group(1)
                    clause_title = tab_clause_match.group(2).strip()
                    
                    # Clean up special Unicode characters that might interfere
                    # Remove zero-width space and other problematic Unicode chars
                    clause_title = re.sub(r'[\u200b-\u200f\ufeff]', '', clause_title).strip()
                    
                    # Check if this is a numbered list item (same logic as for regular numbered patterns)
                    # Exception: Long clause titles (BIMCO, ETS, FuelEU, "for Time Charter Parties") are always headers
                    is_list_item = False
                    t_upper = clause_title.strip().upper()
                    is_long_clause_header = (
                        t_upper.startswith('BIMCO') or t_upper.startswith('ETS') or
                        t_upper.startswith('DURATION') or  # Clause 1: "Duration / Trip Description / Trading exclusions / Trading limits"
                        ' FOR TIME CHARTER PARTIES' in t_upper or ' FOR VOYAGE CHARTER PARTIES' in t_upper
                    )
                    if is_long_clause_header:
                        pass  # These are always headers regardless of length
                    elif any(header.lower() in clause_title.lower() for header in clause_headers):
                        pass  # Known clause headers (Rate of Hire, ISPS, etc.) are always headers regardless of length
                    elif len(clause_title) > 50 or (len(clause_title) > 0 and clause_title[0].islower()):
                        # Exception: Long titles that look like clause headers (contain "Clause", parentheticals)
                        looks_like_clause_title = (
                            clause_title[0].isupper() and
                            ('clause' in clause_title.lower() or
                             '(s)' in clause_title or '(including' in clause_title.lower() or
                             re.search(r'\([^)]{2,50}\)', clause_title)))
                        if looks_like_clause_title:
                            pass  # Treat as clause header
                        else:
                            is_list_item = True
                            logging.info(f"Detected tab-separated numbered list item (treating as content): '{line_stripped[:60]}...'")
                    elif not any(header.lower() in clause_title.lower() for header in clause_headers):
                        list_indicators = ['shall', 'will', 'must', 'may', 'should', 'can', 'comply', 'ensure']
                        if any(indicator in clause_title.lower() for indicator in list_indicators):
                            is_list_item = True
                            logging.info(f"Detected tab-separated numbered list item by verb (treating as content): '{line_stripped[:60]}...'")
                    
                    # If this is a list item, treat it as content
                    if is_list_item:
                        if current_clause_number is not None:
                            current_clause_content.append(line_stripped)
                        else:
                            non_clause_content.append(line_stripped)
                        # Skip further processing for this line
                        continue
                    
                    # Tab-separated clauses are very likely legitimate, so be less restrictive
                    # Increased limit to accommodate longer clause titles like clause 52
                    if (len(clause_title) <= 200 and len(clause_title) > 1):
                        is_clause_header = True
                        # Save the previous clause BEFORE updating variables
                        if current_clause_number and current_clause_title:
                            clause_key = f"clause_{current_clause_number}"
                            logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                            if clause_key not in clauses:
                                # Only save if we have content, OR if this is a different clause number
                                if current_clause_content or clause_num != current_clause_number:
                                    clauses[clause_key] = {
                                        "title": f"{current_clause_number}) {current_clause_title}",
                                        "content": current_clause_content
                                    }
                                    logging.info(f"Successfully saved {clause_key}")
                                else:
                                    logging.warning(f"Skipping save of {clause_key} - no content and same clause number detected again")
                            else:
                                # If clause already exists, only update if we have content to add
                                if current_clause_content:
                                    # Merge content if the existing clause has empty content
                                    if not clauses[clause_key]["content"]:
                                        clauses[clause_key]["content"] = current_clause_content
                                        logging.info(f"Merged content into existing {clause_key} (was empty)")
                                    else:
                                        logging.warning(f"Duplicate clause number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title']}")
                                else:
                                    # Don't overwrite existing clause with empty content
                                    logging.warning(f"Duplicate clause number {current_clause_number} detected with no new content. Keeping existing clause: {clauses[clause_key]['title']}")
                                    # Don't reset - keep collecting content for the existing clause
                                    # Process this line as content for the existing clause instead of as a new clause header
                                    is_clause_header = False
                                    if current_clause_number is not None:
                                        current_clause_content.append(line_stripped)
                                    continue  # Move to next line - treat as content, don't start new clause
                        if is_clause_header:
                            # Update for new clause (only if we're not skipping due to duplicate)
                            current_clause_number = clause_num
                            current_clause_title = clause_title
                            current_clause_content = []  # Reset for new clause
                            logging.info(f"Found tab-separated Charter Party clause {current_clause_number}: {clause_title[:50]}...")
                
                elif numbered_match:
                    clause_num = numbered_match.group(1)
                    clause_title = numbered_match.group(2).strip()
                    
                    # Special check 1: Detect numbered list items within clause content
                    # List items typically:
                    # - Start with lowercase (e.g., "1. the Vessel shall...")
                    # - Are longer sentences (> 50 chars)
                    # - Contain verbs like "shall", "will", "must", "may"
                    # - Don't match known clause header keywords
                    # Exception: Long clause titles (BIMCO, ETS, Duration, "for Time Charter Parties") are always headers
                    is_list_item = False
                    t_upper = clause_title.strip().upper()
                    is_long_clause_header = (
                        t_upper.startswith('BIMCO') or t_upper.startswith('ETS') or
                        t_upper.startswith('DURATION') or  # Clause 1: "Duration / Trip Description / Trading exclusions / Trading limits"
                        ' FOR TIME CHARTER PARTIES' in t_upper or ' FOR VOYAGE CHARTER PARTIES' in t_upper
                    )
                    if is_long_clause_header:
                        pass  # These are always headers regardless of length
                    elif any(header.lower() in clause_title.lower() for header in clause_headers):
                        pass  # Known clause headers (Rate of Hire, ISPS, etc.) are always headers regardless of length
                    elif len(clause_title) > 50 or (len(clause_title) > 0 and clause_title[0].islower()):
                        # Exception: Long titles that look like clause headers (contain "Clause", parentheticals like (s) or (including...))
                        looks_like_clause_title = (
                            clause_title[0].isupper() and
                            ('clause' in clause_title.lower() or
                             '(s)' in clause_title or '(including' in clause_title.lower() or
                             re.search(r'\([^)]{2,50}\)', clause_title))  # Parenthetical phrase
                        )
                        if looks_like_clause_title:
                            pass  # Treat as clause header
                        else:
                            is_list_item = True
                            logging.info(f"Detected numbered list item (treating as content): '{line_stripped[:60]}...'")
                    elif not any(header.lower() in clause_title.lower() for header in clause_headers):
                        # If it doesn't match any clause headers and contains common list item indicators
                        list_indicators = ['shall', 'will', 'must', 'may', 'should', 'can', 'comply', 'ensure']
                        if any(indicator in clause_title.lower() for indicator in list_indicators):
                            is_list_item = True
                            logging.info(f"Detected numbered list item by verb (treating as content): '{line_stripped[:60]}...'")
                    
                    # Special check 2: If this looks like a multi-topic document label (e.g., "2) Duration / Trip Description / Trading Limits / Trading exclusions")
                    # and the NEXT line is a simpler version (e.g., "Duration/Trip Description"), 
                    # skip this line and let the next line be the actual clause header
                    is_document_label = False
                    if not is_list_item and clause_title.count('/') >= 2 and idx + 1 < len(charter_party_content):
                        next_line = charter_party_content[idx + 1].strip()
                        # Check if next line is a simpler section header (one of the main topics from this label)
                        label_topics = [topic.strip() for topic in clause_title.split('/')]
                        for topic in label_topics:
                            if topic and topic.lower() in next_line.lower() and len(next_line) < len(clause_title):
                                # Next line is likely the actual section header, this is just a document label
                                is_document_label = True
                                logging.info(f"Detected document label (will skip): '{line_stripped[:60]}...' - next line is likely the actual clause header")
                                break
                    
                    # If this is a list item, treat it as content for the current clause
                    if is_list_item:
                        if current_clause_number is not None:
                            current_clause_content.append(line_stripped)
                        else:
                            non_clause_content.append(line_stripped)
                        # Skip further processing for this line
                        continue
                    
                    # Apply character limit filter - legitimate clause titles can be longer (up to 200 chars)
                    # Apply colon filter - legitimate clause titles don't contain colons (which indicate content text)
                    # For numbered patterns, be less restrictive since "X) Title" format is very likely a clause
                    # Increased limit to accommodate longer clause titles like clause 52
                    header_match = any(header.lower() in clause_title.lower() for header in clause_headers)
                    title_length_ok = len(clause_title) >= 2  # Allow short abbreviations like "P&C", "ISPS"
                    if (len(clause_title) <= 200 and ':' not in clause_title and 
                        (header_match or title_length_ok) and not is_document_label):
                        is_clause_header = True
                        # Save the previous clause BEFORE updating variables
                        if current_clause_number and current_clause_title:
                            clause_key = f"clause_{current_clause_number}"
                            logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                            if clause_key not in clauses:
                                # Only save if we have content, OR if this is a different clause number
                                # (Don't save empty clauses unless we're moving to a different clause)
                                if current_clause_content or clause_num != current_clause_number:
                                    clauses[clause_key] = {
                                        "title": f"{current_clause_number}) {current_clause_title}",
                                        "content": current_clause_content
                                    }
                                    logging.info(f"Successfully saved {clause_key}")
                                else:
                                    logging.warning(f"Skipping save of {clause_key} - no content and same clause number detected again")
                            else:
                                # If clause already exists, only update if we have content to add
                                if current_clause_content:
                                    # Merge content if the existing clause has empty content
                                    if not clauses[clause_key]["content"]:
                                        clauses[clause_key]["content"] = current_clause_content
                                        logging.info(f"Merged content into existing {clause_key} (was empty)")
                                    else:
                                        logging.warning(f"Duplicate clause number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title'][:50]}...")
                                else:
                                    # Don't overwrite existing clause with empty content
                                    logging.warning(f"Duplicate clause number {current_clause_number} detected with no new content. Keeping existing clause: {clauses[clause_key]['title'][:50]}...")
                                    # Don't reset - keep collecting content for the existing clause
                                    # Process this line as content for the existing clause instead of as a new clause header
                                    is_clause_header = False
                                    if current_clause_number is not None:
                                        current_clause_content.append(line_stripped)
                                    continue  # Move to next line - treat as content, don't start new clause
                        if is_clause_header:
                            # NOW update to new clause (only if we're not skipping due to duplicate)
                            current_clause_number = clause_num
                            current_clause_title = clause_title
                            current_clause_content = []  # Reset for new clause
                            logging.info(f"Found numbered Charter Party clause {current_clause_number}: {clause_title[:50]}...")
                
                # Third, check for standalone section headers
                elif any(line_stripped.startswith(header) for header in clause_headers):
                    # Apply character limit and colon filter for consistency
                    if len(line_stripped) <= 85 and ':' not in line_stripped:
                        # This is a section header - it's NOT a clause delimiter, just a sub-heading within the current clause
                        # Add it as content to the current clause
                        if current_clause_number is not None:
                            current_clause_content.append(line_stripped)
                            logging.info(f"Found section sub-header '{line_stripped[:50]}...' - adding to current clause {current_clause_number}")
                        else:
                            non_clause_content.append(line_stripped)
                            logging.info(f"Found section sub-header '{line_stripped[:50]}...' - adding to non-clause content")
                
                # Fourth, check for "Charter Party – Clause X" patterns
                elif re.match(r'Charter Party [–-] Clause \d+', line_stripped):
                    # Apply character limit and colon filter for consistency
                    if len(line_stripped) <= 85 and ':' not in line_stripped:
                        is_clause_header = True
                        # Save the previous clause BEFORE updating variables
                        if current_clause_number and current_clause_title:
                            clause_key = f"clause_{current_clause_number}"
                            logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                            if clause_key not in clauses:
                                clauses[clause_key] = {
                                    "title": f"{current_clause_number}) {current_clause_title}",
                                    "content": current_clause_content
                                }
                                logging.info(f"Successfully saved {clause_key}")
                            else:
                                logging.warning(f"Duplicate clause number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title'][:50]}...")
                        # NOW update to new clause
                        current_clause_number = str(clause_counter)
                        current_clause_title = line_stripped
                        current_clause_content = []  # Reset for new clause
                        clause_counter += 1
                        logging.info(f"Found formal Charter Party clause {current_clause_number}: {current_clause_title[:50]}...")
                
                # If none of the above patterns match, this is regular content
                if not is_clause_header:
                    # Add content to the current clause, or to non-clause content if no clause is active
                    if current_clause_number is not None:
                        current_clause_content.append(line_stripped)
                    else:
                        # This content doesn't belong to any numbered clause (e.g., header info, company details)
                        non_clause_content.append(line_stripped)
        
        # Don't forget the last clause
        if current_clause_number and current_clause_title:
            clause_key = f"clause_{current_clause_number}"
            logging.info(f"Final clause - attempting to save {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
            # Check if clause already exists - if so, keep the first one (likely the main clause header)
            if clause_key not in clauses:
                clauses[clause_key] = {
                    "title": f"{current_clause_number}) {current_clause_title}",
                    "content": current_clause_content
                }
                logging.info(f"Final clause - successfully saved {clause_key}")
            else:
                logging.warning(f"Final clause duplicate number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title'][:50]}...")
        
        # Post-processing: split misplaced clause headers that were merged into previous clause content
        clauses = self._split_misplaced_clause_headers(clauses)
        
        return clauses, non_clause_content
    
    def _split_misplaced_clause_headers(self, clauses: Dict[str, Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
        """Post-process clauses: find lines like '10. Rate of Hire; Hold Cleaning...' inside another clause's content
        and split them into their own clause. Fixes recurring issue where long clause titles are merged into prev clause."""
        import re
        # Patterns that indicate a clause header was wrongly treated as content (number. or number) + known title)
        misplaced_header_patterns = [
            r'^(\d+)[\.\)]\s*(Rate of Hire[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:ISPS Code|MTSA|Maritime Transportation Security)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*Solid Bulk Cargoes[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*Dangerous Goods[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*War Cancellation[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*FuelEU Maritime[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*EU ETS[^.]*)$',
            r'^(\d+)[\.\)]\s*(Speed and Consumption[^.]*)$',
            r'^(\d+)[\.\)]\s*(Notices[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*Reefer Container[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:Charterers to Provide|Owners to Provide)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:High Risk Area|Canadian Clause Paramount|CONWARTIME)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:Bill Of Lading Identifier|C-TPAT|CUSTOMS-TRADE)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:Additional Information|Owners responsibility)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*(?:Tweendeck|Grain Bulkhead|Stowaway)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*\(including[^)]*\)[^.]*)$',
            r'^(\d+)[\.\)]\s*(.*\s+Clause\s*)$',  # Titles ending with "Clause"
        ]
        
        # Sort clause keys by number to process in order
        def clause_sort_key(k):
            m = re.match(r'clause_(\d+)', k)
            return int(m.group(1)) if m else 0
        
        sorted_keys = sorted(clauses.keys(), key=clause_sort_key)
        
        for clause_key in sorted_keys:
            content = clauses[clause_key].get("content", [])
            if not content:
                continue
            
            for i, line in enumerate(content):
                line_stripped = line.strip()
                for pattern in misplaced_header_patterns:
                    m = re.match(pattern, line_stripped, re.IGNORECASE)
                    if m:
                        clause_num = m.group(1)
                        clause_title = m.group(2).strip()
                        # Only split if this would create a new clause (num doesn't exist or is different)
                        new_clause_key = f"clause_{clause_num}"
                        if new_clause_key not in clauses or clauses[new_clause_key].get("content") == []:
                            # Split: keep content before this line in current clause
                            before_content = content[:i]
                            after_content = content[i + 1:]
                            clauses[clause_key]["content"] = before_content
                            # Create or update the new clause
                            if new_clause_key not in clauses:
                                clauses[new_clause_key] = {
                                    "title": f"{clause_num}) {clause_title}",
                                    "content": after_content
                                }
                                logging.info(f"Split misplaced clause: extracted '{clause_num}) {clause_title[:50]}...' from {clause_key}")
                            else:
                                clauses[new_clause_key]["content"] = after_content + clauses[new_clause_key].get("content", [])
                                clauses[new_clause_key]["title"] = f"{clause_num}) {clause_title}"
                                logging.info(f"Split misplaced clause: merged into existing {new_clause_key}")
                            # Re-sort and continue - need to process new clause's content for further splits
                            return self._split_misplaced_clause_headers(clauses)
                        break
        
        return clauses
    
    def _organize_addendum_clauses(self, addendum_content: List[str]) -> tuple[Dict[str, Dict[str, Any]], List[str]]:
        """Organize Addendum content into individual clauses using various delimiter patterns"""
        import re
        
        clauses = {}
        non_clause_content = []  # Content that doesn't belong to any numbered clause
        current_clause_number = None
        current_clause_title = None
        current_clause_content = []
        clause_counter = 1  # For sequential numbering
        
        for line in addendum_content:
            line_stripped = line.strip()
            if not line_stripped:
                if current_clause_number:
                    current_clause_content.append(line)
                else:
                    non_clause_content.append(line)
                continue
            
            # Check if this line is an addendum clause header (section delimiter)
            is_clause_header = False
            
            # Pattern 1: "ADDENDUM – description" (em dash or hyphen)
            addendum_dash_match = re.match(r'^ADDENDUM\s*[–-]\s*(.+)$', line_stripped, re.IGNORECASE)
            if addendum_dash_match:
                clause_title = addendum_dash_match.group(1).strip()
                if len(clause_title) > 0 and len(clause_title) <= 100:
                    is_clause_header = True
                    # Save the previous clause
                    if current_clause_number and current_clause_title:
                        clause_key = f"addendum_clause_{current_clause_number}"
                        logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                        clauses[clause_key] = {
                            "title": f"ADDENDUM - {current_clause_title}",
                            "content": current_clause_content
                        }
                    # Update to new clause
                    current_clause_number = clause_counter
                    clause_counter += 1
                    current_clause_title = clause_title
                    current_clause_content = []
                    logging.info(f"Found ADDENDUM clause {current_clause_number}: {clause_title[:50]}...")
            
            # Pattern 2: "02 – Account/Charterer:" or "04 - Hire:" (numbered with dash)
            elif re.match(r'^\d{1,2}\s*[–-]\s*.+', line_stripped):
                numbered_dash_match = re.match(r'^(\d{1,2})\s*[–-]\s*(.+)$', line_stripped)
                if numbered_dash_match:
                    clause_num = numbered_dash_match.group(1)
                    clause_title = numbered_dash_match.group(2).strip().rstrip(':')
                    if len(clause_title) > 0 and len(clause_title) <= 100:
                        is_clause_header = True
                        # Save the previous clause
                        if current_clause_number and current_clause_title:
                            clause_key = f"addendum_clause_{current_clause_number}"
                            logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                            clauses[clause_key] = {
                                "title": f"{current_clause_number}) {current_clause_title}",
                                "content": current_clause_content
                            }
                        # Update to new clause
                        current_clause_number = int(clause_num)
                        current_clause_title = clause_title
                        current_clause_content = []
                        logging.info(f"Found numbered addendum clause {current_clause_number}: {clause_title[:50]}...")
            
            # Pattern 3: "Cl. 11- Sub d) - Reinsert initial printed wording"
            elif re.match(r'^Cl\.\s*\d+[-\s]*Sub\s*[a-z]\)', line_stripped, re.IGNORECASE):
                clause_match = re.match(r'^(Cl\.\s*\d+[-\s]*Sub\s*[a-z]\))\s*[-\s]*(.+)$', line_stripped, re.IGNORECASE)
                if clause_match:
                    clause_prefix = clause_match.group(1)
                    clause_title = clause_match.group(2).strip()
                    if len(clause_title) > 0 and len(clause_title) <= 150:
                        is_clause_header = True
                        # Save the previous clause
                        if current_clause_number and current_clause_title:
                            clause_key = f"addendum_clause_{current_clause_number}"
                            logging.info(f"Saving previous {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                            clauses[clause_key] = {
                                "title": f"{current_clause_number}) {current_clause_title}",
                                "content": current_clause_content
                            }
                        # Update to new clause
                        current_clause_number = clause_counter
                        clause_counter += 1
                        current_clause_title = f"{clause_prefix} - {clause_title}"
                        current_clause_content = []
                        logging.info(f"Found clause-sub addendum clause {current_clause_number}: {clause_title[:50]}...")
            
            # If it's a clause header, don't add it to content (it's already captured in the title)
            if is_clause_header:
                continue
            
            # Add content to current clause or non-clause content
            if current_clause_number:
                current_clause_content.append(line_stripped)
            else:
                # Content before first clause (e.g., header info)
                non_clause_content.append(line_stripped)
        
        # Don't forget the last clause
        if current_clause_number and current_clause_title:
            clause_key = f"addendum_clause_{current_clause_number}"
            logging.info(f"Saving final {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
            clauses[clause_key] = {
                "title": f"{current_clause_number}) {current_clause_title}",
                "content": current_clause_content
            }
        
        logging.info(f"Organized Addendum into {len(clauses)} individual clauses with {len(non_clause_content)} non-clause lines")
        return clauses, non_clause_content
    
    def _organize_fixture_recap_clauses(self, fixture_recap_content: List[str]) -> tuple[Dict[str, Dict[str, Any]], List[str]]:
        """Organize Fixture Recap content into individual clauses using 'Fixture recap - X' delimiters"""
        import re
        
        clauses = {}
        non_clause_content = []  # Content that doesn't belong to any numbered clause
        current_clause_number = None
        current_clause_title = None
        current_clause_content = []
        clause_counter = 1  # For non-numbered clauses
        
        for line in fixture_recap_content:
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            # Check for two patterns:
            # Pattern 1: "Fixture recap - 2 account" (numbered)
            # Pattern 2: "Fixture recap – delivery" (non-numbered)
            # Handle both hyphen (-) and en dash (–) characters
            
            # Try numbered pattern first (case insensitive)
            numbered_match = re.match(r'^Fixture recap\s*[-–]\s*(\d+)[\.\s]*(.+)$', line_stripped, re.IGNORECASE)
            # Try non-numbered pattern (case insensitive)
            unnumbered_match = re.match(r'^Fixture recap\s*[-–]\s*([^0-9].+)$', line_stripped, re.IGNORECASE)
            # Hammarland, Heinz, etc.: Standalone "Trading Exclusions" or "Contract terms" without prefix
            standalone_fixture_match = (
                line_stripped.lower() in ['trading exclusions', 'trading exclusions:', 'contract terms', 'contract terms:'] or
                (line_stripped.lower().startswith('trading exclusions') and len(line_stripped) < 25) or
                (line_stripped.lower().startswith('contract terms') and len(line_stripped) < 25)
            )
            
            if numbered_match:
                # Save the previous clause if we have one
                if current_clause_number and current_clause_title:
                    clause_key = f"fixture_clause_{current_clause_number}"
                    if clause_key not in clauses:
                        clauses[clause_key] = {
                            "title": f"{current_clause_number}) {current_clause_title}",
                            "content": current_clause_content
                        }
                        logging.info(f"Saving fixture recap clause {current_clause_number}: {current_clause_title}")
                
                # Start new numbered clause
                current_clause_number = numbered_match.group(1)
                current_clause_title = numbered_match.group(2).strip().rstrip(':')  # Remove trailing colon
                current_clause_content = []
                logging.info(f"Found numbered fixture recap clause {current_clause_number}: {current_clause_title}")
                
            elif unnumbered_match:
                # Save the previous clause if we have one
                if current_clause_number and current_clause_title:
                    clause_key = f"fixture_clause_{current_clause_number}"
                    if clause_key not in clauses:
                        clauses[clause_key] = {
                            "title": f"{current_clause_number}) {current_clause_title}" if str(current_clause_number).isdigit() else current_clause_title,
                            "content": current_clause_content
                        }
                        logging.info(f"Saving fixture recap clause {current_clause_number}: {current_clause_title}")
                
                # Start new non-numbered clause (assign sequential number)
                current_clause_number = str(clause_counter)
                current_clause_title = unnumbered_match.group(1).strip().rstrip(':')  # Remove trailing colon
                current_clause_content = []
                clause_counter += 1
                logging.info(f"Found non-numbered fixture recap clause '{current_clause_title}' -> assigned number {current_clause_number}")
                
            elif standalone_fixture_match:
                # Save the previous clause if we have one
                if current_clause_number and current_clause_title:
                    clause_key = f"fixture_clause_{current_clause_number}"
                    if clause_key not in clauses:
                        clauses[clause_key] = {
                            "title": f"{current_clause_number}) {current_clause_title}" if str(current_clause_number).isdigit() else current_clause_title,
                            "content": current_clause_content
                        }
                        logging.info(f"Saving fixture recap clause {current_clause_number}: {current_clause_title}")
                
                # Start new clause from standalone heading (Trading Exclusions, Contract terms)
                current_clause_number = str(clause_counter)
                current_clause_title = line_stripped.rstrip(':').strip()
                current_clause_content = []
                clause_counter += 1
                logging.info(f"Found standalone fixture recap clause '{current_clause_title}' -> assigned number {current_clause_number}")
                
            else:
                # Add content to current clause or non-clause content
                if current_clause_number is not None:
                    current_clause_content.append(line_stripped)
                else:
                    # Content before first clause (e.g., header info)
                    non_clause_content.append(line_stripped)
        
        # Don't forget the last clause
        if current_clause_number and current_clause_title:
            clause_key = f"fixture_clause_{current_clause_number}"
            if clause_key not in clauses:
                clauses[clause_key] = {
                    "title": f"{current_clause_number}) {current_clause_title}" if str(current_clause_number).isdigit() else current_clause_title,
                    "content": current_clause_content
                }
                logging.info(f"Final fixture recap clause saved: {current_clause_number}")
        
        return clauses, non_clause_content
    
    def _extract_from_pdf_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using PyPDF2"""
        try:
            import PyPDF2
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- PAGE {page_num + 1} ---\n{text}")
            
            full_text = '\n\n'.join(text_content)
            
            return {
                'full_text': full_text,
                'pages': text_content,
                'page_count': len(text_content),
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }
        except Exception as e:
            logging.error(f"Error extracting from PDF {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _extract_from_pdf_pdfminer(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files using pdfminer"""
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(file_path)
            
            return {
                'full_text': text,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        except Exception as e:
            logging.error(f"Error extracting from PDF {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _extract_from_rtf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from RTF files"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            
            text = rtf_to_text(rtf_content)
            
            return {
                'full_text': text,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        except Exception as e:
            logging.error(f"Error extracting from RTF {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT files"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            return {
                'full_text': text,
                'line_count': len(text.split('\n')),
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        except Exception as e:
            logging.error(f"Error extracting from TXT {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _extract_from_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract text from CSV files"""
        try:
            import csv
            rows = []
            
            with open(file_path, 'r', encoding='utf-8') as file:
                # Try to detect delimiter
                sample = file.read(1024)
                file.seek(0)
                
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(file, delimiter=delimiter)
                for row in reader:
                    rows.append(row)
            
            # Convert to text
            text_rows = []
            for row in rows:
                text_rows.append(' | '.join(row))
            
            full_text = '\n'.join(text_rows)
            
            return {
                'full_text': full_text,
                'rows': rows,
                'row_count': len(rows),
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }
        except Exception as e:
            logging.error(f"Error extracting from CSV {file_path}: {str(e)}")
            return {'full_text': '', 'error': str(e)}
    
    def _extract_document_name(self, text: str, filename: str) -> str:
        """Dynamically extract document/vessel name from text or filename"""
        # First, try to find name patterns in the text with context awareness
        lines = text.split('\n')
        
        # Look for vessel details section first for better context
        vessel_details_start = -1
        vessel_details_end = -1
        
        for i, line in enumerate(lines[:200]):  # Check first 200 lines
            line_lower = line.strip().lower()
            # Find vessel details section
            if re.match(r'^\s*1[\.\)]\s*vessel\s+details?', line_lower) or \
               'vessel details' in line_lower or 'vessel specifications' in line_lower:
                vessel_details_start = i
                logging.info(f"Found vessel details section at line {i}")
                break
        
        # Define search range - prefer vessel details section if found
        if vessel_details_start >= 0:
            # Look within vessel details section (next 50 lines)
            search_lines = lines[vessel_details_start:vessel_details_start + 50]
            search_context = "vessel_details"
        else:
            # Fallback to first 100 lines
            search_lines = lines[:100]
            search_context = "general"
        
        # Look for common vessel name patterns (more specific first)
        name_patterns = [
            r'vessel\s+specifications\s*-?\s*vessel\s+name[\s:]+([^\n\r]+)',  # "Vessel specifications -Vessel Name: NAME"
            r'vessel\s+name[\s:]+([^\n\r]+)',        # "Vessel name: NAME" or "Vessel Name: NAME"
            r'ship\s+name[\s:]+([^\n\r]+)',          # "Ship name: NAME"
            r'^MV\s+([A-Z][^\n\r]+)',                # "MV VESSEL NAME" at start of line
            r'^\s*([MV\s]+[A-Z][A-Z\s]+)$',          # Line with just "MV VESSEL NAME"
            r'name[\s:]+([MV\s]+[^\n\r]+)',          # "Name: MV VESSEL"
        ]
        
        logging.info(f"Searching for vessel name in {search_context} context ({len(search_lines)} lines)")
        
        # First, try line-by-line pattern matching
        for line in search_lines:
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Skip very short lines
            if len(line_stripped) < 3:
                continue
            
            for pattern in name_patterns:
                match = re.search(pattern, line_lower)
                if match:
                    name = match.group(1).strip()
                    # Clean up the name but preserve MV prefix
                    name = re.sub(r'[^\w\s-]', ' ', name)
                    name = re.sub(r'\s+', ' ', name).strip()
                    
                    # Validate the name
                    if len(name) > 3 and len(name) < 100:
                        # If it doesn't start with MV but looks like a vessel name, add MV
                        if not name.upper().startswith('MV') and len(name.split()) > 1:
                            name = f"MV {name}"
                        logging.info(f"Found vessel name via pattern matching: '{name}'")
                        return name.upper()
        
        # Second, look for "Vessel Name:" followed by name on next line
        for i, line in enumerate(search_lines[:-1]):  # Don't check last line since we need next line
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            if 'vessel name' in line_lower and ':' in line_lower:
                next_line = search_lines[i + 1].strip()
                if len(next_line) > 3 and len(next_line) < 100:
                    # Clean up the name
                    name = re.sub(r'[^\w\s-]', ' ', next_line)
                    name = re.sub(r'\s+', ' ', name).strip()
                    
                    # Add MV prefix if not present
                    if not name.upper().startswith('MV') and len(name.split()) > 1:
                        name = f"MV {name}"
                    logging.info(f"Found vessel name on next line: '{name}'")
                    return name.upper()
        
        # Second pass: look for standalone vessel names that might not have explicit labels
        for line in lines:
            line_stripped = line.strip()
            if len(line_stripped) < 5:
                continue
                
            # Check if line looks like a vessel name (starts with MV and has 2+ words)
            if re.match(r'^MV\s+[A-Z][A-Z\s]+$', line_stripped.upper()):
                name = re.sub(r'[^\w\s-]', ' ', line_stripped)
                name = re.sub(r'\s+', ' ', name).strip()
                if len(name.split()) >= 2:
                    return name.upper()
        
        # Fallback to filename
        base_name = Path(filename).stem
        # Clean filename - if it looks like it contains a vessel name, extract it
        if 'ARA' in base_name.upper() or 'MV' in base_name.upper():
            # Try to extract vessel name from filename
            clean_name = re.sub(r'[^\w\s-]', ' ', base_name)
            clean_name = re.sub(r'\s+', ' ', clean_name).strip()
            # Remove common suffixes
            clean_name = re.sub(r'(?i)(done|final|draft|copy|\d+)$', '', clean_name).strip()
            if clean_name:
                if not clean_name.upper().startswith('MV') and len(clean_name.split()) > 1:
                    clean_name = f"MV {clean_name}"
                return clean_name.upper()
        
        # Final fallback
        name = re.sub(r'[^\w\s-]', ' ', base_name)
        name = re.sub(r'\s+', ' ', name).strip()
        return name.upper()
    
    def _organize_into_chapters(self, extraction_result: Dict[str, Any], doc_name: str) -> Dict[str, Any]:
        """Organize extracted content using chapter titles as delimiters"""
        # Initialize chapter structure
        chapters = {
            "1_vessel_details": {
                "title": "Vessel Details",
                "content": [],
                "tables": []
            },
            "2_contract_details": {
                "title": "Contract Details", 
                "content": [],
                "tables": []
            },
            "3_delivery_details": {
                "title": "Delivery Figures",
                "content": [],
                "tables": []
            },
            "4_speed_and_consumption": {
                "title": "Speed and Consumption",
                "content": [],
                "tables": []
            },
            "5_lifting_equipment": {
                "title": "Lifting Equipment",
                "content": [],
                "tables": []
            },
            "6_hseq_documentation": {
                "title": "HSEQ Documentation",
                "content": [],
                "tables": []
            },
            "7_vessel_communication_details": {
                "title": "Vessel Communication Details",
                "content": [],
                "tables": []
            },
            "7_vessel_and_owners_details": {
                "title": "Vessel and Owners details",
                "content": [],
                "tables": []
            }
        }
        
        # Get full text and split into lines
        full_text = extraction_result.get('full_text', '')
        lines = full_text.split('\n')
        
        # Define chapter title patterns to look for
        chapter_patterns = [
            (r'^1[\.\)]\s*Vessel Details?', "1_vessel_details"),
            (r'^2[\.\)]\s*Contract Details?', "2_contract_details"),
            # Baltic Iron: Contract Details as Heading 2 without number prefix, or "3) Contract Details"
            (r'^Contract Details?\s*$', "2_contract_details"),
            (r'^3\)\s*Contract Details?', "2_contract_details"),
            # Updated to catch "Delivery Details", "Delivery Figures", or just "Delivery"
            (r'^3[\.\)]\s*Delivery(?:\s+(?:Details?|Figures?))?', "3_delivery_details"),
            # Use period only for 4-6 to avoid matching clause numbers (e.g. "4) Speed" in charter party)
            (r'^4\.\s*Speed and Consumption', "4_speed_and_consumption"),
            (r'^5\.\s*Lifting Equipment', "5_lifting_equipment"),
            # Specific pattern for "5. Vessels communication details" in new documents
            (r'^5\.\s*Vessels?\s+[Cc]ommunication\s+[Dd]etails?', "7_vessel_communication_details"),
            (r'^6\.\s*HSEQ Documentation', "6_hseq_documentation"),
            # Vessel communication details - can be numbered or not
            (r'^[7-9]?[\.\)]?\s*Vessels?\s+[Cc]ommunication\s+[Dd]etails?', "7_vessel_communication_details"),
            (r'^Vessels?\s+[Cc]ommunication\s+[Dd]etails?', "7_vessel_communication_details"),
            # Vessel and Owners details - newer variant (different key for Langdock)
            (r'^[5-9][\.\)]\s*Vessels?\s+[Aa]nd\s+[Oo]wners?\s+[Dd]etails?', "7_vessel_and_owners_details"),
            (r'^Vessels?\s+[Aa]nd\s+[Oo]wners?\s+[Dd]etails?', "7_vessel_and_owners_details")
        ]
        
        # Find chapter positions
        chapter_positions = []
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            for pattern, chapter_key in chapter_patterns:
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    # Special handling: If we're extracting chapter 3 (Delivery) but we're still
                    # within what looks like charter party content (contains clause patterns),
                    # don't treat it as a chapter delimiter - it's likely a clause within the charter party
                    if chapter_key == "3_delivery_details":
                        # Check if we're still within chapter 2 content (contract details)
                        # by looking backwards for chapter 2 marker and forward for chapter 4
                        should_skip = False
                        
                        # Strategy: If we find "Charter Party" text before this "3) Delivery" line,
                        # then this "3) Delivery" is a clause within the charter party, not a chapter
                        for j in range(max(0, i - 200), i):  # Look back up to 200 lines
                            if j < len(lines):
                                check_line = lines[j].strip().lower()
                                # If we find "charter party" text, this is within charter party section
                                if 'charter party' in check_line and len(check_line) < 100:
                                    should_skip = True
                                    logging.info(f"SKIPPING chapter 3 delimiter '{line_stripped}' at line {i} - found 'Charter Party' at line {j}, treating as clause")
                                    break
                                # If we find "2) Contract Details", check if Charter Party is mentioned after it
                                if re.match(r'^2[\.\)]\s*contract details?', check_line):
                                    # Check if "Charter Party" appears between chapter 2 and here
                                    for k in range(j + 1, i):
                                        if 'charter party' in lines[k].strip().lower():
                                            should_skip = True
                                            logging.info(f"SKIPPING chapter 3 delimiter '{line_stripped}' at line {i} - Charter Party section found between Contract Details and here")
                                            break
                                    if should_skip:
                                        break
                        
                        if not should_skip:
                            # Not part of charter party - treat as chapter delimiter
                            chapter_positions.append((i, chapter_key, line_stripped))
                            logging.info(f"Found chapter delimiter: '{line_stripped}' at line {i}")
                        break
                    else:
                        # Not chapter 3, proceed normally
                        chapter_positions.append((i, chapter_key, line_stripped))
                        logging.info(f"Found chapter delimiter: '{line_stripped}' at line {i}")
                        break
        
        # Sort by position
        chapter_positions.sort(key=lambda x: x[0])
        
        # Extract content between chapter delimiters
        for i, (start_pos, chapter_key, chapter_title) in enumerate(chapter_positions):
            # Determine end position
            if i + 1 < len(chapter_positions):
                end_pos = chapter_positions[i + 1][0]
            else:
                end_pos = len(lines)
            
            # Extract content for this chapter (skip the title line itself)
            chapter_content = []
            for line_idx in range(start_pos + 1, end_pos):
                if line_idx < len(lines):
                    line = lines[line_idx].strip()
                    if line:  # Only add non-empty lines
                        chapter_content.append(line)
            
            # For chapter 7 variants, use the actual title from the source document
            if chapter_key in ("7_vessel_communication_details", "7_vessel_and_owners_details"):
                # Strip leading number like "7. " or "7) " to get the display title
                display_title = re.sub(r'^\d+[\.\)]\s*', '', chapter_title).strip()
                if display_title:
                    chapters[chapter_key]["title"] = display_title

            # Special handling for Contract Details - organize into sub-chapters
            if chapter_key == "2_contract_details" and chapter_content:
                # For short documents, use simplified structure; for full documents, use sub-chapters
                document_type = self._detect_document_type(chapter_content)
                if document_type == "basic_vessel_info":
                    # Use simplified structure for basic vessel info documents
                    chapters[chapter_key] = self._organize_basic_vessel_info(chapter_content)
                    logging.info(f"Organized basic vessel info with {len(chapter_content)} content lines (no sub-chapters)")
                else:
                    # Use full sub-chapter structure for complex charter party documents
                    chapters[chapter_key] = self._organize_contract_details_subchapters(chapter_content)
                    logging.info(f"Organized Contract Details into sub-chapters with {len(chapter_content)} total lines")
            else:
                # Regular chapter handling
                if chapter_content:
                    chapters[chapter_key]["content"] = chapter_content
                    logging.info(f"Assigned {len(chapter_content)} lines to {chapter_key}")
        
        # Handle tables - ONLY assign tables to speed and consumption chapter
        # All other "tables" are likely clause headers or text that shouldn't be treated as tables
        tables = extraction_result.get('tables', [])
        for table in tables:
            if not table:
                continue
                
            # Check table headers to determine if it's a legitimate speed/consumption table
            table_text = ' '.join([' '.join(row) for row in table if row]).lower()
            
            # ONLY Speed and Consumption tables are real tables
            if any(keyword in table_text for keyword in [
                'speed', 'consumption', 'fuel', 'knots', 'mt/day', 'hfo', 'engine load'
            ]):
                chapters["4_speed_and_consumption"]["tables"].append(table)
                logging.info("Assigned table to Speed and Consumption chapter")
            else:
                # All other "tables" are ignored - they're likely clause headers or text
                logging.info(f"Ignoring non-speed/consumption table: {table_text[:100]}...")
        
        # Fallback: If no vessel details chapter found, try to extract vessel details from entire document
        if "1_vessel_details" not in chapters or not chapters["1_vessel_details"].get("content"):
            logging.info("No dedicated Vessel Details chapter found, attempting to extract vessel details from full document")
            chapters["1_vessel_details"] = self._extract_vessel_details_fallback(lines)
            logging.info("Created vessel details chapter from document content")
        
        # Extract vessel communication details from other chapters if they were misplaced
        self._extract_misplaced_communication_details(chapters)
        
        # Clean up empty chapters
        cleaned_chapters = {}
        for chapter_key, chapter_data in chapters.items():
            # Check for content, tables, or sub_chapters
            has_content = (chapter_data.get("content") or 
                          chapter_data.get("tables") or 
                          chapter_data.get("sub_chapters"))
            if has_content:
                cleaned_chapters[chapter_key] = chapter_data
        
        return cleaned_chapters
    
    def _extract_misplaced_communication_details(self, chapters: Dict[str, Any]) -> None:
        """Extract vessel communication details that may have been misplaced in other chapters"""
        communication_content = []
        
        # Look through all chapters for communication details
        for chapter_key, chapter_data in chapters.items():
            if chapter_key in ("7_vessel_communication_details", "7_vessel_and_owners_details"):
                continue  # Skip the target chapters
                
            content = chapter_data.get("content", [])
            if not content:
                continue
                
            # Find lines that indicate communication details section
            new_content = []
            found_communication_section = False
            
            for i, line in enumerate(content):
                line_lower = line.strip().lower()
                
                # Check if this line starts a communication details section (either variant)
                if (re.match(r'^\d+\.?\s*vessels?\s+communication\s+details?', line_lower) or
                    re.match(r'^\d+\.?\s*vessels?\s+and\s+owners?\s+details?', line_lower)):
                    found_communication_section = True
                    logging.info(f"Found misplaced communication details in {chapter_key}: '{line.strip()}'")
                    # Add the header line to communication content
                    communication_content.append(line)
                    continue
                
                if found_communication_section:
                    # Check if we've reached the end of communication section
                    # More comprehensive list of communication-related patterns
                    comm_patterns = [
                        'email', '@', 'phone', 'mobile', 'cell', 'vsat', 'inmarsat', 'iridium', 'master', 'captain',
                        'ph:', 'tel:', 'fax:', 'e:', 'sat phone', 'ip phone', 'office', 'address', 'www.', 'http',
                        'imo', 'mv ', 'm/v', 'vessel', 'ship', '+31', '+47', '+44', '+1', '+86', '+33', '+49',  # Common country codes
                        'netherlands', 'norway', 'uk', 'usa', 'china', 'france', 'germany',  # Countries
                        '.com', '.net', '.org', '.no', '.nl', '.uk', '.de'  # Domain extensions
                    ]
                    
                    # Also check if line looks like address components (postal codes, cities)
                    address_pattern = re.match(r'^\d{4,6}\s+[A-Z]{2,3}\s+\w+', line.strip())  # e.g., "8321 DX Urk"
                    vessel_name_pattern = re.match(r'^(mv|m/v)\s+\w+', line_lower)  # e.g., "MV Atlantic"
                    street_address_pattern = re.match(r'^[a-zA-Z][a-zA-Z\s\d]+$', line.strip())  # e.g., "Sluisweg 11"
                    
                    is_communication_line = (
                        any(keyword in line_lower for keyword in comm_patterns) or
                        address_pattern or
                        vessel_name_pattern or
                        street_address_pattern or
                        line.strip().isupper() or  # All caps lines often contain addresses/names
                        re.match(r'^[a-zA-Z\s]+$', line.strip())  # Lines with only letters and spaces (country names, etc.)
                    )
                    
                    if line.strip() and not is_communication_line:
                        # This line doesn't look like communication details, stop collecting
                        found_communication_section = False
                        new_content.append(line)
                    else:
                        # This line is part of communication details
                        communication_content.append(line)
                else:
                    # Keep this line in the original chapter
                    new_content.append(line)
            
            # Update the chapter content if we extracted communication details
            if len(new_content) != len(content):
                chapter_data["content"] = new_content
                logging.info(f"Removed {len(content) - len(new_content)} communication lines from {chapter_key}")
        
        # If we found communication details, assign them to the dedicated chapter
        if communication_content:
            # Determine which key based on header: "Vessel and Owners" vs "Vessel Communication"
            first_line = communication_content[0].strip() if communication_content else ""
            if re.search(r'vessels?\s+and\s+owners?\s+details?', first_line, re.IGNORECASE):
                target_key = "7_vessel_and_owners_details"
            else:
                target_key = "7_vessel_communication_details"
            chapters[target_key]["content"] = communication_content
            # Use the actual title from the header line if it's the first line
            if re.match(r'^\d+\.?\s*vessels?\s+(?:communication|and\s+owners?)\s+details?', first_line, re.IGNORECASE):
                display_title = re.sub(r'^\d+[\.\)]?\s*', '', first_line).strip()
                if display_title:
                    chapters[target_key]["title"] = display_title
            logging.info(f"Moved {len(communication_content)} lines to {target_key} chapter")
    
    def _separate_rider_clause_content(self, charter_party_content: List[str]) -> Tuple[List[str], List[str]]:
        """Separate rider clause content from main charter party content"""
        main_content = []
        rider_content = []
        in_rider_section = False
        
        logging.info(f"Checking {len(charter_party_content)} charter party lines for rider clause section")
        
        for i, line in enumerate(charter_party_content):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Check for rider clause section header
            # Look for patterns like "2.3.1 RIDER CLAUSE OF CHARTER PARTY" or "RIDER CLAUSE OF CHARTER PARTY"
            if ('rider clause' in line_lower and 'charter party' in line_lower) or \
               re.match(r'^\d+\.\d+\.\d+.*rider.*clause.*charter.*party', line_lower):
                in_rider_section = True
                rider_content.append(line)
                logging.info(f"Found rider clause section: '{line_stripped}'")
                continue
            
            # Debug: Check if line contains "rider" to see what we're missing
            if 'rider' in line_lower:
                logging.info(f"DEBUG: Found line with 'rider': '{line_stripped[:60]}...'")
            
            # Also check for patterns that might indicate the section start
            if '2.3.1' in line_stripped and 'rider' in line_lower:
                logging.info(f"DEBUG: Found 2.3.1 with rider: '{line_stripped[:60]}...'")
                in_rider_section = True
                rider_content.append(line)
                continue
            
            # If we're in rider section, add content to rider_content
            if in_rider_section:
                rider_content.append(line)
            else:
                main_content.append(line)
        
        return main_content, rider_content
    
    def _organize_rider_clauses(self, rider_content: List[str]) -> Tuple[Dict[str, Any], List[str]]:
        """Organize rider clause content into individual clauses using numbered delimiters"""
        import re
        
        clauses = {}
        non_clause_content = []
        current_clause_number = None
        current_clause_title = None
        current_clause_content = []
        
        # Define clause headers for better detection
        clause_headers = [
            "BLACK LIST", "EXTRA INSURANCE", "STOWAGE PLANS", "HOLDS ON DELIVERY", "BUNKERING",
            "P&I INSURANCE", "EQUIPMENT", "BILL OF LADING", "WEATHER ROUTING", "DECK CARGO",
            "PLANS", "SERVICES", "DOUBLE BANKING", "PAD-EYES", "REEFER CONTAINER", "SUEZ CANAL",
            "PANAMA CANAL", "WATCHMEN", "TAX RETURN", "AGENCY", "ARREST", "SEIZURE", "CANCELLATION",
            "LANGUAGE", "CONFIDENTIAL", "GEAR", "CERTIFICATES", "DAMAGES", "LIABILITY", "POLLUTION",
            "SAFETY", "MANAGEMENT", "PERFORMANCE", "ITF", "CLAUSE"
        ]
        
        logging.info(f"Processing {len(rider_content)} lines for rider clause organization")
        if rider_content:
            logging.info(f"First few rider content lines: {rider_content[:3]}")
        
        for i, line in enumerate(rider_content):
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            is_clause_header = False
            
            # Skip the main rider clause header (e.g., "2.3 RIDER CLAUSE OF CHARTER PARTY :")
            if ('rider clause' in line_stripped.lower() and 'charter party' in line_stripped.lower()) or \
               re.match(r'^\d+\.\d+\.\d+.*rider.*clause.*charter.*party', line_stripped.lower()):
                logging.info(f"Skipping rider clause header line: '{line_stripped}'")
                continue
            
            # Check for numbered patterns like "1." or "1)" at start of line
            numbered_match = re.match(r'^(\d+)[\.\)]\s*(.+)$', line_stripped)
            if numbered_match:
                clause_num = numbered_match.group(1)
                clause_title = numbered_match.group(2).strip()
                
                # For rider clauses, be more permissive since they're in a dedicated section
                if len(clause_title) <= 85 and len(clause_title) > 1:
                    is_clause_header = True
                    # Save the previous clause BEFORE updating variables
                    if current_clause_number and current_clause_title:
                        clause_key = f"rider_clause_{current_clause_number}"
                        logging.info(f"Saving previous rider {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
                        if clause_key not in clauses:
                            clauses[clause_key] = {
                                "title": f"{current_clause_number}) {current_clause_title}",
                                "content": current_clause_content
                            }
                            logging.info(f"Successfully saved rider {clause_key}")
                        else:
                            logging.warning(f"Duplicate rider clause number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title']}")
                    
                    # Update for new clause
                    current_clause_number = clause_num
                    current_clause_title = clause_title
                    current_clause_content = []  # Reset for new clause
                    logging.info(f"Found rider clause {current_clause_number}: {clause_title[:50]}...")
            
            # If this line is not a clause header, add it to current clause content or non-clause content
            if not is_clause_header:
                if current_clause_number is not None:
                    current_clause_content.append(line_stripped)
                else:
                    # This content doesn't belong to any numbered clause (e.g., header info)
                    non_clause_content.append(line_stripped)
        
        # Don't forget to save the last clause
        if current_clause_number and current_clause_title:
            clause_key = f"rider_clause_{current_clause_number}"
            logging.info(f"Final rider clause - attempting to save {clause_key} with title '{current_clause_title}' and {len(current_clause_content)} content lines")
            
            # Check if clause already exists - if so, keep the first one
            if clause_key not in clauses:
                clauses[clause_key] = {
                    "title": f"{current_clause_number}) {current_clause_title}",
                    "content": current_clause_content
                }
                logging.info(f"Final rider clause - successfully saved {clause_key}")
            else:
                logging.warning(f"Final rider clause duplicate number {current_clause_number} detected. Keeping first occurrence: {clauses[clause_key]['title'][:50]}...")
        
        return clauses, non_clause_content
    
    def _detect_document_type(self, contract_content: List[str]) -> str:
        """Detect if this is a full charter party document or a basic vessel info document"""
        content_text = ' '.join(contract_content).lower()
        
        # Count indicators of full charter party documents
        full_document_indicators = 0
        
        # Check for typical charter party sections
        if any(keyword in content_text for keyword in ['addendum', 'addenda']):
            full_document_indicators += 1
        if any(keyword in content_text for keyword in ['fixture recap', 'fixture', 'recap']):
            full_document_indicators += 1
        if len([line for line in contract_content if any(keyword in line.lower() for keyword in ['clause', 'charter party'])]) > 5:
            full_document_indicators += 1
        
        # Check content volume
        if len(contract_content) > 50:
            full_document_indicators += 1
        if len(content_text.split()) > 200:
            full_document_indicators += 1
            
        # Determine document type
        if full_document_indicators >= 2:
            return "full_charter_party"
        else:
            return "basic_vessel_info"

    def _organize_basic_vessel_info(self, contract_content: List[str]) -> Dict[str, Any]:
        """Organize basic vessel info documents with minimal contract details structure"""
        import re
        
        # Initialize simplified structure for basic vessel info documents
        contract_details = {
            "title": "Contract Details",
            "content": [],
            "tables": []
        }
        
        # For basic documents, put all content directly without any sub-chapter structure or clause detection
        for line in contract_content:
            line_stripped = line.strip()
            if line_stripped:
                contract_details["content"].append(line_stripped)
        
        logging.info(f"Organized basic vessel info with {len(contract_details['content'])} content lines (no clause processing)")
        
        return contract_details

    def _organize_contract_details_subchapters(self, contract_content: List[str]) -> Dict[str, Any]:
        """Organize Contract Details content into sub-chapters: Addendum, Fixture Recap, and Charter Party"""
        
        # Detect document type first
        document_type = self._detect_document_type(contract_content)
        logging.info(f"Detected document type: {document_type}")
        
        if document_type == "basic_vessel_info":
            # For basic vessel info documents, use simplified structure
            return self._organize_basic_vessel_info(contract_content)
        
        # Initialize sub-chapter structure for full charter party documents
        contract_details = {
            "title": "Contract Details",
            "sub_chapters": {
                "addendum": {
                    "title": "Addendum",
                    "content": []
                },
                "fixture_recap": {
                    "title": "Fixture Recap", 
                    "content": []
                },
                "charter_party": {
                    "title": "Charter Party",
                    "content": [],
                    "clauses": {}
                }
            },
            "tables": []
        }
        
        # Find sub-chapter positions - only use main section headers as delimiters
        subchapter_positions = []
        for i, line in enumerate(contract_content):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # More flexible pattern matching for sub-chapters
            # Check for numbered sub-chapters with flexible formatting: 2.1, 2.1., 2.2, 2.2., etc.
            import re
            subchapter_match = re.match(r'^2\.([123])\.?\s*(.+)$', line_stripped, re.IGNORECASE)
            # Also check for patterns like "3) Charter Party" but be very specific
            alt_subchapter_match = re.match(r'^([23])\)\s*(.+)$', line_stripped, re.IGNORECASE)
            # Only consider if it's likely a sub-chapter header (short and contains keywords)
            if alt_subchapter_match and len(line_stripped) < 50:
                alt_subchapter_match = alt_subchapter_match
            else:
                alt_subchapter_match = None
            
            if subchapter_match:
                section_num = subchapter_match.group(1)
                section_title = subchapter_match.group(2).strip()
                section_title_lower = section_title.lower()
                
                # Identify sub-chapter type by title content (more flexible)
                if any(keyword in section_title_lower for keyword in ['fixture recap', 'fixture', 'recap']):
                    subchapter_positions.append((i, "fixture_recap", line))
                    logging.info(f"Found Fixture Recap sub-chapter: '{line}' at position {i}")
                elif any(keyword in section_title_lower for keyword in ['charter party', 'charter', 'party']):
                    # Special case: Don't treat "RIDER CLAUSE OF CHARTER PARTY" as a separate sub-chapter
                    if 'rider' in section_title_lower:
                        logging.info(f"Skipping rider clause header as sub-chapter: '{line}' at position {i}")
                    else:
                        subchapter_positions.append((i, "charter_party", line))
                        logging.info(f"Found Charter Party sub-chapter: '{line}' at position {i}")
                elif any(keyword in section_title_lower for keyword in ['addendum', 'addenda']):
                    subchapter_positions.append((i, "addendum", line))
                    logging.info(f"Found Addendum sub-chapter: '{line}' at position {i}")
                else:
                    # Log unrecognized sub-chapter for debugging
                    logging.info(f"Found unrecognized 2.x sub-chapter: '{line}' at position {i}")
            
            elif alt_subchapter_match:
                section_num = alt_subchapter_match.group(1)
                section_title = alt_subchapter_match.group(2).strip()
                section_title_lower = section_title.lower()
                
                # Check for Charter Party pattern like "3) Charter Party"
                if any(keyword in section_title_lower for keyword in ['charter party', 'charter', 'party']):
                    subchapter_positions.append((i, "charter_party", line))
                    logging.info(f"Found alternative Charter Party sub-chapter: '{line}' at position {i}")
                elif any(keyword in section_title_lower for keyword in ['fixture recap', 'fixture', 'recap']):
                    subchapter_positions.append((i, "fixture_recap", line))
                    logging.info(f"Found alternative Fixture Recap sub-chapter: '{line}' at position {i}")
                elif any(keyword in section_title_lower for keyword in ['addendum', 'addenda']):
                    subchapter_positions.append((i, "addendum", line))
                    logging.info(f"Found alternative Addendum sub-chapter: '{line}' at position {i}")
                else:
                    logging.info(f"Found unrecognized alternative sub-chapter: '{line}' at position {i}")
            
            # Also check for standalone sub-chapter headers (without numbering) - MUCH more strict
            elif line_stripped and len(line_stripped) < 50:  # Likely a header
                # Only match if it's EXACTLY or VERY close to the sub-chapter names (not just contains keywords)
                if (line_lower == 'fixture recap' or 
                    line_lower == 'fixture recap:' or
                    line_lower == 'fixture recap.' or
                    # Only match simple "Fixture recap" headers, not "Fixture recap - X" item lines
                    (line_lower.startswith('fixture recap') and not ' - ' in line_lower and len(line_stripped) < 20)):
                    subchapter_positions.append((i, "fixture_recap", line))
                    logging.info(f"Found standalone Fixture Recap sub-chapter: '{line}' at position {i}")
                elif (line_lower == 'charter party' or 
                      line_lower == 'charter party:' or
                      line_lower == 'charter party.' or
                      line_lower.startswith('charter party -')):
                    subchapter_positions.append((i, "charter_party", line))
                    logging.info(f"Found standalone Charter Party sub-chapter: '{line}' at position {i}")
                elif (line_lower == 'addendum' or 
                      line_lower == 'addendums' or
                      line_lower == 'addendum:' or
                      line_lower == 'addendum.' or
                      line_lower.startswith('addendum -')):
                    subchapter_positions.append((i, "addendum", line))
                    logging.info(f"Found standalone Addendum sub-chapter: '{line}' at position {i}")
        
        # Sort by position
        subchapter_positions.sort(key=lambda x: x[0])
        
        # Debug: Log detected sub-chapter positions for FREDENSBORG
        logging.info(f"Found {len(subchapter_positions)} sub-chapter positions:")
        for pos, section, line in subchapter_positions[:10]:  # Show first 10
            logging.info(f"  Position {pos}: {section} - '{line[:60]}...'")
        if len(subchapter_positions) > 10:
            logging.info(f"  ... and {len(subchapter_positions) - 10} more positions")
        


        
        # If no sub-chapter delimiters found, categorize by content patterns
        if not subchapter_positions:
            logging.info("No explicit sub-chapter delimiters found, using content-based categorization")
            for line in contract_content:
                line_lower = line.lower()
                line_stripped = line.strip()
                
                # Check for numbered clauses first (highest priority - these should ALWAYS go to charter_party)
                import re
                numbered_clause_match = re.match(r'^\d+[\.\)]\s*[A-Z]', line_stripped)
                clause_dash_match = re.match(r'^Clause\s+\d+\s*[-–]', line_stripped, re.IGNORECASE)
                
                if numbered_clause_match or clause_dash_match:
                    # This is definitely a charter party clause
                    contract_details["sub_chapters"]["charter_party"]["content"].append(line)
                
                # Fixture Recap items (only basic recap content, not numbered clauses)
                elif any(pattern in line_lower for pattern in [
                    'fixture recap', 'fixture', 'recap'
                ]) and not numbered_clause_match and not clause_dash_match:
                    contract_details["sub_chapters"]["fixture_recap"]["content"].append(line)
                
                # Addendum items
                elif any(pattern in line_lower for pattern in [
                    'addendum', 'amendment', 'additional', 'supplement'
                ]):
                    contract_details["sub_chapters"]["addendum"]["content"].append(line)
                
                # Charter Party items (legal clauses, terms, etc.)
                elif any(pattern in line_lower for pattern in [
                    'charter party', 'clause', 'owners', 'charterers', 'vessel shall',
                    'liability', 'indemnify', 'charter', 'sub-clause', 'delivery', 'redelivery',
                    'hire', 'bunkers', 'speed', 'consumption', 'arbitration', 'bimco'
                ]):
                    contract_details["sub_chapters"]["charter_party"]["content"].append(line)
                
                # Default to Charter Party for unclassified content (this ensures we don't lose clauses)
                else:
                    contract_details["sub_chapters"]["charter_party"]["content"].append(line)
        
        else:
            # Extract content between sub-chapter delimiters
            for i, (start_pos, subchapter_key, subchapter_title) in enumerate(subchapter_positions):
                # Determine end position
                if i + 1 < len(subchapter_positions):
                    end_pos = subchapter_positions[i + 1][0]
                else:
                    end_pos = len(contract_content)
                
                # Extract content for this sub-chapter (skip the title line itself)
                subchapter_content = []
                for line_idx in range(start_pos + 1, end_pos):
                    if line_idx < len(contract_content):
                        line = contract_content[line_idx].strip()
                        if line:  # Only add non-empty lines
                            subchapter_content.append(line)
                
                if subchapter_content:
                    # If this sub-chapter type already has content, append to it rather than overwrite
                    if contract_details["sub_chapters"][subchapter_key]["content"]:
                        contract_details["sub_chapters"][subchapter_key]["content"].extend(subchapter_content)
                        logging.info(f"Appended {len(subchapter_content)} lines to existing {subchapter_key} sub-chapter")
                    else:
                        contract_details["sub_chapters"][subchapter_key]["content"] = subchapter_content
                    logging.info(f"Assigned {len(subchapter_content)} lines to {subchapter_key} sub-chapter")
        
        # Clean up empty sub-chapters
        cleaned_subchapters = {}
        for subchapter_key, subchapter_data in contract_details["sub_chapters"].items():
            if subchapter_data["content"]:
                cleaned_subchapters[subchapter_key] = subchapter_data
        
        contract_details["sub_chapters"] = cleaned_subchapters
        
        # NEW: Organize Charter Party content into individual clauses using the numbered delimiters
        if "charter_party" in contract_details["sub_chapters"]:
            charter_party_content = contract_details["sub_chapters"]["charter_party"]["content"]

            if charter_party_content:
                # First, check for and separate rider clause content
                main_charter_content, rider_clause_content = self._separate_rider_clause_content(charter_party_content)
                
                # Organize main charter party clauses
                clauses, non_clause_content = self._organize_charter_party_clauses(main_charter_content)
                contract_details["sub_chapters"]["charter_party"]["clauses"] = clauses
                
                # If rider clause content exists, organize it separately
                if rider_clause_content:
                    rider_clauses, rider_non_clause_content = self._organize_rider_clauses(rider_clause_content)
                    contract_details["sub_chapters"]["charter_party"]["rider_clause"] = {
                        "title": "Rider Clause of Charter Party",
                        "clauses": rider_clauses,
                        "content": rider_non_clause_content
                    }
                    logging.info(f"Organized Rider Clause into {len(rider_clauses)} individual clauses and preserved {len(rider_non_clause_content)} non-clause lines")
                
                # Keep only non-clause content in the main content array (header info, company details, etc.)
                contract_details["sub_chapters"]["charter_party"]["content"] = non_clause_content
                logging.info(f"Organized Charter Party into {len(clauses)} individual clauses and preserved {len(non_clause_content)} non-clause lines")
                
                # Post-processing: If clause 2 has empty content, it might have been incorrectly split
                # Check if we need to merge content from later clauses or if content was missed
                if "clause_2" in clauses and not clauses["clause_2"]["content"]:
                    logging.warning(f"Clause 2 '{clauses['clause_2']['title']}' has empty content - checking for missed content")
                    # Look for content that should belong to clause 2 by checking subsequent clauses
                    # This handles cases where clause 2 content was incorrectly separated
                    if "clause_3" in clauses and clauses["clause_3"]["content"]:
                        # Check if clause 3 title suggests it might be part of clause 2 (e.g., "Delivery" when clause 2 is about Duration/Trip Description)
                        clause_2_title_lower = clauses["clause_2"]["title"].lower()
                        clause_3_title_lower = clauses["clause_3"]["title"].lower()
                        # If clause 2 is about Duration/Trip Description and clause 3 is about Delivery,
                        # and clause 2 is empty, this might be a mis-detection
                        if ("duration" in clause_2_title_lower or "trip description" in clause_2_title_lower) and "delivery" in clause_3_title_lower:
                            logging.info(f"Clause 2 is empty but clause 3 contains delivery content - this may indicate a parsing issue")
                            # Don't auto-merge, but log for investigation

            else:
                logging.info("Charter party content is empty")
        else:
            logging.info("No charter_party sub-chapter found")
        
        # NEW: Organize Addendum content into individual clauses using various delimiter patterns
        if "addendum" in contract_details["sub_chapters"]:
            addendum_content = contract_details["sub_chapters"]["addendum"]["content"]
            if addendum_content:
                addendum_clauses, addendum_non_clause_content = self._organize_addendum_clauses(addendum_content)
                contract_details["sub_chapters"]["addendum"]["clauses"] = addendum_clauses
                # Keep the full addendum content AND add the organized clauses
                # (Unlike charter party, we preserve all content for backward compatibility)
                logging.info(f"Organized Addendum into {len(addendum_clauses)} individual clauses and preserved {len(addendum_non_clause_content)} non-clause lines")
        
        # NEW: Organize Fixture Recap content into individual clauses using the "Fixture recap - X" delimiters
        if "fixture_recap" in contract_details["sub_chapters"]:
            fixture_recap_content = contract_details["sub_chapters"]["fixture_recap"]["content"]
            if fixture_recap_content:
                fixture_clauses, fixture_non_clause_content = self._organize_fixture_recap_clauses(fixture_recap_content)
                contract_details["sub_chapters"]["fixture_recap"]["clauses"] = fixture_clauses
                # Keep the full fixture recap content AND add the organized clauses
                # (Unlike charter party, we preserve all content for backward compatibility)
                logging.info(f"Organized Fixture Recap into {len(fixture_clauses)} individual clauses and preserved {len(fixture_non_clause_content)} non-clause lines")
        
        # Ensure compatibility - add empty content list for backward compatibility
        contract_details["content"] = []
        
        return contract_details
    
    def _extract_vessel_details_fallback(self, lines: List[str]) -> Dict[str, Any]:
        """
        Fallback method to extract vessel details when no dedicated chapter exists.
        Scans the entire document for vessel specification patterns.
        """
        vessel_details_content = []
        
        # Patterns that indicate vessel specifications
        vessel_patterns = [
            r'vessel\s+name:?\s*(.+)',
            r'call\s+sign:?\s*(.+)',
            r'imo\s+number:?\s*(.+)',
            r'port\s+of\s+registry:?\s*(.+)',
            r'flag\s+state:?\s*(.+)',
            r'shipyard:?\s*(.+)',
            r'year.*delivery:?\s*(.+)',
            r'class:?\s*(.+)',
            r'dwat:?\s*(.+)',
            r'gross\s+tonnage:?\s*(.+)',
            r'net\s+tonnage:?\s*(.+)',
            r'length\s+over\s+all:?\s*(.+)',
            r'breadth:?\s*(.+)',
            r'draught:?\s*(.+)',
            r'depth:?\s*(.+)',
            r'main\s+engine:?\s*(.+)',
            r'service\s+speed:?\s*(.+)',
            r'fuel\s+oil:?\s*(.+)',
            r'ballast\s+water:?\s*(.+)',
            r'cranes:?\s*(.+)',
            r'holds?:?\s*(.+)',
            r'hatches?:?\s*(.+)',
            r'vessel\s+specifications',
            r'vessel\s+description',
            r'vessel\s+particulars'
        ]
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            # Check if line contains vessel specification patterns
            for pattern in vessel_patterns:
                if re.search(pattern, line_stripped, re.IGNORECASE):
                    vessel_details_content.append(line_stripped)
                    break
        
        # If no specific patterns found, add a placeholder
        if not vessel_details_content:
            vessel_details_content = ["[Vessel details not found in dedicated section]"]
            
        return {
            "title": "Vessel Details",
            "content": vessel_details_content,
            "tables": []
        }
    
    def _discover_files(self) -> List[Dict[str, str]]:
        """Dynamically discover all processable files in source directory"""
        files_to_process = []
        
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return files_to_process
        
        for filename in os.listdir(self.source_dir):
            # Skip temporary files
            if filename.startswith('~$') or filename.startswith('.'):
                continue
            
            file_path = os.path.join(self.source_dir, filename)
            if not os.path.isfile(file_path):
                continue
            
            # Get file extension
            ext = Path(filename).suffix.lower()
            
            if ext in self.supported_formats:
                files_to_process.append({
                    'filename': filename,
                    'path': file_path,
                    'extension': ext,
                    'extractor': self.supported_formats[ext]
                })
                logging.info(f"Found processable file: {filename} ({ext})")
            else:
                logging.warning(f"Unsupported file format: {filename} ({ext})")
        
        return files_to_process
    
    def process_files(self) -> Dict[str, Dict[str, Any]]:
        """Process all discovered files dynamically"""
        processed_data = {}
        
        files_to_process = self._discover_files()
        
        if not files_to_process:
            logging.warning("No processable files found")
            return processed_data
        
        logging.info(f"Processing {len(files_to_process)} files...")
        
        for file_info in files_to_process:
            filename = file_info['filename']
            file_path = file_info['path']
            extractor_func = file_info['extractor']
            
            logging.info(f"Processing {filename}...")
            
            try:
                # Extract text using appropriate method
                extraction_result = extractor_func(file_path)
                
                if not extraction_result.get('full_text'):
                    logging.warning(f"No text extracted from {filename}")
                    continue
                
                # Extract document name dynamically
                doc_name = self._extract_document_name(
                    extraction_result['full_text'], 
                    filename
                )
                
                # Organize content into chapters
                organized_content = self._organize_into_chapters(
                    extraction_result, doc_name
                )
                
                # Create document data structure - organized by chapters
                document_data = {
                    'document_name': doc_name,
                    'chapters': organized_content
                }
                
                # Clean up empty fields
                document_data = {k: v for k, v in document_data.items() if v}
                
                # When multiple files map to same doc_name (e.g. AI_MORGENSTOND_II.docx and
                # AI_MORGENSTOND_II_missingaddendum.docx), prefer the complete version.
                # Skip overwriting with "missing" variant if we already have addendum content.
                if doc_name in processed_data and 'missing' in filename.lower():
                    existing = processed_data[doc_name]
                    existing_addendum = (existing.get('chapters', {}).get('2_contract_details', {})
                                         .get('sub_chapters', {}).get('addendum', {}))
                    existing_content = existing_addendum.get('content', [])
                    existing_clauses = existing_addendum.get('clauses', {})
                    existing_has_addendum = (
                        bool(existing_clauses) or
                        (existing_content and 
                         (len(existing_content) > 1 or existing_content[0].strip().upper() != 'N/A'))
                    )
                    if existing_has_addendum:
                        logging.info(f"Skipping {filename} - keeping existing '{doc_name}' with complete addendum")
                        continue
                
                processed_data[doc_name] = document_data
                logging.info(f"Successfully processed {filename} as '{doc_name}'")
                
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
                continue
        
        logging.info(f"Processing complete. Extracted data for {len(processed_data)} documents")
        return processed_data
    
    def save_outputs(self, processed_data: Dict[str, Dict[str, Any]]):
        """Save processed data to JSON and Markdown files"""
        os.makedirs(self.output_dir, exist_ok=True)
        
        for doc_name, doc_data in processed_data.items():
            # Create directory for this document
            doc_dir = os.path.join(self.output_dir, doc_name)
            os.makedirs(doc_dir, exist_ok=True)
            
            # Save JSON
            json_file = os.path.join(doc_dir, f"{doc_name}_data.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, indent=2, ensure_ascii=False)
            
            # Save Markdown
            md_file = os.path.join(doc_dir, f"{doc_name}_data.md")
            self._save_markdown(doc_data, md_file)
            
            logging.info(f"Saved outputs for {doc_name}")
    
    def _save_markdown(self, doc_data: Dict[str, Any], md_file: str):
        """Save document data as markdown - organized by chapters with sub-chapters"""
        with open(md_file, 'w', encoding='utf-8') as f:
            doc_name = doc_data.get('document_name', 'Document')
            
            # Document header
            f.write(f"# {doc_name}\n\n")
            
            chapters = doc_data.get('chapters', {})
            if chapters:
                # Write each chapter
                for chapter_key in sorted(chapters.keys()):
                    chapter = chapters[chapter_key]
                    chapter_title = chapter.get('title', 'Unknown Chapter')
                    
                    f.write(f"## {chapter_title}\n\n")
                    
                    # Special handling for Contract Details with sub-chapters
                    if chapter_key == "2_contract_details" and chapter.get('sub_chapters'):
                        # Write sub-chapters
                        for subchapter_key in ['fixture_recap', 'charter_party', 'addendum']:
                            if subchapter_key in chapter['sub_chapters']:
                                subchapter = chapter['sub_chapters'][subchapter_key]
                                subchapter_title = subchapter.get('title', 'Unknown Sub-Chapter')
                                
                                f.write(f"### {subchapter_title}\n\n")
                                
                                # Write sub-chapter content
                                for content in subchapter.get('content', []):
                                    if content.strip():
                                        f.write(f"{content}\n\n")
                                
                                f.write("\n")
                        
                        # Write chapter tables (if any)
                        for i, table in enumerate(chapter.get('tables', [])):
                            if table:
                                f.write(f"### Table {i + 1}\n\n")
                                for row in table:
                                    f.write("| " + " | ".join(str(cell) for cell in row) + " |\n")
                                f.write("\n")
                    
                    else:
                        # Regular chapter handling
                        # Write chapter content
                        for content in chapter.get('content', []):
                            if content.strip():
                                f.write(f"{content}\n\n")
                        
                        # Write chapter tables
                        for i, table in enumerate(chapter.get('tables', [])):
                            if table:
                                f.write(f"### Table {i + 1}\n\n")
                                for row in table:
                                    f.write("| " + " | ".join(str(cell) for cell in row) + " |\n")
                                f.write("\n")
                    
                    f.write("\n")
            else:
                # Fallback to full text if no chapters
                full_text = doc_data.get('full_text', '')
                if full_text:
                    f.write(full_text)
                else:
                    f.write("*No text content extracted*")

def main():
    """Main function to run the dynamic text extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Dynamic Text Extraction System')
    parser.add_argument('--source', default='source/vessels', 
                        help='Source directory containing files to process')
    parser.add_argument('--output', default='output/vessels', 
                        help='Output directory for processed data')
    args = parser.parse_args()
    
    # Create extractor and process files
    extractor = DynamicTextExtractor(args.source, args.output)
    processed_data = extractor.process_files()
    
    if processed_data:
        extractor.save_outputs(processed_data)
        print(f"\nSuccessfully processed {len(processed_data)} documents!")
        print(f"Output saved to: {args.output}")
        
        # Display summary
        for doc_name, doc_data in processed_data.items():
            full_text = doc_data.get('full_text', '')
            word_count = len(full_text.split()) if full_text else 0
            char_count = len(full_text) if full_text else 0
            print(f"  - {doc_name}: {word_count} words, {char_count} characters")
    else:
        print("No documents were processed")

if __name__ == "__main__":
    main()
