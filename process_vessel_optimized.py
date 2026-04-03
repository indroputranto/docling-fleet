#!/usr/bin/env python3
"""
Optimized Vessel Data Extractor - Enhanced for AI Search
Extracts vessel information with semantic enhancement for better AI chatbot performance
"""

import os
import json
import logging
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class TwoTierSectionDetector:
    """Two-tier section detection: heading patterns first, content analysis fallback"""
    
    def __init__(self):
        # Content keyword mapping for fallback detection
        self.content_patterns = {
            'vessel_specifications': {
                'keywords': ['vessel name', 'imo number', 'dwt', 'loa', 'beam', 'draft', 'class notation'],
                'subsections': {
                    'basic_info': ['vessel name', 'imo number', 'flag state'],
                    'dimensions': ['loa', 'beam', 'draft', 'depth'],
                    'capacities': ['dwt', 'deadweight', 'cargo capacity'],
                    'equipment': ['crane', 'cargo gear', 'container']
                }
            },
            'bunkers': {
                'keywords': ['bunker', 'fuel', 'vlsfo', 'mgo', 'delivery', 'redelivery', 'sampling'],
                'strong_indicators': ['bunker quantities', 'fuel specifications', 'iso 8217']
            },
            'trading_exclusions': {
                'keywords': ['trading exclusions', 'prohibited', 'sanctions', 'iran', 'cuba', 'north korea'],
                'strong_indicators': ['inl/iwl', 'sanctioned countries', 'blacklisting']
            },
            'cargo_exclusions': {
                'keywords': ['cargo exclusions', 'dangerous', 'hazardous', 'explosive', 'radioactive'],
                'strong_indicators': ['livestock', 'ammunition', 'nuclear material']
            },
            'delivery': {
                'keywords': ['delivery', 'dlosp', 'port', 'on-hire'],
                'strong_indicators': ['delivery port', 'delivery notice']
            },
            'hire': {
                'keywords': ['hire', 'daily', 'usd', 'payment'],
                'strong_indicators': ['rate of hire', 'hire payment']
            }
        }
    
    def detect_by_heading_patterns(self, content: str) -> tuple:
        """Primary detection using document heading patterns"""
        
        # Fixture Recap patterns (very specific)
        if re.match(r"Fixture recap - \d+\.?\s*-?\s*(.+?)[:.]?", content):
            match = re.search(r"Fixture recap - \d+\.?\s*-?\s*(.+?)[:.]?", content)
            subsection = self.clean_subsection_name(match.group(1)) if match else "unknown"
            return "fixture_recap", subsection, "pattern_match"
        
        # Charter Party patterns (Heading 1)
        elif re.match(r"Charter Party – Clause \d+ - (.+)", content):
            match = re.search(r"Charter Party – Clause \d+ - (.+)", content)
            subsection = self.clean_subsection_name(match.group(1)) if match else "unknown"
            return "charter_party", subsection, "pattern_match"
        
        # Vessel Specifications patterns (Heading 3, blue text)
        elif re.match(r"Vessels? [Ss]pecifications? - (.+)", content):
            match = re.search(r"Vessels? [Ss]pecifications? - (.+)", content)
            subsection = self.clean_subsection_name(match.group(1)) if match else "unknown"
            return "vessel_specifications", subsection, "pattern_match"
        
        # Addendum patterns
        elif re.match(r"Addendum", content):
            return "addendum", "general", "pattern_match"
        
        # No clear pattern found
        return None, None, "no_pattern"
    
    def detect_by_content_analysis(self, content: str) -> tuple:
        """Fallback detection using content keywords"""
        content_lower = content.lower()
        
        # Check each content pattern
        for section_type, patterns in self.content_patterns.items():
            if self.has_strong_content_match(content_lower, patterns):
                subsection = self.detect_subsection_by_content(content_lower, patterns.get('subsections', {}))
                return section_type, subsection, "content_match"
        
        return "other", "general", "content_fallback"
    
    def has_strong_content_match(self, content: str, patterns: dict) -> bool:
        """Check if content strongly matches a section type"""
        keyword_count = sum(1 for keyword in patterns['keywords'] if keyword in content)
        strong_indicators = sum(1 for indicator in patterns.get('strong_indicators', []) if indicator in content)
        
        # Strong match: multiple keywords OR any strong indicator
        return keyword_count >= 2 or strong_indicators >= 1
    
    def detect_subsection_by_content(self, content: str, subsections: dict) -> str:
        """Detect subsection based on content analysis"""
        for subsection_name, keywords in subsections.items():
            if any(keyword in content for keyword in keywords):
                return subsection_name
        return "general"
    
    def clean_subsection_name(self, name: str) -> str:
        """Clean and normalize subsection names"""
        # Remove punctuation and normalize
        cleaned = re.sub(r'[^\w\s-]', '', name.strip())
        cleaned = re.sub(r'\s+', '_', cleaned.lower())
        return cleaned
    
    def detect_section_integrated(self, content: str) -> tuple:
        """Integrated detection: patterns first, content fallback"""
        
        # TIER 1: Try heading pattern detection
        section, subsection, method = self.detect_by_heading_patterns(content)
        
        if method == "pattern_match":
            print(f"  Pattern detected: {section}/{subsection}")
            return section, subsection, "pattern"
        
        # TIER 2: Fallback to content analysis
        section, subsection, method = self.detect_by_content_analysis(content)
        
        if method in ["content_match", "content_fallback"]:
            print(f"  Content detected: {section}/{subsection}")
            return section, subsection, "content"
        
        # TIER 3: Last resort - generic classification
        print(f"  Generic classification")
        return "other", "general", "generic"
    
    def create_clean_chunk(self, content: str, vessel_name: str) -> dict:
        """Create a clean chunk with minimal metadata"""
        section, subsection, detection_method = self.detect_section_integrated(content)
        
        return {
            "vessel": vessel_name,
            "section": section,
            "subsection": subsection,
            "content": content,
            "heading": content[:100] + "..." if len(content) > 100 else content,
            "detection_method": detection_method
        }

class OptimizedVesselDataExtractor:
    def __init__(self, source_dir: str = 'source/vessels'):
        """Initialize the optimized vessel data extractor"""
        self.source_dir = source_dir
        self.semantic_enhancer = SemanticClauseEnhancer()
        self.extracted_data = {}
        self._strikethrough_cache = {}
        self._text_cache = {}
    
    def _remove_strikethrough_text(self, text: str) -> str:
        """Process strikethrough text patterns from the given text by wrapping them in [STRIKETHROUGH] tags."""
        import re
        
        # Pattern 1: [[STRIKE ... STRIKE]] - common DOCX strike-through format
        text = re.sub(r'\[\[STRIKE\s+(.*?)\s+STRIKE\]\]', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 2: QUOTE STRIKE-THROUGH ... UNQUOTE STRIKE-THROUGH
        text = re.sub(r'QUOTE STRIKE-THROUGH\s+(.*?)\s+UNQUOTE STRIKE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 3: Simple strike-through markers
        text = re.sub(r'STRIKE\s+(.*?)\s+STRIKE', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Pattern 4: Standalone STRIKE markers
        text = re.sub(r'\bSTRIKE\b', '', text, flags=re.IGNORECASE)
        
        # Pattern 5: QUOTE-THROUGH ... UNQUOTE-THROUGH
        text = re.sub(r'QUOTE-THROUGH\s+(.*?)\s+UNQUOTE-THROUGH', r'[STRIKETHROUGH]\1[/STRIKETHROUGH]', text, flags=re.DOTALL | re.IGNORECASE)
        
        return text.strip()

    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to group consecutive strikethrough runs and wrap them in [STRIKETHROUGH]...[/STRIKETHROUGH] tags."""
        if not hasattr(paragraph, 'runs'):
            return paragraph.text if hasattr(paragraph, 'text') else str(paragraph)
        
        result_parts = []
        current_strikethrough_parts = []
        in_strikethrough = False
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Check if this run has strikethrough formatting
            is_strikethrough = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Check for strike elements
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if strike_elem is not None or dstrike_elem is not None:
                        is_strikethrough = True
            
            if is_strikethrough:
                if not in_strikethrough:
                    # Starting a new strikethrough section
                    in_strikethrough = True
                current_strikethrough_parts.append(text)
            else:
                if in_strikethrough:
                    # Ending strikethrough section
                    if current_strikethrough_parts:
                        result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
                        current_strikethrough_parts = []
                    in_strikethrough = False
                result_parts.append(text)
        
        # Handle any remaining strikethrough text
        if current_strikethrough_parts:
            result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
        
        return ''.join(result_parts)

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file, preserving strikethrough formatting."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs, using the strikethrough-aware function
            for paragraph in doc.paragraphs:
                text = self._process_paragraph_for_strikethrough(paragraph)
                if text:
                    text_parts.append(text)
            
            # Extract tables and add them to the text
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- TABLE {i+1} ---")
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # For each cell, process all paragraphs for strikethrough
                        cell_text = []
                        for p in cell.paragraphs:
                            processed_text = self._process_paragraph_for_strikethrough(p)
                            if processed_text:
                                cell_text.append(processed_text)
                        row_cells.append(' '.join(cell_text).strip())
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("--- END TABLE ---\n")
            
            # Join all parts with double newlines for better section separation
            return '\n\n'.join(text_parts)
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_vessel_name(self, text: str, filename: str) -> str:
        """Extract vessel name from text or filename."""
        # Try to find vessel name in text first
        lines = text.split('\n')[:20]  # Check first 20 lines
        for line in lines:
            # Look for patterns like "Vessel Name: MV SOMETHING"
            match = re.search(r'vessel\s+name:?\s*(.+)', line, re.IGNORECASE)
            if match:
                vessel_name = match.group(1).strip()
                if vessel_name and not vessel_name.lower().startswith('insert'):
                    return vessel_name.replace('"', '').replace("'", "")
        
        # Fallback: extract from filename
        base_name = os.path.splitext(filename)[0]
        # Remove common prefixes/suffixes
        vessel_name = re.sub(r'^(mv|m/v|vessel)[\s_-]*', '', base_name, flags=re.IGNORECASE)
        vessel_name = re.sub(r'[\s_-]*(final|ns|done).*$', '', vessel_name, flags=re.IGNORECASE)
        vessel_name = vessel_name.replace('_', ' ').replace('-', ' ').strip()
        
        # Add MV prefix if not present
        if vessel_name and not vessel_name.upper().startswith('MV'):
            vessel_name = f"MV {vessel_name.upper()}"
        
        return vessel_name

    def _extract_vessel_details_structured(self, text: str, vessel_name: str) -> List[Dict[str, Any]]:
        """Extract vessel details as structured, semantic chunks"""
        items = []
        
        # Split text into logical chunks by line breaks and common separators
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        current_chunk = []
        
        for line in lines:
            # Skip headers and section titles
            if re.match(r'^\d+\.?\s*vessel\s*details', line, re.IGNORECASE):
                continue
            
            # Add to current chunk
            if line:
                current_chunk.append(line)
            
            # Check if this line ends a logical group (by looking for common section breaks)
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in ['dimensions', 'capacities', 'propulsion', 'equipment', 'navigation', 'certificates']):
                if current_chunk:
                    # Create enhanced data for the current chunk
                    chunk_content = '\n'.join(current_chunk)
                    enhanced_data = {
                        "type": "specification",
                        "content": chunk_content,
                        "vessel_name": vessel_name,
                        "semantic_keywords": [
                            "vessel specifications",
                            "vessel details",
                            "technical specifications",
                            "vessel particulars",
                            "ship specifications"
                        ],
                        "description": f"Vessel specifications for {vessel_name} containing detailed technical information and characteristics",
                        "priority": "high",
                        "is_primary_spec": True,
                        "enhanced_for_search": True,
                        "is_vessel_specification": True
                    }
                    items.append(enhanced_data)
                    current_chunk = []
        
        # Add any remaining chunk
        if current_chunk:
            chunk_content = '\n'.join(current_chunk)
            enhanced_data = {
                "type": "specification",
                "content": chunk_content,
                "vessel_name": vessel_name,
                "semantic_keywords": [
                    "vessel specifications",
                    "vessel details", 
                    "technical specifications",
                    "vessel particulars",
                    "ship specifications"
                ],
                "description": f"Vessel specifications for {vessel_name} containing detailed technical information and characteristics",
                "priority": "high",
                "is_primary_spec": True,
                "enhanced_for_search": True,
                "is_vessel_specification": True
            }
            items.append(enhanced_data)
        
        return items

    def _extract_contract_details_enhanced(self, doc: Document, vessel_name: str) -> List[Dict[str, Any]]:
        """Extract contract details with semantic enhancement"""
        items = []
        in_contract_section = False
        current_content = []
        current_subheader = None
        special_section = None
        special_content = []
        special_type = None
        in_charter_party = False

        def is_heading(para, levels=(1,2,3,4)):
            if para.style and para.style.name:
                for lvl in levels:
                    if para.style.name.lower().startswith(f'heading {lvl}'):
                        return True
            return False

        def is_addendum_section(text):
            return re.search(r'2\.?\s*1\.?\s*addendum', text, re.IGNORECASE) or (re.search(r'addendum', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_fixture_recap_section(text):
            return re.search(r'2\.?\s*2\.?\s*fixture\s*recap', text, re.IGNORECASE) or (re.search(r'fixture\s*recap', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        def is_charter_party_section(text):
            return re.search(r'2\.?\s*3\.?\s*charter\s*party', text, re.IGNORECASE) or (re.search(r'charter\s*party', text, re.IGNORECASE) and re.match(r'^\d+\.', text.strip()))

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            text = self._process_paragraph_for_strikethrough(paragraph)
            
            if re.search(r'contract\s*details', text, re.IGNORECASE):
                in_contract_section = True
                continue
            
            if not in_contract_section:
                continue

            # Check for special sections
            if is_addendum_section(text):
                if current_content:
                    items.extend(current_content)
                    current_content = []
                if special_section:
                    if special_type == "fixture_recap":
                        # Use new parsing method for fixture recap breakdown
                        fixture_content = '\n'.join(special_content)
                        parsed_sections = self.semantic_enhancer._parse_fixture_recap_sections(fixture_content, vessel_name)
                        items.extend(parsed_sections)
                    else:
                        # Handle other special sections normally
                        items.append({
                            "type": special_type,
                            "content": '\n'.join(special_content),
                            "tag": special_type.upper(),
                            "vessel_name": vessel_name,
                            "priority": "medium",  # Changed from "low" to "medium"
                            "enhanced_for_search": True  # Changed from False to True
                        })
                special_section = "addendums"
                special_type = "addendums"
                special_content = [text]
                continue

            if is_fixture_recap_section(text):
                if current_content:
                    items.extend(current_content)
                    current_content = []
                if special_section:
                    if special_type == "fixture_recap":
                        # Use new parsing method for fixture recap breakdown
                        fixture_content = '\n'.join(special_content)
                        parsed_sections = self.semantic_enhancer._parse_fixture_recap_sections(fixture_content, vessel_name)
                        items.extend(parsed_sections)
                    else:
                        # Handle other special sections normally
                        items.append({
                            "type": special_type,
                            "content": '\n'.join(special_content),
                            "tag": special_type.upper(),
                            "vessel_name": vessel_name,
                            "priority": "medium",  # Changed from "low" to "medium"
                            "enhanced_for_search": True  # Changed from False to True
                        })
                special_section = "fixture_recap"
                special_type = "fixture_recap"
                special_content = [text]
                continue

            if is_charter_party_section(text):
                if special_section:
                    if special_type == "fixture_recap":
                        # Use new parsing method for fixture recap breakdown
                        fixture_content = '\n'.join(special_content)
                        parsed_sections = self.semantic_enhancer._parse_fixture_recap_sections(fixture_content, vessel_name)
                        items.extend(parsed_sections)
                    else:
                        # Handle other special sections normally
                        items.append({
                            "type": special_type,
                            "content": '\n'.join(special_content),
                            "tag": special_type.upper(),
                            "vessel_name": vessel_name,
                            "priority": "medium",  # Changed from "low" to "medium"
                            "enhanced_for_search": True  # Changed from False to True
                        })
                    special_section = None
                in_charter_party = True
                continue

            if special_section:
                special_content.append(text)
                continue

            # Process charter party content with semantic enhancement
            if in_charter_party:
                is_heading1 = is_heading(paragraph, levels=(1,))
                
                if is_heading1:
                    if current_content:
                        items.extend(current_content)
                        current_content = []
                    
                    current_subheader = text
                    
                    # Create enhanced subheader with semantic information
                    clean_heading = re.sub(r'^\d+\.\s*', '', current_subheader)
                    clean_heading = re.sub(r'[^\w\s-]', '', clean_heading).strip()
                    contextual_tag = f"CLAUSE - {clean_heading}" if clean_heading else "CLAUSE"
                    
                    # Add semantic enhancement to subheader
                    clause_type = self.semantic_enhancer.identify_clause_type(current_subheader, "")
                    enhanced_subheader = {
                        "type": "subheader",
                        "content": text,
                        "tag": contextual_tag,
                        "clause_type": clause_type,
                        "vessel_name": vessel_name,
                        "priority": "high",
                        "enhanced_for_search": True
                    }
                    
                    items.append(enhanced_subheader)
                else:
                    # Regular text with semantic enhancement
                    text = re.sub(r'\s+', ' ', text)
                    text = re.sub(r'\s+([.,;:])', r'\1', text)
                    
                    # Create contextual tag with subheader
                    tag = "CLAUSE"
                    if current_subheader:
                        clean_subheader = re.sub(r'^\d+\.\s*', '', current_subheader)
                        clean_subheader = re.sub(r'[^\w\s-]', '', clean_subheader).strip()
                        if clean_subheader:
                            tag = f"CLAUSE - {clean_subheader}"
                    
                    # Apply semantic enhancement
                    enhanced_item = self.semantic_enhancer.enhance_clause_data(
                        vessel_name, current_subheader or "", text, tag
                    )
                    current_content.append(enhanced_item)

        # Flush any remaining content
        if current_content:
            items.extend(current_content)
        if special_section:
            if special_type == "fixture_recap":
                # Use new parsing method for fixture recap breakdown
                fixture_content = '\n'.join(special_content)
                parsed_sections = self.semantic_enhancer._parse_fixture_recap_sections(fixture_content, vessel_name)
                items.extend(parsed_sections)
            else:
                # Handle other special sections normally
                items.append({
                    "type": special_type,
                    "content": '\n'.join(special_content),
                    "tag": special_type.upper(),
                    "vessel_name": vessel_name,
                    "priority": "medium",  # Changed from "low" to "medium"
                    "enhanced_for_search": True  # Changed from False to True
                })

        return items

    def _extract_vessel_details(self, text: str, vessel_name: str, doc: Document = None) -> Dict[str, Any]:
        """Extract vessel details with semantic enhancement"""
        import re
        details = {
            'vessel_details': {'structured': []},  # Changed to structured like contract_details
            'contract_details': {'structured': []},
            'delivery_figures': {'content': ''},
            'speed_and_consumption': {'content': ''},
            'lifting_equipment': {'content': ''},
            'hseq_documentation': {'content': ''}
        }
        
        # Section patterns
        section_patterns = [
            ('vessel_details', re.compile(r'^1\.?\s*vessel details', re.IGNORECASE)),
            ('contract_details', re.compile(r'^2\.?\s*contract details', re.IGNORECASE)),
            ('delivery_figures', re.compile(r'^3\.?\s*delivery figures', re.IGNORECASE)),
            ('speed_and_consumption', re.compile(r'^4\.?\s*speed and consumption', re.IGNORECASE)),
            ('lifting_equipment', re.compile(r'^5\.?\s*lifting equipment', re.IGNORECASE)),
            ('hseq_documentation', re.compile(r'^6\.?\s*hseq documentation', re.IGNORECASE)),
        ]
        
        # Process strikethrough markers
        strikethrough_blocks = []
        def save_strikethrough(match):
            block_id = f"__STRIKETHROUGH_{len(strikethrough_blocks)}__"
            strikethrough_blocks.append(match.group(1))
            return block_id
        
        processed_text = re.sub(r'\[STRIKETHROUGH\](.*?)\[/STRIKETHROUGH\]', save_strikethrough, text, flags=re.DOTALL)
        
        # Find section boundaries
        lines = processed_text.split('\n')
        section_indices = {}
        for i, line in enumerate(lines):
            for key, pattern in section_patterns:
                if key not in section_indices and pattern.match(line.strip()):
                    section_indices[key] = i
        
        # Extract content for each section
        sorted_sections = sorted(section_indices.items(), key=lambda x: x[1])
        for idx, (key, start) in enumerate(sorted_sections):
            end = sorted_sections[idx + 1][1] if idx + 1 < len(sorted_sections) else len(lines)
            section_lines = lines[start + 1:end]
            section_text = '\n'.join(section_lines).strip()
            
            # Restore strikethrough blocks
            for i, block in enumerate(strikethrough_blocks):
                section_text = section_text.replace(f"__STRIKETHROUGH_{i}__", f"[STRIKETHROUGH]{block}[/STRIKETHROUGH]")
            
            if key == 'contract_details' and doc is not None:
                details[key]['structured'] = self._extract_contract_details_enhanced(doc, vessel_name)
            elif key == 'vessel_details':
                details[key]['structured'] = self._extract_vessel_details_structured(section_text, vessel_name)
            else:
                details[key]['content'] = section_text
        
        return details

    def process_files(self) -> Dict[str, Dict[str, Any]]:
        """Process all vessel files with semantic enhancement"""
        vessels_data = {}
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return vessels_data
        
        # Filter for DOCX files
        docx_files = [f for f in os.listdir(self.source_dir) 
                     if f.lower().endswith('.docx') and not f.startswith('~$')]
        logging.info(f"Found {len(docx_files)} DOCX files to process: {docx_files}")
        
        for filename in docx_files:
            file_path = os.path.join(self.source_dir, filename)
            logging.info(f"Processing {filename}")
            
            try:
                # Extract text and document
                full_text = self._extract_text_from_docx(file_path)
                doc = Document(file_path)
                
                vessel_name = self._extract_vessel_name(full_text, filename)
                
                if vessel_name:
                    vessel_details = self._extract_vessel_details(full_text, vessel_name, doc)
                    vessels_data[vessel_name] = vessel_details
                    logging.info(f"Successfully extracted data for {vessel_name} with semantic enhancement")
                else:
                    logging.warning(f"Could not extract vessel name from {filename}")
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
        
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels: {list(vessels_data.keys())}")
        return vessels_data

    def save_outputs(self, vessels_data: Dict[str, Dict[str, Any]], output_dir: str = 'output/vessels'):
        """Save processed vessel data to JSON and markdown files"""
        os.makedirs(output_dir, exist_ok=True)
        
        for vessel_name, vessel_data in vessels_data.items():
            vessel_dir = os.path.join(output_dir, vessel_name)
            os.makedirs(vessel_dir, exist_ok=True)
            
            # Save JSON with semantic enhancements
            json_file = os.path.join(vessel_dir, f"{vessel_name}_data.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(vessel_data, f, indent=2, ensure_ascii=False)
            
            # Save markdown
            md_file = os.path.join(vessel_dir, f"{vessel_name}_data.md")
            self._save_markdown(vessel_data, md_file)
            
            logging.info(f"Saved enhanced data for {vessel_name}")

    def _save_markdown(self, vessel_data: Dict[str, Any], md_file: str):
        """Save vessel data as markdown with semantic structure"""
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(f"# {vessel_data.get('vessel_details', {}).get('content', '').split('Vessel Name:')[-1].split()[0] if 'Vessel Name:' in vessel_data.get('vessel_details', {}).get('content', '') else 'Vessel'} Data\n\n")
            
            # Add structured vessel details
            if vessel_data.get('vessel_details', {}).get('structured'):
                f.write("## Vessel Details\n\n")
                for item in vessel_data['vessel_details']['structured']:
                    if isinstance(item, dict):
                        spec_type = item.get('spec_type', 'general')
                        content = item.get('content', '')
                        if content:
                            f.write(f"### {spec_type.replace('_', ' ').title()}\n\n")
                            if item.get('enhanced_for_search'):
                                f.write(f"*Specification Type: {spec_type}*\n\n")
                            f.write(f"{content}\n\n")
            
            # Add enhanced contract details
            if vessel_data.get('contract_details', {}).get('structured'):
                f.write("## Contract Details\n\n")
                for item in vessel_data['contract_details']['structured']:
                    if isinstance(item, dict):
                        if item.get('type') == 'subheader':
                            f.write(f"### {item['content']}\n\n")
                            if item.get('enhanced_for_search'):
                                f.write(f"*Clause Type: {item.get('clause_type', 'general')}*\n\n")
                        elif item.get('content'):
                            f.write(f"{item['content']}\n\n")
            
            # Add other sections
            for section in ['delivery_figures', 'speed_and_consumption', 'lifting_equipment', 'hseq_documentation']:
                if vessel_data.get(section, {}).get('content'):
                    f.write(f"## {section.replace('_', ' ').title()}\n\n")
                    f.write(vessel_data[section]['content'])
                    f.write("\n\n")

def main():
    """Main function to run the optimized vessel data extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Optimized Vessel Data Extractor with Semantic Enhancement')
    parser.add_argument('--source', default='source/vessels', help='Source directory containing vessel DOCX files')
    parser.add_argument('--output', default='output/vessels', help='Output directory for processed data')
    args = parser.parse_args()
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('vessel_extraction_optimized.log'),
            logging.StreamHandler()
        ]
    )
    
    # Create extractor and process files
    extractor = OptimizedVesselDataExtractor(args.source)
    vessels_data = extractor.process_files()
    
    if vessels_data:
        extractor.save_outputs(vessels_data, args.output)
        print(f"\n✅ Successfully processed {len(vessels_data)} vessels with semantic enhancement!")
        print(f"📁 Output saved to: {args.output}")
        print(f"🚀 Ready for optimized AI search!")
    else:
        print("❌ No vessels were processed")

if __name__ == "__main__":
    main() 