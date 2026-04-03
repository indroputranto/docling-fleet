#!/usr/bin/env python3
"""
Simplified Vessel Data Extractor - Two-Tier Detection System
Uses heading patterns first, content analysis as fallback
"""

import os
import json
import logging
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any, Tuple
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class TwoTierSectionDetector:
    """Two-tier section detection: heading patterns first, content analysis fallback"""
    
    def __init__(self):
        # Content keyword mapping for fallback detection
        self.content_patterns = {
            'vessel_specifications': {
                'keywords': ['vessel name', 'imo number', 'dwt', 'loa', 'beam', 'draft', 'class notation'],
                'subsections': {
                    'basic_info': ['vessel name', 'imo number', 'flag state'],
                    'dimensions': ['loa', 'beam', 'draft', 'depth'],
                    'capacities': ['dwt', 'deadweight', 'cargo capacity'],
                    'equipment': ['crane', 'cargo gear', 'container']
                }
            },
            'bunkers': {
                'keywords': ['bunker', 'fuel', 'vlsfo', 'mgo', 'delivery', 'redelivery', 'sampling'],
                'strong_indicators': ['bunker quantities', 'fuel specifications', 'iso 8217']
            },
            'trading_exclusions': {
                'keywords': ['trading', 'trading exclusions', 'prohibited', 'sanctions', 'iran', 'cuba', 'north korea'],
                'strong_indicators': ['inl/iwl', 'sanctioned countries', 'blacklisting']
            },
            'cargo_exclusions': {
                'keywords': ['cargo exclusions', 'dangerous', 'hazardous', 'explosive', 'radioactive'],
                'strong_indicators': ['livestock', 'ammunition', 'nuclear material']
            },
            'delivery': {
                'keywords': ['delivery', 'dlosp', 'port', 'on-hire'],
                'strong_indicators': ['delivery port', 'delivery notice']
            },
            'hire': {
                'keywords': ['hire', 'daily', 'usd', 'payment'],
                'strong_indicators': ['rate of hire', 'hire payment']
            }
        }
    
    def detect_by_heading_patterns(self, content: str) -> Tuple[str, str, str]:
        """Primary detection using Microsoft Word heading patterns"""
        
        # Get the first line to check for heading patterns
        first_line = content.split('\n')[0].strip()
        
        # Vessel Specifications (consolidate all into one)
        if re.match(r"^Vessels? [Ss]pecifications? - (.+)", first_line):
            return "vessel_specifications", "all_specifications", "pattern_match"
        
        # Fixture Recap (numbered items pattern)
        elif re.match(r"^Fixture recap - (\d+)\.?\s*-?\s*(.+)", first_line):
            match = re.search(r"^Fixture recap - (\d+)\.?\s*-?\s*(.+)", first_line)
            if match:
                item_text = match.group(2).strip().rstrip(':.')
                subsection = self.clean_subsection_name(item_text)
            else:
                subsection = "general"
            return "fixture_recap", subsection, "pattern_match"
        
        # Charter Party (Heading 1 pattern)
        elif re.match(r"^Charter Party – Clause (\d+) - (.+)", first_line):
            match = re.search(r"^Charter Party – Clause (\d+) - (.+)", first_line)
            if match:
                clause_text = match.group(2).strip()
                subsection = self.clean_subsection_name(clause_text)
            else:
                subsection = "general"
            return "charter_party", subsection, "pattern_match"
        
        # Individual Charter Party section headings - exact matches only
        elif first_line == "Redelivery":
            return "charter_party", "redelivery", "pattern_match"
        elif first_line == "Bunkers":
            return "charter_party", "bunkers", "pattern_match"
        elif first_line == "Performance of Voyages":
            return "charter_party", "performance_of_voyages", "pattern_match"
        elif first_line == "Speed and Consumption":
            return "charter_party", "speed_and_consumption", "pattern_match"
        elif first_line == "Rate of Hire":
            return "charter_party", "rate_of_hire", "pattern_match"
        elif first_line == "Hold Cleaning":
            return "charter_party", "hold_cleaning", "pattern_match"
        elif first_line == "Hire Payment":
            return "charter_party", "hire_payment", "pattern_match"
        elif first_line == "Cash Advances":
            return "charter_party", "cash_advances", "pattern_match"
        elif first_line == "Spaces Available":
            return "charter_party", "spaces_available", "pattern_match"
        elif first_line == "Supercargo":
            return "charter_party", "supercargo", "pattern_match"
        elif first_line == "Sailing Orders and Logs":
            return "charter_party", "sailing_orders_and_logs", "pattern_match"
        elif first_line == "Cargo Exclusions":
            return "charter_party", "cargo_exclusions", "pattern_match"
        elif first_line == "OFF HIRE CLAUSE":
            return "charter_party", "off_hire_clause", "pattern_match"
        elif first_line == "Pollution":
            return "charter_party", "pollution", "pattern_match"
        elif first_line == "Drydocking":
            return "charter_party", "drydocking", "pattern_match"
        elif first_line == "Total Loss":
            return "charter_party", "total_loss", "pattern_match"
        elif first_line == "Exceptions":
            return "charter_party", "exceptions", "pattern_match"
        elif first_line == "Liberties":
            return "charter_party", "liberties", "pattern_match"
        elif first_line == "Liens":
            return "charter_party", "liens", "pattern_match"
        elif first_line == "Salvage":
            return "charter_party", "salvage", "pattern_match"
        elif first_line == "General Average":
            return "charter_party", "general_average", "pattern_match"
        elif first_line == "Navigation":
            return "charter_party", "navigation", "pattern_match"
        elif first_line == "Cargo Claims":
            return "charter_party", "cargo_claims", "pattern_match"
        elif first_line == "Cargo Handling Gear and Lights":
            return "charter_party", "cargo_handling_gear_and_lights", "pattern_match"
        elif first_line == "Bills of Lading":
            return "charter_party", "bills_of_lading", "pattern_match"
        elif first_line == "Protective Clauses":
            return "charter_party", "protective_clauses", "pattern_match"
        elif first_line in ["War Risks", "BIMCO War Risks"]:
            return "charter_party", "war_risks", "pattern_match"
        elif first_line == "Ice":
            return "charter_party", "ice", "pattern_match"
        elif first_line == "Requisition":
            return "charter_party", "requisition", "pattern_match"
        elif first_line == "Stevedore Damage":
            return "charter_party", "stevedore_damage", "pattern_match"
        elif first_line == "Slow Steaming":
            return "charter_party", "slow_steaming", "pattern_match"
        elif first_line in ["Piracy", "BIMCO Piracy"]:
            return "charter_party", "piracy", "pattern_match"
        
        # BIMCO Clauses - treat all as charter party clauses
        elif re.match(r"^BIMCO .+ Clause", first_line):
            # Extract clause type from BIMCO clause title
            match = re.search(r"^BIMCO (.+?) Clause", first_line)
            if match:
                clause_type = match.group(1).lower().replace(" ", "_")
                subsection = self.clean_subsection_name(clause_type)
            else:
                subsection = "bimco_general"
            return "charter_party", subsection, "pattern_match"
        elif first_line == "Taxes":
            return "charter_party", "taxes", "pattern_match"
        elif first_line == "Industrial Action":
            return "charter_party", "industrial_action", "pattern_match"
        elif first_line == "Stowaways":
            return "charter_party", "stowaways", "pattern_match"
        elif first_line == "Smuggling":
            return "charter_party", "smuggling", "pattern_match"
        elif first_line in ["International Safety Management", "ISM"]:
            return "charter_party", "international_safety_management", "pattern_match"
        elif first_line in ["International Ship and Port", "ISPS"]:
            return "charter_party", "international_ship_and_port", "pattern_match"
        elif first_line in ["Sanctions", "BIMCO Sanctions"]:
            return "charter_party", "sanctions", "pattern_match"
        
        # Addendum patterns
        elif re.match(r"^Addendum", first_line):
            return "addendum", "general", "pattern_match"
        
        # Other numbered sections
        elif re.match(r"^(\d+)\.?\s+(.+)", first_line):
            match = re.search(r"^(\d+)\.?\s+(.+)", first_line)
            if match:
                section_text = match.group(2).strip()
                subsection = self.clean_subsection_name(section_text)
            else:
                subsection = "general"
            return "other", subsection, "pattern_match"
        
        # No clear pattern found
        return None, None, "no_pattern"
    
    def detect_by_content_analysis(self, content: str) -> Tuple[str, str, str]:
        """Fallback detection using content keywords"""
        content_lower = content.lower()
        
        # Check each content pattern
        for section_type, patterns in self.content_patterns.items():
            if self.has_strong_content_match(content_lower, patterns):
                subsection = self.detect_subsection_by_content(content_lower, patterns.get('subsections', {}))
                return section_type, subsection, "content_match"
        
        return "other", "general", "content_fallback"
    
    def has_strong_content_match(self, content: str, patterns: dict) -> bool:
        """Check if content strongly matches a section type"""
        keyword_count = sum(1 for keyword in patterns['keywords'] if keyword in content)
        strong_indicators = sum(1 for indicator in patterns.get('strong_indicators', []) if indicator in content)
        
        # Strong match: multiple keywords OR any strong indicator
        return keyword_count >= 2 or strong_indicators >= 1
    
    def detect_subsection_by_content(self, content: str, subsections: dict) -> str:
        """Detect subsection based on content analysis"""
        for subsection_name, keywords in subsections.items():
            if any(keyword in content for keyword in keywords):
                return subsection_name
        return "general"
    
    def clean_subsection_name(self, name: str) -> str:
        """Clean and normalize subsection names"""
        # Remove punctuation and normalize
        cleaned = re.sub(r'[^\w\s-]', '', name.strip())
        cleaned = re.sub(r'\s+', '_', cleaned.lower())
        return cleaned
    
    def detect_section_integrated(self, content: str) -> Tuple[str, str, str]:
        """Integrated detection: patterns first, content fallback"""
        
        # TIER 1: Try heading pattern detection
        section, subsection, method = self.detect_by_heading_patterns(content)
        
        if method == "pattern_match":
            print(f"  Pattern detected: {section}/{subsection}")
            return section, subsection, "pattern"
        
        # TIER 2: Fallback to content analysis
        section, subsection, method = self.detect_by_content_analysis(content)
        
        if method in ["content_match", "content_fallback"]:
            print(f"  Content detected: {section}/{subsection}")
            return section, subsection, "content"
        
        # TIER 3: Last resort - generic classification
        print(f"  Generic classification")
        return "other", "general", "generic"
    
    def create_clean_chunk(self, vessel_name: str, section: str, subsection: str, content: str, heading: str, detection_method: str) -> Dict:
        """Create a clean chunk with minimal metadata"""
        
        chunk = {
            "vessel": vessel_name,
            "section": section,
            "subsection": subsection,
            "content": content,
            "heading": heading
        }
        
        # Filter out None or empty string values
        return {k: v for k, v in chunk.items() if v is not None and v != ""}

class SimpleVesselDataExtractor:
    def __init__(self, source_dir: str = 'source/vessels'):
        """Initialize the simplified vessel data extractor"""
        self.source_dir = source_dir
        self.detector = TwoTierSectionDetector()
        self.extracted_data = {}
        self._strikethrough_cache = {}
        self._text_cache = {}
    
    def _process_paragraph_for_strikethrough(self, paragraph) -> str:
        """Process a paragraph to group consecutive strikethrough runs and wrap them in [STRIKETHROUGH]...[/STRIKETHROUGH] tags."""
        if not hasattr(paragraph, 'runs'):
            return paragraph.text if hasattr(paragraph, 'text') else str(paragraph)
        
        result_parts = []
        current_strikethrough_parts = []
        in_strikethrough = False
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Check if this run has strikethrough formatting
            is_strikethrough = False
            if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                rPr = run._element.rPr
                if rPr is not None:
                    # Check for strike elements
                    strike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
                    dstrike_elem = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike')
                    
                    if strike_elem is not None or dstrike_elem is not None:
                        is_strikethrough = True
            
            if is_strikethrough:
                if not in_strikethrough:
                    # Starting a new strikethrough section
                    in_strikethrough = True
                current_strikethrough_parts.append(text)
            else:
                if in_strikethrough:
                    # Ending strikethrough section
                    if current_strikethrough_parts:
                        result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
                        current_strikethrough_parts = []
                    in_strikethrough = False
                result_parts.append(text)
        
        # Handle any remaining strikethrough text
        if current_strikethrough_parts:
            result_parts.append(f"[STRIKETHROUGH]{''.join(current_strikethrough_parts)}[/STRIKETHROUGH]")
        
        return ''.join(result_parts)

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file, preserving strikethrough formatting."""
        if file_path in self._text_cache:
            return self._text_cache[file_path]
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs, using the strikethrough-aware function
            for paragraph in doc.paragraphs:
                text = self._process_paragraph_for_strikethrough(paragraph)
                if text:
                    text_parts.append(text)
            
            # Extract tables and add them to the text
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- TABLE {i+1} ---")
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # For each cell, process all paragraphs for strikethrough
                        cell_text = []
                        for p in cell.paragraphs:
                            processed_text = self._process_paragraph_for_strikethrough(p)
                            if processed_text:
                                cell_text.append(processed_text)
                        row_cells.append(' '.join(cell_text).strip())
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("--- END TABLE ---\n")
            
            # Join all parts with double newlines for better section separation
            processed_text = '\n\n'.join(text_parts)
            self._text_cache[file_path] = processed_text
            return processed_text
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _extract_vessel_name(self, text: str, filename: str) -> str:
        """Extract vessel name from text or filename."""
        # Try to find vessel name from text
        lines = text.split('\n')[:20]  # Check first 20 lines
        for line in lines:
            if 'vessel name' in line.lower():
                parts = line.split(':')
                if len(parts) > 1:
                    name = parts[1].strip()
                    if name and len(name) > 2:
                        return name.upper()
        
        # Fallback to filename
        vessel_name = os.path.splitext(filename)[0]
        vessel_name = re.sub(r'[^a-zA-Z0-9\s]', ' ', vessel_name)
        vessel_name = re.sub(r'\s+', ' ', vessel_name).strip()
        return vessel_name.upper()

    def _split_into_sections(self, text: str, doc: Document = None) -> List[str]:
        """Split text into sections based on Microsoft Word heading patterns"""
        sections = []
        lines = text.split('\n')
        current_section = []
        
        # Create a mapping of line text to paragraph styles for minimal filtering
        style_map = {}
        if doc:
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() and paragraph.style and paragraph.style.name:
                    style_map[paragraph.text.strip()] = paragraph.style.name
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if this line should be ignored for section breaking (only very specific sub-elements)
            should_ignore_for_section_break = False
            if line in style_map:
                style_name = style_map[line]
                # Only ignore List Paragraph style for non-heading content
                # If the line matches heading patterns, don't ignore it even if it has List Paragraph style
                if 'List Paragraph' in style_name:
                    # Check if this line looks like a heading despite having List Paragraph style
                    is_heading_pattern = (
                        (re.match(r"^[A-Z][a-z]+(\s+[A-Z][a-z]+)*$", line) and len(line.split()) <= 4 and len(line) < 50) or
                        (re.match(r"^[A-Z][A-Z\s]+$", line) and len(line.split()) <= 4 and len(line) < 50) or
                        (re.match(r"^\d+\.\s+[A-Z][A-Za-z\s]+$", line) and len(line) < 100) or
                        re.match(r"^BIMCO .+", line) or
                        re.match(r"^(Charter Party|Clause \d+)", line) or
                        re.match(r"^Addendum", line) or
                        (re.match(r"^\d+\.\s+[A-Z][^.]*$", line) and len(line) < 100)
                    )
                    
                    if not is_heading_pattern:
                        should_ignore_for_section_break = True
            
            # Check if this is a heading that should start a new section
            is_heading = False
            
            # Context-aware section detection: check what type of section we're currently in
            in_charter_party_section = False
            in_complex_clause = False
            
            if current_section:
                section_start = current_section[0].lower()
                # Detect if we're currently inside Charter Party content
                in_charter_party_section = (
                    'charter party' in section_start or
                    'bimco' in section_start or
                    any(clause_indicator in section_start for clause_indicator in [
                        'clause', 'redelivery', 'bunkers', 'hire', 'drydocking', 'piracy', 'war risks',
                        'sanctions', 'designated entities', 'advance cargo', 'ballast water'
                    ])
                )
                
                # We're in a complex legal clause if current section starts with BIMCO or complex clause patterns
                # But this is less restrictive now since we have Heading 1 detection for Charter Party
                in_complex_clause = (
                    not in_charter_party_section and (  # Only apply complex clause logic outside Charter Party
                        'bimco' in section_start or
                        any(clause_type in section_start for clause_type in [
                            'piracy clause', 'war risks', 'sanctions clause', 'hull fouling',
                            'carbon intensity', 'emissions trading', 'slow steaming'
                        ]) or
                        len(' '.join(current_section)) > 1500  # Increased threshold since we have better detection
                    )
                )
            
            # Check for section breaks (allowing most heading styles through)
            if not should_ignore_for_section_break:
                # PRIORITY 1: Charter Party Heading 1 detection
                # If we're inside Charter Party content, use Heading 1 styles as primary delimiters
                if in_charter_party_section and line in style_map:
                    style_name = style_map[line]
                    if 'Heading 1' in style_name:
                        is_heading = True
                        print(f"    Heading 1 detected in Charter Party: {line[:50]}...")
                
                # PRIORITY 2: Vessel Specifications - consolidate ALL into one chunk
                elif re.match(r"^Vessels? [Ss]pecifications? - .+", line):
                    # If we're already in a vessel spec section, don't start new one
                    if current_section and any("Vessels Specifications" in s or "Vessel specifications" in s for s in current_section):
                        current_section.append(line)
                        i += 1
                        continue
                    else:
                        is_heading = True
                
                # PRIORITY 3: Fixture Recap items
                elif re.match(r"^Fixture recap - \d+\.?\s*-?\s*.+", line):
                    is_heading = True
                
                # PRIORITY 4: Charter Party main clauses (fallback for when Heading 1 detection fails)
                elif re.match(r"^Charter Party – Clause \d+ - .+", line):
                    is_heading = True
                
                # PRIORITY 5: Numbered clauses - THE PERFECT DELIMITER for Charter Party
                elif re.match(r"^\d+\.\s+", line):
                    # This matches exactly what we see in screenshots: "7. Bunkers", "28. BIMCO...", etc.
                    is_heading = True
                    if in_charter_party_section:
                        print(f"    Charter Party numbered clause detected: {line[:60]}...")
                    else:
                        print(f"    Numbered clause detected: {line[:60]}...")
                
                # PRIORITY 6: BIMCO clauses - can appear anywhere, including Charter Party
                elif re.match(r"^BIMCO .+", line):
                    is_heading = True
                    if in_charter_party_section:
                        print(f"    Charter Party BIMCO clause detected: {line[:60]}...")
                
                # PRIORITY 7: Single capitalized words - but only outside Charter Party and complex clauses
                elif (not in_charter_party_section and
                      re.match(r"^[A-Z][a-z]+(\s+[A-Z][a-z]+)*$", line) and 
                      len(line.split()) <= 4 and len(line) < 50 and
                      not in_complex_clause):
                    # Additional check: avoid breaking on obvious subheadings within clauses
                    if line.lower() not in ['costs', 'wages', 'employment', 'insurance', 'liability', 'claims', 
                                          'damages', 'expenses', 'payments', 'reimbursement', 'indemnity',
                                          'compliance', 'enforcement', 'obligations', 'requirements']:
                        is_heading = True
                
                # PRIORITY 8: All uppercase clause titles - but only outside Charter Party and complex clauses
                elif (not in_charter_party_section and
                      re.match(r"^[A-Z][A-Z\s]+$", line) and 
                      len(line.split()) <= 4 and len(line) < 50 and
                      not in_complex_clause):
                    is_heading = True
                
                # PRIORITY 9: Common charter party patterns (fallback)
                elif re.match(r"^(Charter Party|Clause \d+)", line):
                    is_heading = True
                
                # PRIORITY 10: Addendum patterns
                elif re.match(r"^Addendum", line):
                    is_heading = True
                
                # PRIORITY 11: Other numbered sections (outside Charter Party)
                elif (not in_charter_party_section and
                      re.match(r"^\d+\.\s+[A-Z][^.]*$", line) and len(line) < 100):
                    is_heading = True
            
            # If this is a heading, save previous section and start new one
            if is_heading:
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
                current_section.append(line)
            else:
                current_section.append(line)
            
            i += 1
        
        # Don't forget the last section
        if current_section:
            sections.append('\n'.join(current_section))
        
        return sections

    def _enhance_sections_with_tables(self, sections: List[str], text: str) -> List[str]:
        """Enhance sections with their associated tables using dynamic proximity and relevance analysis"""
        enhanced_sections = []
        
        # Find all tables in the text with their positions
        lines = text.split('\n')
        table_blocks = []
        current_table = None
        
        for i, line in enumerate(lines):
            if '--- TABLE' in line:
                if current_table is None:
                    current_table = {'start': i, 'content': [], 'line_start': i}
            elif '--- END TABLE ---' in line and current_table is not None:
                current_table['end'] = i
                current_table['line_end'] = i
                table_blocks.append(current_table)
                current_table = None
            elif current_table is not None:
                current_table['content'].append(line)
        
        # Map section positions in the text for proximity analysis
        section_positions = []
        current_line = 0
        for section in sections:
            section_lines = section.split('\n')
            section_start = current_line
            section_end = current_line + len(section_lines)
            section_positions.append({
                'start': section_start,
                'end': section_end,
                'content': section
            })
            current_line = section_end
        
        # For each section, find the most relevant table(s) using dynamic analysis
        for i, section in enumerate(sections):
            enhanced_section = section
            section_pos = section_positions[i]
            
            # Find the best matching table for this section
            best_table = self._find_best_table_for_section(section, section_pos, table_blocks)
            
            if best_table:
                table_content = '\n'.join(best_table['content'])
                enhanced_section += f"\n\n{table_content}"
            
            enhanced_sections.append(enhanced_section)
        
        return enhanced_sections
    
    def _find_best_table_for_section(self, section: str, section_pos: dict, table_blocks: list) -> dict:
        """Dynamically find the most relevant table for a section using proximity and content analysis"""
        if not table_blocks:
            return None
            
        section_lower = section.lower()
        best_table = None
        best_score = 0
        
        for table_block in table_blocks:
            table_content = '\n'.join(table_block['content']).lower()
            
            # Skip empty tables
            if not table_content.strip():
                continue
            
            # Calculate relevance score
            score = 0
            
            # 1. Proximity score (closer tables get higher scores)
            proximity_distance = min(
                abs(table_block['line_start'] - section_pos['end']),
                abs(table_block['line_end'] - section_pos['start'])
            )
            # Normalize proximity score (closer = higher score)
            proximity_score = max(0, 100 - proximity_distance)
            
            # 2. Content relevance score based on keyword overlap
            content_score = self._calculate_content_relevance(section_lower, table_content)
            
            # 3. Structural relevance - check if section heading suggests it should have a table
            structural_score = self._calculate_structural_relevance(section_lower, table_content)
            
            # 4. Anti-contamination check - heavily penalize if table seems unrelated
            contamination_penalty = self._calculate_contamination_penalty(section_lower, table_content)
            
            # Combine scores with weights
            total_score = (
                proximity_score * 0.3 + 
                content_score * 0.4 + 
                structural_score * 0.4 - 
                contamination_penalty * 0.6
            )
            
            if total_score > best_score and total_score > 50:  # Minimum threshold
                best_score = total_score
                best_table = table_block
        
        return best_table
    
    def _calculate_content_relevance(self, section_text: str, table_content: str) -> float:
        """Calculate how relevant table content is to section content"""
        # Extract key terms from both section and table
        section_words = set(re.findall(r'\b\w{3,}\b', section_text))
        table_words = set(re.findall(r'\b\w{3,}\b', table_content))
        
        # Calculate overlap
        overlap = len(section_words.intersection(table_words))
        total_unique = len(section_words.union(table_words))
        
        if total_unique == 0:
            return 0
        
        # Jaccard similarity with bonus for important keywords
        base_score = (overlap / total_unique) * 100
        
        # Bonus for exact phrase matches
        important_phrases = ['speed', 'consumption', 'fuel', 'knots', 'lifting', 'equipment', 'crane']
        phrase_bonus = 0
        for phrase in important_phrases:
            if phrase in section_text and phrase in table_content:
                phrase_bonus += 20
        
        return min(100, base_score + phrase_bonus)
    
    def _calculate_structural_relevance(self, section_text: str, table_content: str) -> float:
        """Check if section structure suggests it should have this type of table"""
        section_lines = section_text.split('\n')
        section_heading = section_lines[0].lower() if section_lines else ""
        
        # Check for explicit table-related headings
        if any(term in section_heading for term in ['speed and consumption', 'performance', 'specifications']):
            if any(term in table_content for term in ['speed', 'consumption', 'knots']):
                return 80
        
        if any(term in section_heading for term in ['lifting', 'equipment', 'crane']):
            if any(term in table_content for term in ['lifting', 'equipment', 'crane', 'capacity']):
                return 80
        
        # Check for numbered sections that might reference appendices or tables
        if re.match(r'^\d+\.?\s*', section_heading):
            return 30
        
        return 0
    
    def _calculate_contamination_penalty(self, section_text: str, table_content: str) -> float:
        """Penalize tables that seem completely unrelated to the section"""
        
        # Heavy penalty if section is about legal/clause content but table is technical data
        legal_indicators = ['clause', 'charter party', 'bimco', 'owners', 'charterers', 'liability', 'indemnify']
        technical_indicators = ['speed', 'consumption', 'fuel', 'knots', 'mt/day']
        
        section_is_legal = any(indicator in section_text for indicator in legal_indicators)
        table_is_technical = any(indicator in table_content for indicator in technical_indicators)
        
        if section_is_legal and table_is_technical:
            # Check if the legal section actually discusses technical performance
            if not any(term in section_text for term in ['speed', 'consumption', 'performance', 'fuel']):
                return 100  # High penalty for clear contamination
        
        return 0

    def _extract_vessel_details(self, text: str, vessel_name: str, doc: Document) -> Dict[str, Any]:
        """Extract vessel details using simplified structure"""
        sections = self._split_into_sections(text, doc)
        
        # Enhance sections with their associated tables
        enhanced_sections = self._enhance_sections_with_tables(sections, text)
        
        vessel_data = {
            "vessel_details": {"structured": []},
            "contract_details": {"structured": []}
        }
        
        print(f"Processing {len(enhanced_sections)} sections for {vessel_name}")
        
        for i, section in enumerate(enhanced_sections):
            if not section.strip():
                continue
            
            # Detect section and subsection using two-tier detection
            section_type, subsection, detection_method = self.detector.detect_section_integrated(section)
            
            # Create heading (first 100 chars or full content if shorter)
            heading = section[:100] + "..." if len(section) > 100 else section
            
            # Create clean chunk
            chunk = self.detector.create_clean_chunk(vessel_name, section_type, subsection, section, heading, detection_method)
            
            # Add metadata for fixture recap chunks
            if chunk["section"] == "fixture_recap" and chunk["subsection"] != "complete":
                chunk["metadata"] = {
                    "is_complete": False,
                    "chunk_type": "individual_clause",
                    "priority": "medium"
                }
            
            print(f"  Pattern detected: {section_type}/{subsection}")
            
            # Organize into output structure
            if chunk["section"] == "vessel_specifications":
                vessel_data["vessel_details"]["structured"].append(chunk)
            elif chunk["section"] in ["fixture_recap", "charter_party", "addendum"]:
                vessel_data["contract_details"]["structured"].append(chunk)
            elif chunk["section"] == "other":
                # Create individual chunks instead of merging content
                content_lower = section.lower()
                if any(term in content_lower for term in ["delivery", "figures", "laycan"]):
                    # Create individual delivery chunk
                    delivery_chunk = self.detector.create_clean_chunk(vessel_name, "delivery_figures", "delivery_info", section, heading, detection_method)
                    vessel_data["contract_details"]["structured"].append(delivery_chunk)
                elif any(term in content_lower for term in ["speed", "consumption", "fuel"]):
                    # Create individual speed/consumption chunk
                    speed_chunk = self.detector.create_clean_chunk(vessel_name, "speed_and_consumption", "performance", section, heading, detection_method)
                    vessel_data["contract_details"]["structured"].append(speed_chunk)
                elif any(term in content_lower for term in ["lifting", "equipment", "crane"]):
                    # Create individual lifting equipment chunk
                    lifting_chunk = self.detector.create_clean_chunk(vessel_name, "lifting_equipment", "equipment_info", section, heading, detection_method)
                    vessel_data["contract_details"]["structured"].append(lifting_chunk)
                elif any(term in content_lower for term in ["hseq", "documentation", "certificates"]):
                    # Create individual HSEQ chunk
                    hseq_chunk = self.detector.create_clean_chunk(vessel_name, "hseq_documentation", "documentation", section, heading, detection_method)
                    vessel_data["contract_details"]["structured"].append(hseq_chunk)
                else:
                    # Put unclassified content in contract details as individual chunks
                    vessel_data["contract_details"]["structured"].append(chunk)
        
        # Create consolidated fixture recap chunk
        fixture_recap_chunks = [chunk for chunk in vessel_data["contract_details"]["structured"] 
                               if chunk.get("section") == "fixture_recap"]
        
        if fixture_recap_chunks:
            # Combine all fixture recap content
            combined_content = []
            for chunk in fixture_recap_chunks:
                combined_content.append(chunk.get("content", ""))
            
            # Create consolidated chunk
            consolidated_chunk = {
                "vessel": vessel_name,
                "section": "fixture_recap", 
                "subsection": "complete",
                "content": "\n\n".join(combined_content),
                "heading": "Complete Fixture Recap",
                "metadata": {
                    "is_complete": True,
                    "chunk_type": "complete_section",
                    "priority": "high"
                }
            }
            
            # Add it to the beginning of contract details
            vessel_data["contract_details"]["structured"].insert(0, consolidated_chunk)
        
        return vessel_data

    def process_files(self) -> Dict[str, Dict[str, Any]]:
        """Process all DOCX files in the source directory"""
        vessels_data = {}
        
        if not os.path.exists(self.source_dir):
            logging.error(f"Source directory {self.source_dir} does not exist")
            return vessels_data
        
        # Filter for DOCX files
        docx_files = [f for f in os.listdir(self.source_dir) 
                     if f.lower().endswith('.docx') and not f.startswith('~$')]
        logging.info(f"Found {len(docx_files)} DOCX files to process: {docx_files}")
        
        for filename in docx_files:
            file_path = os.path.join(self.source_dir, filename)
            logging.info(f"Processing {filename}")
            
            try:
                # Extract text and document
                full_text = self._extract_text_from_docx(file_path)
                doc = Document(file_path)
                
                vessel_name = self._extract_vessel_name(full_text, filename)
                
                if vessel_name:
                    vessel_details = self._extract_vessel_details(full_text, vessel_name, doc)
                    vessels_data[vessel_name] = vessel_details
                    logging.info(f"Successfully extracted data for {vessel_name}")
                else:
                    logging.warning(f"Could not extract vessel name from {filename}")
            except Exception as e:
                logging.error(f"Error processing {filename}: {str(e)}")
        
        logging.info(f"Processing complete. Extracted data for {len(vessels_data)} vessels: {list(vessels_data.keys())}")
        return vessels_data

    def save_outputs(self, vessels_data: Dict[str, Dict[str, Any]], output_dir: str = 'output/vessels'):
        """Save processed vessel data to JSON and markdown files"""
        os.makedirs(output_dir, exist_ok=True)
        
        for vessel_name, vessel_data in vessels_data.items():
            vessel_dir = os.path.join(output_dir, vessel_name)
            os.makedirs(vessel_dir, exist_ok=True)
            
            # Save JSON
            json_file = os.path.join(vessel_dir, f"{vessel_name}_data.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(vessel_data, f, indent=2, ensure_ascii=False)
            
            # Save markdown
            md_file = os.path.join(vessel_dir, f"{vessel_name}_data.md")
            self._save_markdown(vessel_data, md_file)
            
            logging.info(f"Saved clean data for {vessel_name}")

    def _save_markdown(self, vessel_data: Dict[str, Any], md_file: str):
        """Save vessel data as markdown"""
        with open(md_file, 'w', encoding='utf-8') as f:
            # Extract vessel name from first vessel spec if available
            vessel_name = "Vessel"
            if vessel_data.get('vessel_details', {}).get('structured'):
                first_spec = vessel_data['vessel_details']['structured'][0]
                vessel_name = first_spec.get('vessel', 'Vessel')
            
            f.write(f"# {vessel_name} Data\n\n")
            
            # Vessel Details
            if vessel_data.get('vessel_details', {}).get('structured'):
                f.write("## Vessel Details\n\n")
                for chunk in vessel_data['vessel_details']['structured']:
                    f.write(f"### {chunk.get('subsection', 'General').replace('_', ' ').title()}\n\n")
                    f.write(f"{chunk.get('content', '')}\n\n")
            
            # Contract Details
            if vessel_data.get('contract_details', {}).get('structured'):
                f.write("## Contract Details\n\n")
                current_section = None
                for chunk in vessel_data['contract_details']['structured']:
                    if chunk.get('section') != current_section:
                        current_section = chunk.get('section')
                        f.write(f"### {current_section.replace('_', ' ').title()}\n\n")
                    
                    f.write(f"#### {chunk.get('subsection', 'General').replace('_', ' ').title()}\n\n")
                    f.write(f"{chunk.get('content', '')}\n\n")
            
            # Other sections
            for section in ['delivery_figures', 'speed_and_consumption', 'lifting_equipment', 'hseq_documentation']:
                if vessel_data.get(section, {}).get('content'):
                    f.write(f"## {section.replace('_', ' ').title()}\n\n")
                    f.write(vessel_data[section]['content'])
                    f.write("\n\n")

def main():
    """Main function to run the simplified vessel data extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Simplified Vessel Data Extractor with Two-Tier Detection')
    parser.add_argument('--source', default='source/vessels', help='Source directory containing vessel DOCX files')
    parser.add_argument('--output', default='output/vessels', help='Output directory for processed data')
    args = parser.parse_args()
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('vessel_extraction_simple.log'),
            logging.StreamHandler()
        ]
    )
    
    # Create extractor and process files
    extractor = SimpleVesselDataExtractor(args.source)
    vessels_data = extractor.process_files()
    
    if vessels_data:
        extractor.save_outputs(vessels_data, args.output)
        print(f"\nSuccessfully processed {len(vessels_data)} vessels with two-tier detection!")
        print(f"Output saved to: {args.output}")
        print(f"Ready for simplified AI search!")
    else:
        print("No vessels were processed")

if __name__ == "__main__":
    main() 